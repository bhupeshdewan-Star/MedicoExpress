from pathlib import Path
from docx import Document
import json
ROOT=Path(__file__).resolve().parents[1];OUT=ROOT/"artifacts"/"product-monograph-qualification"
required=["Executive Clinical Summary","Therapeutic Gap and Clinical Context","Product Identity and Molecule Profile","Approved Indication and Label Boundaries","Mechanism of Action","Pharmacology, PK, and Practical Use","Dosage and Administration Boundary","Clinical Evidence Architecture","Safety, Tolerability, and Monitoring","Place in Therapy and Patient Selection","Fair Comparative Framing","Claim Substantiation Matrix","Required Data Gaps Before Approval","References and Source Register","Review and Approval Status"]
def check(path):
 d=Document(path);text="\n".join([p.text for p in d.paragraphs]+[c.text for t in d.tables for r in t.rows for c in r.cells]);checks=[]
 for x in required:checks.append((x,x.lower() in text.lower()))
 for x in ["not approved for prescribing","External distribution","Prohibited","current India","https://","data gaps"]:checks.append((x,x.lower() in text.lower()))
 for x in ["completely safe","no side effects","guaranteed","preferred over inhaled corticosteroid"]:checks.append(("forbidden:"+x,x.lower() not in text.lower()))
 fails=[x for x,v in checks if not v]
 return {"file":path.name,"checks":len(checks),"passed":len(checks)-len(fails),"failures":fails,"result":"CONDITIONAL_PASS" if not fails else "FAIL","visualQA":"UNVERIFIED"}
res=[check(p) for p in OUT.glob("*_Product_Monograph.docx")];(OUT/"automated_validation.json").write_text(json.dumps(res,indent=2),encoding="utf-8")
for x in res:print(json.dumps(x))
