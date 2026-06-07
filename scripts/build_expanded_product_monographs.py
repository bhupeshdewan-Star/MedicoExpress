from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "product-monograph-expanded"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "18324A"
TEAL = "137C7B"
INK = "24313A"
MUTED = "5D6A73"
PALE = "E8F2F1"
LIGHT = "F4F6F7"
AMBER = "F4E6B9"
RUST = "9A553A"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=10.5, bold=False, italic=False, color=INK, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_lead=None, color=INK, after=7, line=1.18):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        style_run(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_lead):])
        style_run(r, color=color)
    else:
        r = p.add_run(text)
        style_run(r, color=color)
    return p


def add_hook(doc, question, evidence, meaning, boundary):
    text = f"Scientific interpretation: {evidence} {meaning} {boundary}"
    return add_paragraph(doc, text)


def add_callout(doc, title, body, fill=PALE, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    style_run(r, 10.5, True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(body)
    style_run(r, 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None, header_fill=NAVY, font_size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    repeat_table_header(table.rows[0])
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        style_run(r, font_size, True, color=WHITE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else LIGHT)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            style_run(r, font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        style_run(p.add_run(item), 9.8)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(text), 16 if level == 1 else 12.5 if level == 2 else 11, True, color=NAVY if level == 1 else TEAL)
    return p


def chapter(doc, number, title, promise, paragraphs, bullets=None, table=None, page_break=True):
    add_heading(doc, f"{number}. {title}", 1)
    add_callout(doc, "Why this chapter matters", promise)
    for hook, evidence, meaning, boundary in paragraphs:
        add_hook(doc, hook, evidence, meaning, boundary)
    if bullets:
        add_bullets(doc, bullets)
    if table:
        add_table(doc, *table)
    if page_break:
        doc.add_page_break()


def setup_document(product):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 12.5, TEAL, 12, 6),
        ("Heading 3", 11, NAVY, 9, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run(f"ClinCommand OS | {product} | Expanded HCP Scientific Monograph"), 7.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run("Controlled qualification draft | Dr. Bhupesh Dewan | Not for prescribing or external distribution"), 7.2, color=MUTED)
    return doc


COMMON_CONTROL = (
    "This document is an expanded scientific qualification draft. It is designed to support medical review and HCP knowledge development, "
    "but it is not approved prescribing information and must not be used for prescribing, promotion, or external distribution until the current "
    "India-approved package insert, sponsor product particulars, and required medical, regulatory, pharmacovigilance, legal, and compliance approvals are incorporated."
)


TRELAGLIPTIN_REFS = [
    ("1", "PMDA. Zafatek (trelagliptin succinate) review report.", "https://www.pmda.go.jp/files/000213963.pdf"),
    ("2", "PMDA. Trelagliptin renal-impairment and 25-mg strength review.", "https://www.pmda.go.jp/files/000231402.pdf"),
    ("3", "PMDA. Current Zafatek patient guide.", "https://www.info.pmda.go.jp/downfiles/ph/GUI/470310_3969024F1028_2_00G.pdf"),
    ("4", "McKeage K. Trelagliptin: First Global Approval. Drugs. 2015;75:1161-1164. PMID 26115728.", "https://pubmed.ncbi.nlm.nih.gov/26115728/"),
    ("5", "Inagaki N, et al. Once-weekly trelagliptin versus daily alogliptin. Lancet Diabetes Endocrinol. 2015;3:191-197. PMID 25609193.", "https://pubmed.ncbi.nlm.nih.gov/25609193/"),
    ("6", "Inagaki N, et al. Phase 2 randomized dose-ranging study of SYR-472. Diabetes Obes Metab. PMID 24622716.", "https://pubmed.ncbi.nlm.nih.gov/24622716/"),
    ("7", "Inagaki N, et al. Long-term monotherapy and combination therapy study. PMID 27181699.", "https://pubmed.ncbi.nlm.nih.gov/27181699/"),
    ("8", "Kaku K, et al. Trelagliptin add-on to insulin, randomized phase IV study. PMCID PMC6175153.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC6175153/"),
    ("9", "Inagaki N, et al. Trelagliptin in severe renal impairment or ESRD. PMCID PMC7078116.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC7078116/"),
    ("10", "Inagaki N, et al. Switching from daily DPP-4 inhibitor to weekly trelagliptin. PMCID PMC5835476.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC5835476/"),
    ("11", "Oita M, et al. Satisfaction after switching to weekly trelagliptin. PMID 29093280.", "https://pubmed.ncbi.nlm.nih.gov/29093280/"),
    ("12", "Ishii H, et al. Randomized QOL and treatment-satisfaction study. PMCID PMC6612345.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC6612345/"),
    ("13", "Ida S, et al. Vascular endothelial function pilot study. PMCID PMC5096292.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC5096292/"),
    ("14", "Dewan B, et al. Trelagliptin versus vildagliptin in Indian patients. PMCID PMC12215575.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC12215575/"),
    ("15", "DUET-beta study: dulaglutide versus trelagliptin. PMCID PMC11291836.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC11291836/"),
    ("16", "Kaku K. Safety evaluation of trelagliptin. PMID 28829213.", "https://pubmed.ncbi.nlm.nih.gov/28829213/"),
    ("17", "PBPK-DPP-4 occupancy model in renal impairment. Biomed Pharmacother. 2022;153:113509.", "https://www.sciencedirect.com/science/article/pii/S0753332222008988"),
    ("18", "CDSCO. Approved New Drugs portal.", "https://cdsco.gov.in/opencms/opencms/en/Approval_new/Approved-New-Drugs"),
]


SERATRODAST_REFS = [
    ("1", "CDSCO. Approved-drug list: seratrodast 40 mg and 80 mg, add-on therapy in bronchial asthma.", "https://cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadApprovalMarketingFDC/LIST-OF-APPROVED-DRUG-2012%281%29.pdf"),
    ("2", "CDSCO Pulmonary NDAC recommendations for seratrodast.", "https://cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadCommitteeFiles/NDACPulmonary_10.pdf"),
    ("3", "PMDA. Seratrodast hepatic safety bulletin.", "https://www.pmda.go.jp/safety/info-services/drugs/calling-attention/safety-info/0148.html"),
    ("4", "Japanese electronic package insert mirror: Bronica.", "https://shirobon.net/imgview.php?param1=4490018F2020&param2=PDF"),
    ("5", "Samara E, et al. Population PK/PD analysis. Clin Pharmacol Ther. 1997;62:426-435. PMID 9357394.", "https://pubmed.ncbi.nlm.nih.gov/9357394/"),
    ("6", "Tamaoki J, et al. Effect on sputum production and physicochemical properties. Chest. 2000;118:73-79. PMID 10893362.", "https://pubmed.ncbi.nlm.nih.gov/10893362/"),
    ("7", "Aizawa H, et al. Airway hyperresponsiveness, exhaled NO, and sputum eosinophils. PMID 9844991.", "https://pubmed.ncbi.nlm.nih.gov/9844991/"),
    ("8", "Muramatsu H, et al. Clinical effects and 11-dehydrothromboxane B2. PMID 11517517.", "https://pubmed.ncbi.nlm.nih.gov/11517517/"),
    ("9", "Dewan B, et al. Seratrodast versus montelukast in Indian adults with asthma.", "https://imsear.searo.who.int/items/3683f1dc-4268-464d-9b58-424ae89f4cd3"),
    ("10", "Open-access full text: Indian seratrodast versus montelukast study.", "https://imsear.searo.who.int/bitstream/123456789/182422/1/bjmmr2016v12n12p1-13.pdf"),
    ("11", "Discovery of seratrodast as a thromboxane A2 antagonist. J-STAGE.", "https://www.jstage.jst.go.jp/article/yakushi1947/119/5/119_5_377/_article/-char/en"),
    ("12", "Seratrodast inhibition of eosinophil activation by prostaglandin D2.", "https://karger.com/iaa/article/124/1-3/365/171217/Eosinophil-Activation-by-Prostaglandin-D2-and-Its"),
    ("13", "Seratrodast inhibits ferroptosis in experimental systems. PMCID PMC11584883.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC11584883/"),
    ("14", "GINA 2026 Strategy Report.", "https://ginasthma.org/wp-content/uploads/2026/05/GINA-2026-Strategy-Report-WMS.pdf"),
]


def front_matter(doc, product, subtitle, indication, key_boundary):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(85)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("CLINCOMMAND OS"), 11, True, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    style_run(p.add_run(product), 30, True, color=NAVY, font="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    style_run(p.add_run("Expanded HCP Scientific Product Monograph"), 16, color=TEAL, font="Aptos Display")
    add_callout(doc, "Scientific purpose", subtitle, fill=PALE)
    add_table(doc, ["Control field", "Value"], [
        ("Document status", "Expanded qualification draft; not approved for prescribing, promotion, or external distribution"),
        ("Intended audience", "Medical, regulatory, pharmacovigilance, compliance reviewers and qualified HCP-content reviewers"),
        ("Owner", "Dr. Bhupesh Dewan, Mumbai, Maharashtra, India"),
        ("Evidence cut-off", "6 June 2026"),
        ("India indication context", indication),
        ("Critical boundary", key_boundary),
        ("Controlled sources", "Owner-supplied Product Monograph SOP and World-Class HCP Product Monograph Skill consulted"),
    ], widths=[1.65, 4.6], header_fill=TEAL, font_size=8.5)
    add_callout(doc, "Mandatory use boundary", COMMON_CONTROL, fill=AMBER, accent=RUST)
    doc.add_page_break()
    add_heading(doc, "Document Control, Review, and Disclosure", 1)
    add_paragraph(doc, "This monograph deliberately separates established evidence, jurisdiction-specific regulatory information, clinical interpretation, and unresolved gaps. International evidence is included to support scientific understanding; it must not be converted into an India-facing prescribing or promotional claim unless consistent with the current India-approved package insert.")
    add_table(doc, ["Review domain", "Required disposition before release"], [
        ("Medical accuracy", "Independent verification of every material scientific statement, numerical result, and interpretation"),
        ("Regulatory", "Confirmation of current India indication, posology, contraindications, warnings, interactions, and legal requirements"),
        ("Pharmacovigilance", "Reconciliation with current safety information, post-marketing findings, and risk-management requirements"),
        ("Legal and compliance", "Assessment against UCPMP and company policies; approval of intended-use classification"),
        ("Visual quality", "Page-by-page rendered inspection of tables, references, warnings, and navigation"),
    ], widths=[1.45, 4.8], font_size=8.5)
    doc.add_page_break()


def executive_summary(doc, product, paragraphs, quick_rows, limitations):
    add_heading(doc, "Executive Scientific Summary", 1)
    add_callout(doc, "Clinical focus", f"What should an HCP understand about {product} before deciding whether it is relevant for an individual patient?", fill=PALE)
    for row in paragraphs:
        add_hook(doc, *row)
    add_heading(doc, "Executive quick reference", 2)
    add_table(doc, ["Decision dimension", "Evidence-based summary"], quick_rows, widths=[1.7, 4.55], header_fill=TEAL, font_size=8.5)
    add_heading(doc, "What remains uncertain", 2)
    add_bullets(doc, limitations)
    add_callout(doc, "Executive conclusion", f"{product} should be considered only within its locally approved role, after matching the patient, treatment objective, safety requirements, and evidence limitations. Convenience or mechanistic plausibility must never replace outcome evidence or the current approved label.", fill=AMBER, accent=RUST)
    doc.add_page_break()


def study_profile(doc, number, study):
    add_heading(doc, f"{number}. Study Profile: {study['title']}", 1)
    add_callout(doc, "Clinical focus", study["question"])
    add_table(doc, ["Study element", "Detail"], study["details"], widths=[1.65, 4.6], header_fill=TEAL, font_size=8.2)
    for hook, evidence, meaning, boundary in study["analysis"]:
        add_hook(doc, hook, evidence, meaning, boundary)
    add_callout(doc, "Evidence interpretation", study["interpretation"], fill=AMBER, accent=RUST)
    doc.add_page_break()


def add_source_register(doc, refs):
    add_heading(doc, "Reference and Open-Web Source Register", 1)
    add_callout(doc, "Search method", "Sources were discovered through PubMed, PubMed Central, regulator portals, trial-linked articles, open-access journal pages, and open-web citation chaining. Google Scholar is treated only as a discovery route; original publications and regulator records remain the evidence source.")
    add_paragraph(doc, "Reference handling follows a bibliography-first approach. Each line in the source register is written as a controlled citation with a stable locator so the review team can trace the origin quickly, but the narrative sections are still responsible for summarising what the evidence actually means. This keeps the document readable for HCP review while preserving the audit trail needed for scientific, regulatory, and legal checking.")
    add_paragraph(doc, "The register also acts as a control checkpoint for source quality. Regulator documents, sponsor-hosted product pages, package inserts, and peer-reviewed full text are not interchangeable, even when they discuss the same molecule. The review team should therefore check whether a citation is being used to support jurisdiction-specific labelling, mechanistic background, comparative efficacy, safety, or simply historical context. This distinction is especially important when a product monograph is being prepared before the current India-approved package insert has been secured, because the same English sentence can be scientifically true in one jurisdiction and legally unusable in another.")
    add_table(doc, ["No.", "Vancouver-style reference", "Stable URL"], refs, widths=[0.35, 3.65, 2.25], header_fill=TEAL, font_size=7.3)
    doc.add_page_break()


def add_claim_matrix(doc, product, claims):
    add_heading(doc, "Claim-Source and Fair-Balance Matrix", 1)
    add_paragraph(doc, f"The following matrix controls high-value statements about {product}. It is an internal review tool and does not confer promotional approval.")
    add_table(doc, ["ID", "Controlled statement", "Evidence anchor", "Boundary / fair balance", "Status"], claims,
              widths=[0.45, 1.8, 1.35, 1.9, 0.75], header_fill=NAVY, font_size=7.2)
    doc.add_page_break()


def add_prescribing_placeholder(doc, product, gaps):
    add_heading(doc, "Part 2: India Prescribing Information - Controlled Placeholder", 1)
    add_callout(doc, "Release blocker", f"The current authority-approved India package insert for {product} was not supplied or retrieved in a form suitable for verbatim incorporation. No reconstructed or inferred prescribing information is presented as authoritative.", fill=AMBER, accent=RUST)
    add_heading(doc, "Required approved content", 2)
    add_bullets(doc, [
        "Brand and active ingredient presentation, dosage form, strength, route, pack, storage, legal category, MAH, and company address.",
        "Exact approved indication and patient population.",
        "Exact dosage, administration, missed-dose instructions, adjustment rules, and treatment duration.",
        "Contraindications, serious warnings, precautions, interactions, adverse reactions, overdose, and special-population language.",
        "Current approval, revision, and preparation dates plus further-information statement.",
    ])
    add_heading(doc, "Unresolved product-specific fields", 2)
    add_bullets(doc, gaps)
    add_callout(doc, "Control rule", "The approved India prescribing information must be inserted without scientific rewriting, then independently checked against the source document before this monograph can progress beyond qualification draft status.")
    doc.add_page_break()


def add_evidence_methodology_annex(doc, product):
    add_heading(doc, "Appendix: Evidence-Search and Interpretation Methodology", 1)
    add_callout(doc, "Why this appendix matters", f"A detailed monograph on {product} is only as trustworthy as the process used to find, verify, interpret, and update its evidence.")
    methods = [
        ("How was evidence discovered? ", "The search pathway combined regulator portals, PubMed, PubMed Central, open-access journal repositories, trial registries, citation chaining, and broader open-web discovery. Google Scholar can help locate related papers and cited-by networks, but every result must be verified against the original journal, DOI, registry, or regulator record. ", "This layered approach reduces the risk that an important study is missed simply because it is indexed inconsistently. ", "The present qualification run is extensive but is not represented as a formally registered systematic review with duplicate independent screening."),
        ("Why is PubMed Central searched separately? ", "PubMed identifies indexed records, while PubMed Central provides accessible full text for many clinically important articles. Full text permits verification of eligibility criteria, analysis populations, missing-data methods, confidence intervals, adverse events, and limitations that are often absent from abstracts. ", "This is especially important before making comparisons or quoting numerical results. ", "Open access does not itself establish quality; study design and relevance remain decisive."),
        ("How are regulatory sources prioritized? ", "The current authority-approved local package insert is the controlling source for indication, dosage, contraindications, warnings, and mandatory safety language. Other regulator documents can provide development history, review reasoning, pharmacology, and safety context. ", "Jurisdiction separation prevents a scientifically correct foreign-label statement from becoming an incorrect local prescribing claim. ", "Where the current India package insert is unavailable, the monograph must remain a controlled draft."),
        ("How are study designs ranked? ", "Randomized controlled pivotal trials generally provide stronger evidence for causal efficacy than uncontrolled studies, while long-term extensions add exposure information and observational studies add real-world context. Mechanistic, biomarker, and preclinical studies answer different questions and are labelled accordingly. ", "The hierarchy helps readers understand how much confidence to place in each finding. ", "No study is upgraded merely because its result is favorable."),
        ("How are comparisons controlled? ", "A defensible comparison identifies the comparator, population, background therapy, duration, endpoint, analysis method, absolute result, confidence interval, and limitations. Head-to-head randomized evidence is treated differently from cross-trial context. ", "This preserves clinical meaning and prevents hanging comparisons. ", "A statistically significant difference is not automatically clinically important."),
        ("How are non-inferiority findings interpreted? ", "Non-inferiority depends on a prespecified margin, appropriate analysis populations, assay sensitivity, adherence to protocol, and confidence intervals. The conclusion applies to the studied endpoint and conditions. ", "It can support that a product was not unacceptably worse than its comparator within the margin. ", "It does not establish superiority, equivalence for all outcomes, or interchangeability in every patient."),
        ("How are safety data integrated? ", "Safety assessment combines trial adverse events, regulator reviews, post-marketing signals, label changes, pharmacovigilance communications, and clinically relevant class context. Denominators, exposure duration, seriousness, expectedness, and causality limitations are retained. ", "This makes the safety narrative usable without falsely implying certainty. ", "Absence of a signal in a small trial cannot establish absence of risk."),
        ("How are negative and neutral findings handled? ", "A complete monograph includes clinically important findings that did not show benefit, such as neutral pulmonary-function results, absent biomarker changes, or lack of demonstrated hard outcomes. ", "Neutral findings define what the medicine should not be expected to do and protect against overinterpretation. ", "Selective omission of unfavorable evidence would make the monograph scientifically misleading."),
        ("How is clinical relevance separated from statistical significance? ", "P-values and confidence intervals describe statistical evidence, while clinical relevance depends on effect magnitude, baseline risk, endpoint importance, duration, patient burden, and alternatives. ", "The monograph therefore explains what a numerical result could mean in practice. ", "Small statistically significant changes should not be transformed into broad treatment claims."),
        ("How are patient-reported outcomes interpreted? ", "Treatment satisfaction, quality of life, symptom scores, and convenience measures can capture outcomes that laboratory values miss. They are also sensitive to expectations, open-label design, instrument selection, and individual preferences. ", "These outcomes can support shared decision-making when described precisely. ", "They do not automatically prove adherence, persistence, or improved clinical outcomes."),
        ("How are evidence gaps identified? ", "The search actively looks for outcomes and populations that are absent as well as those that are present. Examples include cardiovascular outcomes, exacerbations, pregnancy, pediatrics, severe organ impairment, long-term India evidence, and uncommon adverse events. ", "Visible gaps prevent model-generated assumptions from filling missing knowledge. ", "A gap should trigger caution or future research, not a stronger claim."),
        ("How should this monograph be updated? ", "A pre-approval update search should review regulator changes, new trials, corrections, retractions, safety communications, guidelines, and sponsor pharmacovigilance information. Material changes should update the claim-source matrix and trigger relevant reviewers. ", "A living monograph remains current without losing traceability. ", "No update becomes effective without controlled review and approval."),
    ]
    for item in methods:
        add_hook(doc, *item)
    add_paragraph(doc, "The reference set is reviewed a second time after drafting so that the bibliography, the claim matrix, and the prescribing-information placeholder stay in step. That extra pass matters because small label or citation mismatches can quietly distort the final scientific meaning of a monograph even when the document looks complete on the page.")
    add_paragraph(doc, "A final interpretive sweep checks whether the open-web source set is balanced between regulator records, company disclosures, and independent peer-reviewed literature. Where a sponsor-hosted page is used, the text should be treated as useful supporting material rather than the last word on approved prescribing status. Where an open-access journal article is used, the analysis should still identify the study design, comparator, population, follow-up, and the limits of extrapolation. This layered handling of references is what prevents a polished monograph from drifting into an overconfident brochure.")
    add_paragraph(doc, "The same discipline applies to updates. A monograph should be re-opened whenever a regulator changes the label, a sponsor revises the package insert, a new pivotal study shifts the benefit-risk conversation, or a safety communication adds a meaningful warning. At that point, the task is not merely to append new citations. The reviewer must decide whether the claim matrix, the prescribing placeholder, the HCP question set, and the practical-use guidance all still tell a consistent story. That is the difference between a document that merely looks finished and a document that is genuinely controlled.")
    add_heading(doc, "Evidence hierarchy used in this qualification draft", 2)
    add_table(doc, ["Evidence level", "Typical source", "Permitted use", "Key limitation"], [
        ("1", "Current India-approved package insert and CDSCO records", "Local indication, dosing, warnings, legal prescribing content", "Current package insert not yet available"),
        ("2", "Regulator review reports and safety communications", "Development reasoning, PK, safety signals, jurisdiction-specific context", "May not reflect India label"),
        ("3", "Randomized controlled clinical trials", "Causal efficacy and comparative safety within studied conditions", "May be short or population-specific"),
        ("4", "Long-term extensions and observational studies", "Durability, exposure, practice context, hypotheses", "Confounding and weaker causal inference"),
        ("5", "Systematic reviews and meta-analyses", "Synthesis across eligible studies", "Limited by included evidence and heterogeneity"),
        ("6", "Mechanistic, biomarker, modeling, and preclinical evidence", "Biological rationale and future research", "Cannot establish patient-important clinical benefit"),
    ], widths=[0.65, 1.8, 2.2, 1.6], font_size=7.6)
    doc.add_page_break()


def add_hcp_faq_annex(doc, product, profile):
    add_heading(doc, "Appendix: Detailed HCP Scientific Questions and Answers", 1)
    add_callout(doc, "Purpose", f"These questions are designed to make the scientific evidence on {product} easier to retrieve during medical review. They are not prescribing instructions or promotional rebuttals.")
    questions = [
        ("What is the most defensible one-sentence description of the product?", profile["one_line"]),
        ("What is the locally verified therapeutic role?", profile["role"]),
        ("What is the molecule's principal scientific distinction?", profile["distinction"]),
        ("What is the strongest direct efficacy evidence?", profile["efficacy"]),
        ("What is the most important India-specific evidence?", profile["india"]),
        ("What is the most important international evidence?", profile["international"]),
        ("What is the most important safety issue?", profile["safety"]),
        ("What is the most important monitoring issue?", profile["monitoring"]),
        ("What should never be claimed from the available evidence?", profile["never_claim"]),
        ("How should mechanism of action be explained to an HCP?", profile["mechanism"]),
        ("What does the pharmacokinetic profile mean clinically?", profile["pk"]),
        ("Which evidence is pivotal and which is supportive?", profile["hierarchy"]),
        ("How should comparative findings be communicated?", profile["comparison"]),
        ("What are the central limitations of the evidence base?", profile["limitations"]),
        ("What is known about long-term use?", profile["long_term"]),
        ("What is known about use in older adults?", profile["elderly"]),
        ("What is known about renal impairment?", profile["renal"]),
        ("What is known about hepatic impairment?", profile["hepatic"]),
        ("What is known about pregnancy and lactation?", profile["pregnancy"]),
        ("What is known about pediatric use?", profile["pediatric"]),
        ("How should concomitant therapy be reviewed?", profile["concomitant"]),
        ("How should a patient be selected?", profile["selection"]),
        ("Which patients may require another strategy?", profile["alternative"]),
        ("What should be assessed before treatment?", profile["before"]),
        ("What should be reviewed after treatment begins?", profile["after"]),
        ("What should the patient be counselled about?", profile["counselling"]),
        ("What does the evidence say about quality of life or convenience?", profile["qol"]),
        ("Are hard clinical outcomes established?", profile["outcomes"]),
        ("How should statistically significant findings be interpreted?", profile["statistics"]),
        ("How should neutral findings be presented?", profile["neutral"]),
        ("What role does pharmacovigilance play?", profile["pv"]),
        ("How should guideline context be discussed?", profile["guidelines"]),
        ("Can foreign-label information be used?", profile["foreign"]),
        ("What must be obtained before external use?", profile["external"]),
        ("What should a reviewer verify in every study table?", profile["study_table"]),
        ("What future evidence would most improve confidence?", profile["future"]),
        ("How should new evidence be added to the monograph?", profile["update"]),
        ("What is the correct conclusion for a busy HCP?", profile["conclusion"]),
    ]
    for i, (question, answer) in enumerate(questions, 1):
        add_heading(doc, f"Q{i}. {question}", 2)
        add_paragraph(doc, answer)
        add_paragraph(doc, f"Control boundary: This answer summarizes the current evidence cut-off for {product}. It must be reconciled with the current India-approved package insert and reviewed for the intended use before distribution.", color=MUTED, after=9)
        if i % 5 == 0:
            doc.add_page_break()
    doc.add_page_break()


def add_scientific_glossary(doc, product, extra_terms):
    add_heading(doc, "Appendix: Scientific and Clinical Interpretation Glossary", 1)
    add_callout(doc, "Purpose", f"This glossary standardizes the language used to interpret the {product} evidence base. It is intended to prevent important statistical, regulatory, and clinical terms from being used loosely.")
    terms = [
        ("Absolute effect", "The observed difference expressed in the original measurement unit, such as percentage-point HbA1c change, liters per second of peak flow, or number of patients with an event. Absolute effects are often more clinically interpretable than relative descriptions."),
        ("Active-controlled trial", "A study comparing an investigational or study medicine with another active medicine. It can answer comparative questions but must be interpreted according to design, comparator dose, background therapy, endpoint, and statistical objective."),
        ("Add-on therapy", "A treatment used in addition to an existing regimen rather than as a replacement. The term does not mean the background therapy can be stopped, reduced, or ignored unless the approved label and evidence specifically support that action."),
        ("Adverse event", "Any unfavorable medical occurrence after treatment, regardless of whether it is considered caused by the medicine. Adverse-event frequency alone does not establish causality, and comparison requires similar ascertainment and exposure."),
        ("Adverse drug reaction", "An adverse event for which a causal relationship to the medicine is considered at least reasonably possible under the applicable assessment framework. Regulatory definitions and frequency categories should be followed precisely."),
        ("Analysis population", "The group of randomized or treated participants included in a statistical analysis, such as intention-to-treat, modified intention-to-treat, per-protocol, or safety populations. Different populations can produce different estimates."),
        ("Assay sensitivity", "The ability of a trial to distinguish an effective treatment from a less effective or ineffective one. It is especially important in non-inferiority studies because an insensitive trial can make treatments appear similar."),
        ("Background therapy", "Treatment continued by participants while the study medicine or comparator is added. Background therapy shapes applicability, treatment effect, and safety, and must be stated when interpreting add-on studies."),
        ("Bias", "A systematic influence that can distort a study result. Open-label treatment, selective reporting, loss to follow-up, inadequate randomization, unblinded assessment, and post-hoc analysis are common sources requiring consideration."),
        ("Clinical significance", "The practical importance of an observed effect to patients and clinicians. Clinical significance depends on effect magnitude, endpoint importance, duration, burden, safety, and alternatives, not only on statistical significance."),
        ("Confidence interval", "A range expressing uncertainty around an estimated treatment effect. The interval helps assess precision and, in superiority or non-inferiority studies, whether the result crosses clinically or statistically important boundaries."),
        ("Controlled claim", "A statement whose exact wording, evidence source, label status, limitations, fair-balance text, review disposition, and intended use are recorded. A controlled claim is not automatically approved for promotion."),
        ("Effect modifier", "A patient or study characteristic associated with different treatment effects across subgroups. Apparent effect modification requires careful interpretation and preferably prespecified, adequately powered analysis."),
        ("Evidence cut-off", "The date through which the evidence search and safety review are considered current. New studies, regulatory changes, corrections, or safety signals after this date require an update assessment."),
        ("Exploratory endpoint", "An endpoint intended mainly to generate hypotheses rather than support definitive conclusions. Exploratory findings are useful for research planning but should not be presented with the certainty of a prespecified primary endpoint."),
        ("External validity", "The extent to which a study result can be applied beyond the studied participants, sites, treatment context, and duration. Geography, disease severity, background treatment, and eligibility criteria strongly influence generalizability."),
        ("Fair balance", "Presentation of benefit, risk, limitations, uncertainty, and alternatives with appropriate prominence and proximity. Fair balance is a scientific and compliance requirement, not merely a formatting exercise."),
        ("Hard clinical outcome", "A patient-important event such as death, hospitalization, severe exacerbation, kidney failure, or cardiovascular event. A surrogate or biomarker may correlate with risk but does not automatically prove improvement in hard outcomes."),
        ("Hazard or safety signal", "Information suggesting a new or incompletely characterized possible causal association between a medicine and an event. Signals require evaluation and do not always establish causality or incidence."),
        ("Intention-to-treat", "An analysis principle that generally includes randomized participants according to assigned treatment. It preserves randomization benefits, although handling of missing data and protocol deviations still matters."),
        ("Jurisdiction-specific label", "The approved prescribing information applicable in a particular country or region. Indications, doses, contraindications, warnings, and legal requirements can differ between jurisdictions and must not be blended."),
        ("Mechanistic plausibility", "A biologically reasonable explanation for how a medicine might produce an effect. Plausibility supports research and interpretation but cannot replace direct clinical evidence of benefit or safety."),
        ("Meta-analysis", "A statistical synthesis of results from multiple studies. Its validity depends on the quality, similarity, completeness, and reporting of included studies; pooled precision cannot correct biased or weak primary evidence."),
        ("Missing-data method", "The statistical approach used when outcome data are unavailable. Attrition and the chosen method can influence results, particularly in non-inferiority trials and studies with substantial discontinuation."),
        ("Non-inferiority", "A conclusion that a new treatment is not worse than an active comparator by more than a prespecified acceptable margin for a defined endpoint. It does not establish superiority or equality for all outcomes."),
        ("Open-label study", "A study in which participants and investigators know the assigned treatment. Open-label designs can be appropriate but are more vulnerable to expectation and assessment bias, especially for subjective outcomes."),
        ("Patient-reported outcome", "A measure reported directly by a patient, such as symptoms, treatment satisfaction, or quality of life. Validated instruments, blinding, missing data, and clinical meaning should be considered."),
        ("Per-protocol analysis", "An analysis focused on participants who sufficiently followed the protocol. It can be important in non-inferiority assessment but may introduce bias because treatment adherence and exclusions are not random."),
        ("Pharmacodynamics", "What the medicine does to the body, target, pathway, biomarker, or physiological response. Pharmacodynamic activity can support dosing rationale but is not always equivalent to clinical benefit."),
        ("Pharmacokinetics", "What the body does to the medicine through absorption, distribution, metabolism, and excretion. PK becomes clinically meaningful when connected to dose, timing, interactions, organ impairment, and safety."),
        ("Post-marketing evidence", "Safety or effectiveness information generated after approval through surveillance, reports, studies, or real-world data. It can identify uncommon risks and broader-use patterns not captured in trials."),
        ("Primary endpoint", "The main outcome prespecified to answer a trial's central question. Interpretation should consider the statistical plan, effect size, confidence interval, missing data, and whether the endpoint is clinically important."),
        ("Randomized controlled trial", "A study assigning participants to interventions by a random process. Proper randomization reduces confounding, but trial credibility also depends on allocation concealment, blinding, conduct, analysis, and reporting."),
        ("Real-world evidence", "Clinical evidence derived from routine-care data sources. It can improve generalizability and safety understanding but is vulnerable to confounding, measurement error, selection, and incomplete data."),
        ("Responder analysis", "An analysis classifying participants according to a response threshold. Threshold justification, missing-data handling, baseline risk, and absolute responder numbers are needed for meaningful interpretation."),
        ("Statistical significance", "A finding unlikely under a specified null hypothesis at a chosen threshold. It does not measure effect size, clinical importance, study quality, or the probability that a claim is true."),
        ("Surrogate endpoint", "A laboratory, biomarker, or intermediate outcome used in place of a direct patient-important outcome. Surrogates can support efficacy assessment but may not reliably predict clinical benefit in every setting."),
        ("Systematic review", "A structured evidence synthesis using a prespecified question, search, eligibility criteria, appraisal, and transparent methods. A broad literature search is valuable but should not be called systematic unless those standards are met."),
        ("Treatment-emergent adverse event", "An event beginning or worsening after treatment exposure under a study definition. Frequency comparisons require attention to exposure duration, ascertainment, baseline events, and analysis population."),
        ("Uncertainty", "The range of plausible interpretations remaining because of sampling variation, bias, limited duration, missing evidence, or generalizability concerns. A trustworthy monograph states uncertainty explicitly."),
    ] + extra_terms
    add_table(doc, ["Term", "Working definition and interpretation control"], terms, widths=[1.55, 4.7], header_fill=TEAL, font_size=7.7)
    doc.add_page_break()


def add_clinical_review_lenses(doc, product):
    add_heading(doc, "Appendix: Fifteen Clinical Review Lenses", 1)
    add_callout(doc, "How to use these lenses", f"Before using any statement from this monograph, reviewers should examine the {product} evidence through each lens below. Together they protect against attractive but incomplete interpretation.")
    lenses = [
        ("1. Patient lens", "Is the patient in front of the clinician meaningfully similar to the studied population, including disease severity, age, organ function, background therapy, and treatment objective? If not, the uncertainty should be stated rather than hidden."),
        ("2. Endpoint lens", "Does the study measure a direct patient-important outcome, a validated clinical outcome, a symptom scale, a physiological measure, or an exploratory biomarker? The wording should become more cautious as the endpoint becomes more indirect."),
        ("3. Comparator lens", "Was the comparator appropriate, correctly dosed, and relevant to current practice? A favorable result against one comparator does not establish superiority to all alternatives or to the current standard of care."),
        ("4. Duration lens", "Was follow-up long enough to assess the intended benefit and important harms? Short studies can characterize early efficacy but usually cannot establish durability, uncommon safety events, or long-term clinical outcomes."),
        ("5. Magnitude lens", "How large is the observed effect in absolute and clinically understandable terms? A small difference can be statistically convincing but of limited practical importance, while a meaningful effect can remain imprecise in a small study."),
        ("6. Precision lens", "How wide is the confidence interval, and what clinically important effects remain compatible with it? Precision should influence the certainty of conclusions and the prominence of limitations."),
        ("7. Design lens", "Was the study randomized, blinded, controlled, and prospectively analyzed? Design strengths increase confidence, while open-label, uncontrolled, retrospective, or post-hoc findings require explicit restraint."),
        ("8. Missing-data lens", "How many participants discontinued or lacked outcome data, why were data missing, and how were they handled? Missing data can materially alter both superiority and non-inferiority conclusions."),
        ("9. Safety lens", "Were adverse events actively collected, were exposure times comparable, and was the sample large enough to identify uncommon risks? Safety conclusions should integrate trials with post-marketing and regulator evidence."),
        ("10. Jurisdiction lens", "Is the statement a scientific finding, a foreign-label fact, or a current India-approved prescribing statement? These categories must remain visibly separate in wording, layout, review, and intended use."),
        ("11. Guideline lens", "How does the medicine's demonstrated evidence fit with current treatment priorities and recommended pathways? Local approval permits use but does not automatically establish guideline preference or broad treatment priority."),
        ("12. Alternative-therapy lens", "What other therapies address the same clinical problem, and which outcomes have they demonstrated? Balanced selection considers what the patient may gain or forgo with each option."),
        ("13. Practical-use lens", "Can the regimen be administered, monitored, understood, and sustained safely in the actual care setting? Convenience, complexity, caregiver support, medication errors, and monitoring burden all influence real clinical value."),
        ("14. Conflict and transparency lens", "Who funded, designed, conducted, analyzed, and published the study? Sponsorship does not invalidate evidence, but transparent disclosure and independent verification strengthen trust."),
        ("15. Update lens", "Could new safety information, label revisions, guidelines, or studies have changed the conclusion since the evidence cut-off? The most polished monograph becomes unreliable if it is not actively maintained."),
    ]
    for title, body in lenses:
        add_heading(doc, title, 2)
        add_paragraph(doc, body)
        add_paragraph(doc, f"Application to {product}: record the reviewer's conclusion and any required wording change in the claim-source matrix before approval.", color=MUTED)
    doc.add_page_break()


def add_seratrodast_hepatic_annex(doc):
    add_heading(doc, "Appendix: Detailed Hepatic-Risk Review for Seratrodast", 1)
    add_callout(doc, "Why this appendix is mandatory", "The hepatic-risk signal is the most consequential safety issue identified for seratrodast. It must be understood in detail, remain prominent, and be reconciled with the current India-approved package insert before external use.", fill=AMBER, accent=RUST)
    rows = [
        ("What was observed before approval? ", "Japanese regulatory safety information reports abnormal liver-function tests in 28 of 824 patients, or 3.4%, during preapproval clinical development. This denominator provides useful signal context, but it does not by itself establish the incidence of serious hepatic injury or predict risk in Indian practice. The result should remain linked to its jurisdiction, development period, and ascertainment method."),
        ("What changed after marketing? ", "Post-marketing experience identified serious hepatic dysfunction, jaundice, and fulminant hepatitis, including fatal cases. These observations changed the practical safety narrative from routine laboratory abnormality to a potentially life-threatening risk requiring active monitoring and clinical response. Individual reports do not prove causality in every case, but their seriousness justifies prominent risk mitigation."),
        ("How many serious cases were described in the PMDA communication? ", "The reviewed PMDA safety communication describes 49 serious hepatic-dysfunction cases during the early post-marketing period, including four cases of fulminant hepatitis. Four deaths involving fulminant hepatitis were reported within approximately two years and three months after launch, although causality was not established in every case. These numbers must be reproduced only with their source and limitations."),
        ("When can hepatic injury occur? ", "Reported onset extended across treatment periods up to approximately 33 weeks, showing that early normal tests cannot be treated as permanent reassurance. Monitoring and patient education must continue according to the approved local requirements. The time-course evidence also argues against restricting vigilance to the first few days or weeks."),
        ("How severe were reported laboratory abnormalities? ", "The PMDA communication includes cases with very high aminotransferases and bilirubin, including a reported maximum ALT/GPT of 1,760 and total bilirubin of 38.7 mg/dL. A described fatal case presented after 197 days of 80-mg daily treatment with marked aminotransferase elevation, hyperbilirubinemia, progressive coagulopathy, encephalopathy, and death. Such examples illustrate severity but do not define typical presentation."),
        ("What happened after recognition and management? ", "Among evaluable non-fulminant cases with known outcomes in the reviewed communication, most improved within approximately 12 weeks. This supports the importance of timely recognition, discontinuation when indicated, and appropriate management. It must not be interpreted as reassurance that all cases are reversible or that delayed action is acceptable."),
        ("What monitoring is stated in Japanese information? ", "Japanese product information requires liver-function testing once monthly and careful use in hepatic impairment. This clear jurisdiction-specific instruction should be visible in the scientific monograph because it informs risk understanding. Whether the same frequency and action thresholds apply in India must be verified from the current India-approved package insert."),
        ("Which symptoms deserve patient education? ", "Although final counselling language must follow the approved India label, clinicians generally need to consider symptoms compatible with hepatic dysfunction, such as unusual fatigue, anorexia, nausea, dark urine, jaundice, or right-upper-quadrant discomfort. Patient education should support prompt assessment without creating false reassurance or alarm. The monograph must not replace clinical evaluation."),
        ("How should concomitant medicines be considered? ", "A full medication review should identify other medicines, supplements, alcohol exposure, and conditions that can affect the liver or complicate causal assessment. Concomitant risk does not prove that seratrodast caused an event, but it influences patient selection, monitoring, differential diagnosis, and benefit-risk review. India-specific interaction and contraindication language remains package-insert controlled."),
        ("How should liver safety affect patient selection? ", "A credible add-on rationale must be weighed against hepatic history, baseline status, monitoring feasibility, available alternatives, and the clinical importance of the residual asthma problem. The known signal makes seratrodast unsuitable for casual or poorly monitored use. Exact exclusions and contraindications must come from the current India label."),
        ("How should the benefit-risk conversation be framed? ", "The potential benefit is supported mainly by selected symptom, airway-response, secretion, and lung-function findings, whereas the hepatic signal includes rare but serious post-marketing events. The HCP should therefore define a clear treatment objective and reassess whether observed benefit justifies continued exposure. This is a clinical decision framework, not a prescribing recommendation."),
        ("What evidence remains missing? ", "The public evidence reviewed does not provide a complete contemporary India estimate of hepatic-event incidence, validated risk factors, optimal monitoring effectiveness, or outcomes across routine use. Sponsor pharmacovigilance data, current periodic safety reports, and the approved India package insert are therefore essential before finalization. Absence of public evidence must not be interpreted as absence of risk."),
        ("How should abnormal laboratory results be interpreted? ", "An abnormal liver test requires interpretation in the context of baseline values, symptoms, timing, concomitant medicines, alcohol exposure, viral and autoimmune causes, biliary disease, and other clinical findings. The monograph should encourage timely assessment without pretending that a single value proves or excludes drug-induced liver injury. Exact interruption and discontinuation criteria must follow approved local guidance and qualified clinical judgment."),
        ("Why is regular monitoring not a guarantee of safety? ", "Periodic testing can support earlier detection, but injury may evolve between scheduled assessments and some patients may develop symptoms before the next planned test. Monitoring therefore works together with patient education, clinician vigilance, access to testing, and prompt action. A monitoring recommendation must never be written in a way that implies serious hepatic injury is fully preventable."),
        ("How should hepatic risk appear visually? ", "The risk must be visible in the executive summary, safety chapter, practical-use pathway, patient-counselling section, and claim-source matrix. It should use clear hierarchy and readable typography without sensationalism. If benefit statements are easy to find while the hepatic signal is buried in an appendix, the monograph has failed fair balance even if the words are technically present."),
        ("What is the correct final safety conclusion? ", "Seratrodast has a recognized serious hepatic-risk signal that materially affects selection, monitoring, counselling, and continuation decisions. The available evidence supports prominent caution and jurisdiction-specific verification rather than an absolute statement about incidence or predictability. External use remains blocked until current India prescribing and pharmacovigilance information is fully reconciled and approved."),
        ("What should an audit trail demonstrate? ", "The final audit trail should show which hepatic sources were consulted, which numerical details were independently verified, how current India safety information was reconciled, why each risk statement was worded as written, and who approved it. It should also record unresolved uncertainty and future update triggers. This makes the safety chapter reproducible, reviewable, and resistant to gradual dilution in later derivative materials and future field-support documents."),
    ]
    for hook, body in rows:
        add_paragraph(doc, hook + body, bold_lead=hook)
    add_table(doc, ["Hepatic-risk control", "Required action before release"], [
        ("India label verification", "Insert exact contraindication, warning, monitoring, discontinuation, and special-population language"),
        ("PV reconciliation", "Review current sponsor safety database, PSUR/PBRER, regulatory communications, and signal status"),
        ("Claim review", "Ensure every efficacy or place-in-therapy summary carries proportionate hepatic-risk context"),
        ("Visual review", "Confirm the hepatic warning is not buried, minimized, clipped, or less legible than efficacy content"),
        ("Human approval", "Obtain medical, regulatory, pharmacovigilance, legal, and compliance approval"),
    ], widths=[1.65, 4.6], header_fill=RUST, font_size=8.0)
    doc.add_page_break()


def build_trelagliptin():
    doc = setup_document("Trelagliptin")
    front_matter(
        doc, "Trelagliptin",
        "A detailed scientific reference on a long-acting DPP-4 inhibitor developed for once-weekly administration in type 2 diabetes mellitus.",
        "An India approval listing and an Indian phase 3 study were identified; exact current India package-insert wording remains to be verified.",
        "Japan-label dose and safety statements are scientific context and must not be silently transferred into India-facing prescribing information."
    )
    executive_summary(doc, "Trelagliptin", [
        ("Why does dosing frequency matter in type 2 diabetes? ", "Long-term glycaemic management often requires multiple medicines and sustained daily routines, making treatment burden a relevant clinical consideration. ", "Trelagliptin introduces a once-weekly oral DPP-4 inhibitor option whose principal distinction is administration frequency. ", "Reduced frequency is not, by itself, proof of better adherence, persistence, cardiovascular outcomes, or long-term glycaemic durability."),
        ("What does the pivotal efficacy evidence show? ", "Japanese phase II and phase III studies support glucose-lowering efficacy, including non-inferiority to daily alogliptin in the studied population; an Indian phase III study reported non-inferiority to twice-daily vildagliptin over 16 weeks. ", "These data establish a clinical evidence base for HbA1c lowering within the studied populations and durations. ", "The evidence does not establish superiority, cardiorenal benefit, mortality benefit, or equivalence across all patient groups."),
        ("What is the central safety consideration? ", "Trelagliptin's long pharmacodynamic action and predominant renal elimination make renal-function assessment, dose selection, hypoglycaemia awareness, and medication-error prevention especially important. ", "The weekly schedule changes how clinicians should think about missed doses, duplicate doses, switching, and persistence of adverse effects. ", "Exact India dosing and safety wording must come from the current approved package insert."),
        ("Where might the product fit? ", "Within a locally approved role, trelagliptin may be relevant when an oral DPP-4 inhibitor is clinically appropriate and a weekly schedule aligns with the patient's treatment plan. ", "Patient selection should reflect glycaemic goals, comorbidities, renal function, concomitant therapy, and the need for therapies with proven outcome benefits. ", "The product should not displace guideline-preferred cardiorenal therapies when those benefits are the dominant treatment objective."),
    ], [
        ("Molecule and class", "Trelagliptin succinate; selective oral DPP-4 inhibitor."),
        ("Core clinical distinction", "Sustained DPP-4 inhibition supporting once-weekly administration."),
        ("Verified international role", "Approved in Japan for type 2 diabetes mellitus; an India approval listing was identified."),
        ("Evidence base", "Dose-ranging, active-controlled, long-term, switching, insulin add-on, renal-impairment, QOL, and Indian comparative studies."),
        ("Primary evidence-supported outcome", "Reduction in HbA1c and related glycaemic measures."),
        ("Principal practical controls", "Renal assessment, correct weekly-day administration, missed-dose education, duplicate-dose prevention, and review of concomitant hypoglycaemic therapy."),
    ], [
        "Current complete India-approved prescribing information was not available for verbatim incorporation.",
        "No dedicated cardiovascular-outcomes trial was identified.",
        "Evidence for renal, cardiovascular, mortality, or microvascular outcome improvement is not established.",
        "Much of the pivotal evidence was generated in Japanese patients; generalizability requires clinical judgment.",
        "Long-term controlled evidence and broad real-world India evidence remain limited.",
    ])
    chapters = [
        ("1", "Type 2 Diabetes: The Clinical Problem Beyond HbA1c", "Understanding the disease context prevents a glucose-lowering medicine from being considered in isolation.", [
            ("Why is type 2 diabetes more than a glucose number? ", "The disease is a progressive metabolic disorder involving insulin resistance, beta-cell dysfunction, dysregulated glucagon secretion, and multi-organ complications. ", "Treatment decisions must therefore balance glycaemic control with cardiovascular, renal, weight, hypoglycaemia, and patient-burden priorities. ", "A product monograph should make clear which of these outcomes are directly supported and which are not."),
            ("What makes treatment selection increasingly individualized? ", "Contemporary care considers atherosclerotic cardiovascular disease, heart failure, chronic kidney disease, obesity, frailty, hypoglycaemia risk, cost, preferences, and regimen complexity. ", "A DPP-4 inhibitor may be suitable in selected clinical circumstances because of oral administration and glucose-dependent pharmacology. ", "Class suitability does not imply that every molecule has identical evidence or that DPP-4 inhibition addresses every treatment priority."),
            ("Where does treatment burden enter the clinical conversation? ", "Polypharmacy and repeated daily dosing can create practical friction for patients managing a chronic asymptomatic condition. ", "A weekly oral schedule may be a meaningful operational feature for some patients and care teams. ", "Its value must be assessed patient by patient and must not be converted into an unproven adherence or outcome claim."),
        ]),
        ("2", "The Incretin System and the DPP-4 Therapeutic Class", "Class context helps an HCP understand what trelagliptin can reasonably be expected to do.", [
            ("What biological problem does DPP-4 inhibition address? ", "DPP-4 rapidly degrades endogenous incretin hormones including GLP-1 and GIP, which participate in glucose-dependent insulin secretion and glucagon regulation. ", "Inhibition prolongs endogenous incretin activity and supports glycaemic control with a generally low intrinsic hypoglycaemia propensity when used without insulin or secretagogues. ", "This mechanism does not create proven cardiovascular, renal, or weight-loss benefits for trelagliptin."),
            ("How should the class be differentiated from GLP-1 receptor agonists? ", "DPP-4 inhibitors enhance endogenous incretin activity, whereas GLP-1 receptor agonists directly stimulate the receptor pharmacologically and often produce larger HbA1c and weight effects. ", "The distinction matters when matching treatment intensity and outcome priorities. ", "Cross-class comparisons must rely on comparable evidence and should not imply that convenience compensates for differences in demonstrated outcomes."),
            ("Why is within-class evidence still necessary? ", "DPP-4 inhibitors differ in dosing frequency, elimination pathways, dose-adjustment requirements, and the extent of molecule-specific outcome evidence. ", "Trelagliptin's weekly profile and renal handling make its practical controls distinct. ", "Class assumptions cannot replace molecule-specific label and safety review."),
        ]),
        ("3", "Discovery, Development, and Regulatory History", "The development pathway explains why once-weekly dosing became the molecule's defining scientific question.", [
            ("What was trelagliptin designed to achieve? ", "Trelagliptin, also known during development as SYR-472, was optimized as a potent and selective long-acting DPP-4 inhibitor. ", "Its development programme evaluated whether sustained inhibition could permit weekly oral dosing while maintaining glycaemic efficacy. ", "A longer interval increases the importance of correct dose selection and error prevention."),
            ("Where was the first approval obtained? ", "PMDA materials document Japanese approval in 2015, with subsequent regulatory work supporting a 25-mg strength for severe renal impairment and ESRD. ", "This provides detailed scientific and regulatory context. ", "Japanese approved conditions must remain visibly separate from India-approved conditions."),
            ("What is known about India? ", "An Indian phase III study and CDSCO approval-list context were identified during the evidence search. ", "This supports India relevance and mandates careful integration of local data. ", "The current regulator-approved India package insert remains the controlling document for all prescribing claims and must still be obtained."),
        ]),
        ("4", "Molecular Identity and Pharmaceutical Profile", "Molecule-level facts anchor the monograph before clinical interpretation begins.", [
            ("What exactly is the active substance? ", "Regulatory materials identify trelagliptin succinate, providing the active moiety trelagliptin. ", "Salt-to-active-equivalence and product strength presentation are essential for accurate prescribing information. ", "Final strengths, excipients, pack sizes, storage, and sponsor-specific product particulars require the current India product file."),
            ("What distinguishes the pharmacological profile? ", "The molecule demonstrates potent, highly selective DPP-4 inhibition with slow dissociation and sustained pharmacodynamic activity. ", "This supports the scientific rationale for a weekly schedule. ", "Mechanistic distinctiveness should not be expressed as clinical superiority unless supported by comparative outcomes."),
            ("Why does formulation detail matter? ", "A weekly medicine can create consequences if the wrong strength is selected or if doses are repeated too closely. ", "Clear product identity, tablet differentiation, and dispensing instructions become patient-safety controls. ", "These controls cannot be finalized without sponsor product particulars and approved packaging information."),
        ]),
        ("5", "Mechanism of Action: From DPP-4 Binding to Glycaemic Effect", "The mechanism chapter connects molecular action to realistic clinical expectations.", [
            ("How does sustained inhibition translate into weekly action? ", "Trelagliptin binds DPP-4 non-covalently and demonstrates prolonged enzyme inhibition across the dosing interval; PMDA review data describe substantial inhibition seven days after a 100-mg dose. ", "The pharmacodynamic duration supports once-weekly administration. ", "The weekly schedule should be understood as sustained target inhibition, not as permission to improvise dosing."),
            ("What happens downstream of DPP-4 inhibition? ", "Reduced incretin degradation supports glucose-dependent insulin secretion and restrains inappropriate glucagon secretion. ", "The glucose-dependent mechanism helps explain the class's generally low intrinsic hypoglycaemia tendency. ", "Hypoglycaemia can still occur, particularly when combined with insulin or insulin secretagogues."),
            ("What should not be inferred from mechanism? ", "Experimental and biomarker studies may explore endothelial, inflammatory, or neurological effects. ", "Such work can generate hypotheses about pleiotropic biology. ", "It must not be used to claim prevention of cardiovascular, renal, cognitive, or microvascular outcomes without direct clinical evidence."),
        ]),
        ("6", "Pharmacodynamics Across the Weekly Interval", "Weekly administration is credible only if pharmacodynamic activity persists predictably.", [
            ("What is the key pharmacodynamic observation? ", "Regulatory review describes sustained DPP-4 inhibition throughout the seven-day interval after therapeutic dosing. ", "This bridges the gap between plasma exposure and weekly clinical use. ", "The degree of inhibition is a surrogate pharmacodynamic measure and is not itself a patient-important outcome."),
            ("How does pharmacodynamics influence missed-dose thinking? ", "Because biological activity persists beyond the administration day, a missed or duplicate dose has different implications from a short-acting daily medicine. ", "HCPs need precise, label-aligned instructions and patient counselling. ", "No locally applicable missed-dose instruction should be reconstructed from memory or another jurisdiction."),
            ("Why can adverse effects persist after stopping? ", "Long-acting target inhibition means the medicine's pharmacological influence may not disappear immediately after discontinuation. ", "Clinical assessment should consider the previous weekly dose and timing. ", "Urgent adverse-event management must follow the approved label and clinical judgment."),
        ]),
        ("7", "Pharmacokinetics: Absorption, Distribution, Metabolism, and Excretion", "PK becomes clinically useful when numbers are translated into prescribing consequences.", [
            ("How rapidly is trelagliptin absorbed? ", "Human studies and regulatory review describe a median time to maximum concentration of approximately one to two hours, with food producing no material change in overall exposure. ", "This supports practical administration without a strong food-dependent exposure concern in the verified Japan context. ", "Exact India administration wording still requires the local package insert."),
            ("How is the medicine distributed and metabolized? ", "Protein binding is relatively low and metabolism is limited, with unchanged trelagliptin predominating; a minor active metabolite contributes little to overall exposure. ", "The profile reduces reliance on extensive hepatic metabolism. ", "It does not eliminate the need to review interactions or severe hepatic impairment."),
            ("Why is renal function central? ", "A large proportion of administered drug is excreted unchanged in urine, and exposure rises as renal function declines. ", "Renal assessment is therefore directly relevant to dose selection and safety. ", "The precise India adjustment algorithm must be verified from the approved India label."),
        ], None, (["PK dimension", "Evidence-based finding", "Clinical consequence"], [
            ("Absorption", "Tmax approximately 1-2 hours; total exposure minimally affected by food in studied conditions", "Administration practicalities should follow local label"),
            ("Distribution", "Relatively low protein binding", "Low binding does not remove interaction review"),
            ("Metabolism", "Limited; unchanged drug predominates", "Hepatic metabolism is not the primary elimination route"),
            ("Excretion", "Predominantly renal; substantial unchanged urinary excretion", "Assess renal function and apply label dose rules"),
            ("Duration", "Terminal half-life and sustained target inhibition support weekly interval", "Prevent duplicate dosing and manage missed doses precisely"),
        ], [1.25, 2.6, 2.4])),
        ("8", "Renal Impairment and Dose-Selection Logic", "Renal dosing is not a footnote for trelagliptin; it is a core safety chapter.", [
            ("What happens to exposure as renal function declines? ", "Regulatory analyses report progressively increased trelagliptin exposure across mild, moderate, severe renal impairment, and ESRD. ", "This exposure-response relationship explains the need for lower strengths in more advanced impairment. ", "The HCP must use the current locally approved thresholds rather than a generalized DPP-4 class rule."),
            ("Is dialysis expected to remove the drug effectively? ", "PMDA review data indicate only limited removal during hemodialysis. ", "This reinforces the importance of selecting the correct dose before administration. ", "Dialysis should not be assumed to reverse an administration error."),
            ("What did the severe-renal-impairment trial add? ", "A randomized phase III study evaluated 25 mg once weekly in Japanese patients with severe impairment or ESRD and demonstrated HbA1c lowering versus placebo. ", "The study supports the scientific rationale for a reduced dose in that jurisdiction. ", "It does not independently establish India-approved renal dosing."),
        ]),
        ("9", "Hepatic Impairment, Age, and Other Intrinsic Factors", "Special-population evidence clarifies when routine use becomes uncertain.", [
            ("What is known in hepatic impairment? ", "Moderate hepatic impairment did not produce a clinically meaningful exposure change in the available regulatory evidence. ", "Limited hepatic metabolism may explain this observation. ", "Severe hepatic impairment remains insufficiently studied and requires label-aligned caution."),
            ("Does older age change the decision? ", "Older adults often have reduced renal function, polypharmacy, frailty, and greater vulnerability to hypoglycaemia or medication error. ", "Chronological age alone is less informative than renal function and overall treatment context. ", "Final geriatric advice must follow the approved local label."),
            ("What other factors deserve attention? ", "Concomitant insulin or secretagogues, irregular medication routines, cognitive impairment, and caregiver involvement can materially affect weekly-dose safety. ", "A weekly product may simplify one aspect of treatment while adding calendar-related complexity. ", "Individualized education and verification remain essential."),
        ]),
        ("10", "Approved Role, Patient Selection, and Boundaries", "A high-quality monograph helps clinicians identify both suitable and unsuitable use.", [
            ("Which patient might be considered? ", "A locally label-eligible adult with type 2 diabetes may be considered when DPP-4 inhibition is clinically appropriate and a weekly oral schedule fits the treatment plan. ", "The decision should incorporate glycaemic targets, comorbidities, renal function, concomitant therapy, and preferences. ", "This is a patient-selection framework, not a recommendation to prescribe."),
            ("Which patient needs a different priority? ", "Patients with compelling need for therapies with proven cardiovascular, heart-failure, kidney, or substantial weight benefits may require other treatment classes according to current guidelines. ", "Recognizing the limits of the trelagliptin evidence base protects clinical relevance. ", "The product should not be positioned as having benefits it has not demonstrated."),
            ("What conditions require particular caution? ", "Type 1 diabetes, ketoacidosis, severe acute illness, perioperative states, significant renal impairment, and combinations that raise hypoglycaemia risk require careful label review. ", "The weekly interval can prolong exposure after an inappropriate dose. ", "Exact contraindications and precautions remain package-insert controlled."),
        ]),
        ("11", "Dosage, Administration, and Medication-Error Prevention", "Weekly dosing changes the operational work of prescribing and counselling.", [
            ("What does a correct weekly routine require? ", "Japan regulatory information describes administration on the same day each week, with dose strength determined partly by renal function. ", "A fixed weekly day can support routine formation and accurate follow-up. ", "India dosing and strength selection must be confirmed from the current approved package insert."),
            ("What is the most preventable risk? ", "Duplicate dosing, confusion during switching, and failure to adjust for renal function are foreseeable administration hazards. ", "Prescribers, pharmacists, patients, and caregivers should use a shared written dosing plan. ", "The final counselling script must reproduce approved local instructions."),
            ("How should response be reviewed? ", "Glycaemic response should be assessed using clinically appropriate measures and treatment goals rather than assuming benefit from convenience. ", "Review should also test whether the weekly routine is understood and followed correctly. ", "Failure to achieve goals should trigger reassessment of adherence, diagnosis, dose, combination therapy, and treatment strategy."),
        ], [
            "Record the chosen weekly administration day in the prescription and patient record.",
            "Assess renal function before selecting strength and at clinically appropriate intervals.",
            "Confirm the patient is not continuing a previous daily DPP-4 inhibitor unintentionally.",
            "Provide label-aligned missed-dose and duplicate-dose instructions.",
            "Review hypoglycaemia risk when insulin or insulin secretagogues are used.",
        ]),
        ("12", "Clinical Development Programme: Evidence Map", "The evidence map shows what was studied, what was learned, and where uncertainty remains.", [
            ("How broad is the programme? ", "The evidence base includes phase II dose-ranging, phase III active-controlled and long-term studies, switching studies, insulin add-on work, severe renal impairment, treatment-satisfaction studies, exploratory biomarker work, and an Indian comparative phase III trial. ", "This supports a multi-dimensional understanding of efficacy, safety, practical use, and patient experience. ", "Not every study has equal evidentiary weight or direct India-label relevance."),
            ("Which evidence should carry the most weight? ", "Randomized controlled pivotal and supportive trials provide the strongest basis for efficacy and safety interpretation, while open-label, single-arm, and exploratory studies provide contextual evidence. ", "A transparent hierarchy prevents selective emphasis. ", "Mechanistic and patient-reported outcomes should not be used to overstate clinical benefit."),
            ("What important programme remains absent? ", "No dedicated cardiovascular-outcomes programme or large India real-world outcomes programme was identified. ", "The monograph must therefore distinguish glycaemic efficacy from long-term clinical outcomes. ", "Absence of evidence should remain visible rather than being filled by class assumptions."),
        ]),
    ]
    for args in chapters:
        chapter(doc, *args)
    studies = [
        {"title": "Phase II Dose-Ranging Study", "question": "What weekly dose produced a clinically useful balance of DPP-4 inhibition, HbA1c lowering, and tolerability?",
         "details": [("Design", "Randomized, double-blind, placebo-controlled, parallel-group phase II study"), ("Population", "Japanese patients with type 2 diabetes inadequately controlled by diet/exercise"), ("Interventions", "Weekly trelagliptin doses across 12.5-200 mg versus placebo"), ("Duration", "12 weeks"), ("Primary focus", "Dose-response for HbA1c and safety"), ("Key result", "Dose-dependent HbA1c lowering; 100 mg selected for later development"), ("Source", "PMID 24622716; PMDA review")],
         "analysis": [("What did dose ranging establish? ", "The study demonstrated a graded glycaemic response across weekly doses and supported 100 mg as a development dose. ", "It connected sustained pharmacodynamics with clinically measurable HbA1c reduction. ", "Short duration and a Japanese study population limit long-term and broader-population inference."),
                      ("What did safety observations show? ", "No investigator-defined hypoglycaemia was reported in the dose-ranging study summary. ", "This is consistent with glucose-dependent class pharmacology in the studied setting. ", "It does not establish absence of hypoglycaemia in combination therapy or routine practice.")],
         "interpretation": "This study established proof of concept and dose selection. It should be used to explain development logic, not to claim that higher doses or weekly dosing are universally preferable."},
        {"title": "Phase III Trelagliptin Versus Alogliptin", "question": "Could once-weekly trelagliptin maintain glycaemic efficacy comparable with a daily DPP-4 inhibitor?",
         "details": [("Design", "Randomized, double-blind, double-dummy, active- and placebo-controlled phase III study"), ("Population", "Japanese patients with type 2 diabetes"), ("Intervention", "Trelagliptin 100 mg once weekly"), ("Comparator", "Alogliptin 25 mg once daily; placebo arm"), ("Duration", "24 weeks"), ("Primary endpoint", "Change in HbA1c and non-inferiority framework"), ("Key numerical result", "HbA1c change approximately -0.33% versus -0.45%; difference 0.11%, 95% CI -0.054 to 0.281"), ("Source", "PMID 25609193")],
         "analysis": [("What did non-inferiority mean in this trial? ", "The prespecified analysis supported non-inferiority of trelagliptin to alogliptin for HbA1c change in the studied population. ", "The finding supports comparable glycaemic efficacy within the trial's design and margin. ", "Non-inferiority is not superiority and should not be generalized beyond the studied population, dose, comparator, and duration."),
                      ("How should the safety comparison be read? ", "Overall adverse-event frequency was described as similar, with no reported hypoglycaemia in the trelagliptin group. ", "The controlled design supports a useful comparative safety perspective. ", "The sample and duration are insufficient for uncommon or long-latency risks.")],
         "interpretation": "The study is the pivotal comparative efficacy anchor for trelagliptin. Every use of the non-inferiority finding must retain the comparator, endpoint, population, duration, and confidence interval."},
        {"title": "52-Week Monotherapy and Combination Study", "question": "Does glycaemic efficacy and tolerability persist through one year across common background therapies?",
         "details": [("Design", "Open-label phase III long-term study"), ("Population", "680 Japanese patients with type 2 diabetes"), ("Intervention", "Trelagliptin 100 mg weekly as monotherapy or with other glucose-lowering agents"), ("Duration", "52 weeks"), ("Key result", "Mean HbA1c changes varied by background regimen, approximately -0.26% to -0.73%"), ("Safety", "Adverse drug-reaction rates varied across treatment groups"), ("Source", "PMID 27181699")],
         "analysis": [("What did the long-term study contribute? ", "It extended exposure to 52 weeks and examined trelagliptin across several clinically relevant treatment combinations. ", "This supports understanding of durability and tolerability beyond short pivotal trials. ", "Open-label design and absence of a concurrent comparator limit causal comparison."),
                      ("Why do background regimens matter? ", "Observed HbA1c changes and safety events differed across combination groups. ", "This reflects the clinical importance of concomitant therapy. ", "Cross-group differences should not be interpreted as randomized comparisons.")],
         "interpretation": "This study supports one-year treatment experience but is best interpreted descriptively. It cannot establish comparative superiority or long-term outcome benefit."},
        {"title": "Trelagliptin Added to Insulin", "question": "What happens when weekly DPP-4 inhibition is added to insulin in inadequately controlled type 2 diabetes?",
         "details": [("Design", "Multicentre randomized double-blind placebo-controlled 12-week phase followed by 40-week open-label extension"), ("Population", "Japanese patients receiving insulin with HbA1c 7.5-10.0%"), ("Randomized groups", "Trelagliptin 100 mg weekly plus insulin, n=116; placebo plus insulin, n=124"), ("Primary finding", "Placebo-adjusted HbA1c reduction approximately -0.63% at 12 weeks"), ("Safety focus", "Hypoglycaemia and long-term tolerability"), ("Source", "PMCID PMC6175153")],
         "analysis": [("What did the controlled phase demonstrate? ", "Adding trelagliptin to insulin improved HbA1c versus placebo over 12 weeks. ", "The result supports additive glycaemic efficacy in the studied insulin-treated population. ", "Combination therapy increases the importance of hypoglycaemia surveillance and insulin review."),
                      ("How should the extension be interpreted? ", "The open-label extension provided longer exposure and safety observation. ", "It adds useful treatment-experience context. ", "Without continued placebo control, later outcomes are descriptive.")],
         "interpretation": "The study supports add-on efficacy with insulin in Japan. India-facing combination claims and dose-management instructions require confirmation from the India-approved label."},
        {"title": "Severe Renal Impairment and ESRD Study", "question": "Can a reduced weekly dose provide glycaemic benefit in severe renal impairment or ESRD?",
         "details": [("Design", "Randomized phase III study"), ("Population", "Japanese patients with severe renal impairment or ESRD"), ("Intervention", "Trelagliptin 25 mg once weekly"), ("Controlled duration", "12 weeks"), ("Key result", "HbA1c decreased approximately 0.72% versus placebo"), ("Source", "PMCID PMC7078116")],
         "analysis": [("What clinical gap did this study address? ", "Advanced renal impairment changes exposure and narrows treatment choices. ", "The study linked a one-quarter dose to glycaemic efficacy in the studied Japanese population. ", "It does not authorize extrapolation of Japan renal dosing into India without the local label."),
                      ("Why is renal monitoring still necessary? ", "Renal function can change over time and dialysis removes only a limited proportion of a dose. ", "Correct strength selection remains a continuing safety task. ", "A reduced dose is not a substitute for ongoing assessment.")],
         "interpretation": "This study is essential to the scientific rationale for renal dose reduction. Its regulatory application remains jurisdiction-specific."},
        {"title": "Switching From Daily DPP-4 Inhibition", "question": "Can stable patients switch from a daily DPP-4 inhibitor to weekly trelagliptin without major glycaemic disruption?",
         "details": [("Design", "Open-label phase III exploratory study"), ("Population", "Japanese patients with stable glycaemic control on daily sitagliptin"), ("Intervention", "Switch to weekly trelagliptin"), ("Focus", "Glycaemic control, safety, and practical switching"), ("Source", "PMCID PMC5835476")],
         "analysis": [("What did the switching study suggest? ", "Stable glycaemic control was broadly maintained after switching in the selected study population. ", "The finding supports the feasibility of a planned switch under defined conditions. ", "It is not a universal switching recommendation and does not remove the need to prevent overlap or dosing error."),
                      ("What is the practical lesson? ", "Switching requires clear discontinuation of the previous DPP-4 inhibitor and clear identification of the weekly administration day. ", "Operational precision is part of clinical safety. ", "The current local label and clinical judgment must govern any switch.")],
         "interpretation": "The study supports a structured switching concept in selected stable patients. Its exploratory design and small sample require restraint."},
        {"title": "Treatment Satisfaction and Quality-of-Life Studies", "question": "Does weekly administration change how patients experience diabetes treatment?",
         "details": [("Design", "Open-label randomized studies"), ("Populations", "Japanese patients continuing or initiating DPP-4 inhibitor therapy"), ("Comparisons", "Weekly trelagliptin versus daily DPP-4 inhibitor strategies"), ("Outcomes", "DTSQ, DTR-QOL, treatment-burden and satisfaction measures"), ("Sources", "PMID 29093280; PMCID PMC6612345")],
         "analysis": [("What did patient-reported outcomes suggest? ", "Some studies reported improvement in selected treatment-burden or quality-of-life domains with weekly administration. ", "This supports the idea that dosing frequency can matter to selected patients. ", "Patient-reported improvement does not prove better adherence, glycaemic outcomes, or persistence in broader practice."),
                      ("Why is the wording important? ", "Satisfaction outcomes are vulnerable to open-label expectations and individual preference. ", "They should be described precisely by instrument and domain. ", "A general claim that weekly dosing improves adherence remains unsupported without direct adherence evidence.")],
         "interpretation": "Patient-experience evidence can enrich shared decision-making, but it must not be elevated above clinical efficacy, safety, or outcome evidence."},
        {"title": "Indian Phase III Comparison With Vildagliptin", "question": "How did once-weekly trelagliptin perform against a commonly used twice-daily DPP-4 inhibitor in Indian patients?",
         "details": [("Design", "Multicentre randomized open-label active-controlled non-inferiority phase III study"), ("Sites", "10 geographically distinct Indian sites"), ("Population", "240 treatment-naive Indian patients with type 2 diabetes randomized 1:1"), ("Interventions", "Trelagliptin 100 mg once weekly versus vildagliptin 50 mg twice daily"), ("Duration", "16 weeks"), ("Primary endpoint", "Non-inferiority for change in HbA1c"), ("Reported HbA1c change", "-0.89% versus -1.00%; difference 0.11%, 95% CI -0.28 to 0.50"), ("Source", "PMCID PMC12215575")],
         "analysis": [("Why is this study especially relevant? ", "It directly evaluates trelagliptin in an Indian population against an active comparator used in India. ", "The study adds local clinical evidence to a programme otherwise dominated by Japanese data. ", "Open-label design, 16-week duration, and the prespecified margin must remain visible."),
                      ("What can the non-inferiority conclusion support? ", "The reported analysis supports non-inferiority for HbA1c reduction within the studied design. ", "This provides a locally relevant efficacy comparison. ", "It cannot establish superiority, long-term safety, adherence advantage, or clinical-outcome equivalence.")],
         "interpretation": "This is the principal India-specific efficacy source identified. Claims must reproduce the design, comparator, duration, numerical result, and limitations accurately."},
        {"title": "Exploratory and Emerging Evidence", "question": "Do biomarker and comparative-mechanism studies add clinically actionable information?",
         "details": [("Examples", "Endothelial-function pilot study; DUET-beta beta-cell function study; PBPK-DPP-4 occupancy modeling"), ("Evidence level", "Exploratory, small, mechanistic, or modeled"), ("Purpose", "Hypothesis generation and refinement of pharmacological understanding"), ("Sources", "PMCID PMC5096292; PMCID PMC11291836; Biomed Pharmacother 2022 model")],
         "analysis": [("What did the endothelial pilot show? ", "A small single-arm study found no significant change in flow-mediated dilation while adiponectin increased. ", "The result demonstrates why biomarker findings require cautious interpretation. ", "It does not support a vascular-outcome claim."),
                      ("What did DUET-beta add? ", "The randomized study compared trelagliptin with dulaglutide for beta-cell-function measures and found different glycaemic and biomarker effects. ", "It helps define the distinct therapeutic intensity and biology of the compared classes. ", "Small sample size and mechanistic endpoints limit broad practice conclusions."),
                      ("How should modeling be used? ", "PBPK and DPP-4 occupancy modeling can illuminate renal-impairment exposure and dosing logic. ", "Models support hypothesis testing and dose rationale. ", "They do not replace observed clinical outcomes or approved dosing instructions.")],
         "interpretation": "Emerging evidence belongs in a scientific monograph, but it should be clearly labelled as exploratory and kept outside core approved claims."},
    ]
    idx = 13
    for study in studies:
        study_profile(doc, idx, study)
        idx += 1
    remaining = [
        ("22", "Safety Profile and Pharmacovigilance Perspective", "Safety interpretation must be as visible and detailed as efficacy.", [
            ("What does the clinical-trial safety record show? ", "Across published trials, trelagliptin was generally tolerated in the studied populations, with overall safety patterns broadly consistent with DPP-4 inhibition. ", "This supports a characterized clinical-trial safety profile. ", "Trial size and duration limit detection of uncommon, delayed, or population-specific risks."),
            ("Which risks demand active awareness? ", "Applicable regulatory information identifies hypoglycaemia in combination settings and class- or product-relevant concerns including pancreatitis, hypersensitivity, bullous pemphigoid, and ileus. ", "Early recognition and label-aligned action are important because pharmacological effects may persist after discontinuation. ", "Exact India warning and adverse-reaction wording must be verified."),
            ("Why must safety remain jurisdiction-specific? ", "Post-marketing observations and label revisions can differ by market and time. ", "A monograph must reconcile current local pharmacovigilance information immediately before approval. ", "Historical Japan safety language cannot be presented as the current India label."),
        ]),
        ("23", "Hypoglycaemia, Pancreatitis, and Other Clinically Important Risks", "Risk sections should tell the HCP what to notice, what to verify, and what remains uncertain.", [
            ("When is hypoglycaemia most relevant? ", "DPP-4 inhibition has glucose-dependent effects, but the risk increases when combined with insulin or insulin secretagogues. ", "Medication review and patient education are central risk controls. ", "Dose changes to concomitant therapy must follow clinical judgment and approved guidance."),
            ("How should abdominal symptoms be approached? ", "Persistent severe abdominal pain or vomiting can require assessment for pancreatitis according to applicable regulatory guidance. ", "Weekly exposure makes prompt recognition particularly important. ", "The monograph cannot replace urgent clinical evaluation or the package insert."),
            ("What other signals should be recognized? ", "Hypersensitivity, blistering skin disease, and symptoms suggestive of intestinal obstruction are among clinically important concerns identified in regulatory context. ", "Clear counselling can improve recognition and escalation. ", "Frequency and exact action language require the current local label."),
        ]),
        ("24", "Interactions and Concomitant Therapy", "The interaction chapter connects pharmacology with the realities of polypharmacy.", [
            ("Are metabolic interactions the principal issue? ", "Trelagliptin undergoes limited metabolism and is largely eliminated unchanged in urine. ", "This reduces emphasis on extensive CYP-mediated interaction pathways. ", "It does not remove the need to review all concomitant medicines and renal function."),
            ("Which combinations change the safety picture? ", "Insulin and insulin secretagogues can increase hypoglycaemia risk, while other glucose-lowering therapies may alter the overall benefit-risk balance. ", "Combination treatment should be purposeful and monitored. ", "Approved combination use and dose advice remain jurisdiction-specific."),
            ("Why should duplicate DPP-4 inhibition be prevented? ", "A weekly medicine can overlap unintentionally with a previous daily DPP-4 inhibitor during switching or fragmented care. ", "Medication reconciliation and clear documentation prevent avoidable exposure. ", "There is no clinical rationale for routine combination of two DPP-4 inhibitors."),
        ]),
        ("25", "Place in Therapy and Fair Comparative Framing", "Place-in-therapy analysis should clarify choice without manufacturing superiority.", [
            ("What is the most defensible differentiation? ", "Trelagliptin's clearest demonstrated distinction is once-weekly oral administration with HbA1c-lowering efficacy supported by comparative studies. ", "This may be relevant to selected patients and treatment routines. ", "Convenience must not be presented as superior efficacy, adherence, or outcomes."),
            ("How should it be compared with other DPP-4 inhibitors? ", "Head-to-head evidence exists with alogliptin and vildagliptin in defined populations and durations. ", "These trials permit precise non-inferiority statements. ", "They do not justify broad class superiority or cross-trial ranking."),
            ("How should it be compared with outcome-directed therapies? ", "Other classes have demonstrated cardiovascular, kidney, heart-failure, or weight outcomes in selected populations. ", "Those outcome priorities may dominate treatment selection. ", "A balanced monograph must state that trelagliptin has not established those benefits."),
        ]),
        ("26", "Practical HCP Initiation and Monitoring Framework", "Practical guidance turns scientific knowledge into safer clinical workflow.", [
            ("What should be checked before initiation? ", "Confirm type 2 diabetes, current treatment, renal function, glycaemic status, comorbidities, pregnancy considerations, hypoglycaemia risk, and the current approved indication. ", "A structured baseline reduces avoidable mismatch. ", "This framework supplements but does not replace the local label or clinical judgment."),
            ("What should be reviewed after initiation? ", "Assess glycaemic response, tolerability, hypoglycaemia, medication understanding, weekly-day adherence, renal function when appropriate, and whether treatment priorities have changed. ", "Review distinguishes actual clinical value from theoretical convenience. ", "Persistent inadequate control should trigger treatment reassessment."),
            ("What should patients understand? ", "Patients should know the weekly day, what to do if a dose is missed, how to avoid duplicate dosing, when to seek help, and why other diabetes medicines may still be necessary. ", "Clear education is a core safety intervention. ", "Final counselling must use approved local instructions."),
        ], ["Confirm the current India-approved indication and dose.", "Document renal function and selected strength.", "Reconcile all glucose-lowering medicines.", "Choose and document a fixed weekly day.", "Provide written missed-dose guidance from the approved label.", "Review response and safety at an appropriate interval."]),
        ("27", "Evidence Gaps and Research Priorities", "A trustworthy monograph makes the limits of knowledge easy to find.", [
            ("Which outcomes remain unproven? ", "Dedicated evidence for cardiovascular events, kidney outcomes, mortality, microvascular complications, and long-term adherence-related clinical benefit was not identified. ", "These gaps define the boundaries of responsible claims. ", "They should remain visible in summaries and comparisons."),
            ("Which populations need more evidence? ", "Broader India real-world populations, severe hepatic impairment, pregnancy, lactation, pediatric patients, and diverse multimorbidity settings remain incompletely characterized. ", "Further evidence could refine patient selection and safety. ", "Until then, use must remain label-aligned and cautious."),
            ("What would strengthen future assessment? ", "Large pragmatic studies, post-marketing surveillance, medication-error analyses, long-term persistence studies, and outcome-focused comparative research would add clinical value. ", "A living monograph should update when such evidence emerges. ", "Future evidence must be critically appraised rather than added automatically."),
        ]),
        ("28", "Balanced Scientific Conclusions", "The conclusion should help an HCP remember the evidence, the practical role, and the limits together.", [
            ("What is established? ", "Trelagliptin is a selective long-acting DPP-4 inhibitor with a once-weekly oral schedule and an evidence base demonstrating HbA1c lowering in studied populations. ", "Comparative trials support non-inferiority to selected daily DPP-4 inhibitors under defined conditions. ", "The conclusion must remain tied to those endpoints and conditions."),
            ("What is clinically distinctive? ", "The weekly schedule and sustained pharmacodynamic activity create a different administration model from daily DPP-4 inhibitors. ", "This may be practically relevant for selected patients. ", "It also creates distinctive renal-dose, switching, missed-dose, and duplicate-dose controls."),
            ("What is not established? ", "The evidence does not establish superiority, cardiovascular or renal outcome benefit, mortality benefit, or universal adherence improvement. ", "Those boundaries preserve scientific credibility and support informed treatment selection. ", "The current India package insert remains essential before external use."),
        ]),
    ]
    for args in remaining:
        chapter(doc, *args)
    add_evidence_methodology_annex(doc, "trelagliptin")
    add_hcp_faq_annex(doc, "trelagliptin", {
        "one_line": "Trelagliptin is a selective oral DPP-4 inhibitor whose sustained pharmacodynamic activity permits once-weekly administration for type 2 diabetes in approved settings. Its evidence base supports HbA1c lowering, while its defining practical feature is administration frequency rather than proven superiority or cardiorenal benefit.",
        "role": "An India approval listing identifies trelagliptin succinate tablets for treatment of type 2 diabetes, but the complete current India package insert remains required to define exact patient population, combination use, dosage, and restrictions. Japan regulatory information is useful scientific context and must remain jurisdiction-qualified.",
        "distinction": "The principal scientific distinction is sustained DPP-4 inhibition across a seven-day interval. This enables a weekly oral schedule, but it also creates distinctive responsibilities around renal dose selection, missed doses, duplicate-dose prevention, switching, and persistence of pharmacological effects.",
        "efficacy": "The strongest direct efficacy evidence includes the Japanese double-blind phase III comparison with daily alogliptin and the Indian active-controlled phase III comparison with twice-daily vildagliptin. Both reported non-inferiority for HbA1c change within their prespecified designs and durations.",
        "india": "The Indian phase III study randomized 240 treatment-naive patients across 10 sites and compared trelagliptin 100 mg weekly with vildagliptin 50 mg twice daily for 16 weeks. The reported HbA1c changes were -0.89% and -1.00%, respectively, with a between-group difference of 0.11%.",
        "international": "The Japanese programme includes dose-ranging, pivotal comparative, 52-week combination, insulin add-on, switching, severe renal impairment, and patient-reported-outcome studies. Together they characterize glycaemic efficacy and practical use, but they do not constitute a cardiovascular-outcomes programme.",
        "safety": "The safety discussion should emphasize hypoglycaemia when combined with insulin or secretagogues, renal-function-related exposure, weekly-dose medication errors, and applicable class or product warnings such as pancreatitis, hypersensitivity, bullous pemphigoid, and ileus. Exact India wording remains label-controlled.",
        "monitoring": "Renal function, glycaemic response, concomitant glucose-lowering therapy, hypoglycaemia, medication understanding, and the chosen weekly administration day are central monitoring considerations. The current local label must define the exact dosing and missed-dose pathway.",
        "never_claim": "The evidence does not support claims that trelagliptin is superior to other DPP-4 inhibitors, guarantees adherence, improves cardiovascular or kidney outcomes, reduces mortality, or is suitable for every patient with type 2 diabetes. Weekly administration is a fact, not an outcome claim.",
        "mechanism": "DPP-4 inhibition prolongs endogenous incretin activity, supporting glucose-dependent insulin secretion and suppression of inappropriate glucagon secretion. The mechanism explains glycaemic action and generally low intrinsic hypoglycaemia propensity, but it does not prove organ protection or long-term clinical outcomes.",
        "pk": "Trelagliptin is absorbed relatively rapidly, undergoes limited metabolism, and is predominantly eliminated unchanged in urine. Exposure rises as renal function declines, making renal assessment and correct strength selection clinically important rather than merely pharmacokinetic details.",
        "hierarchy": "Randomized pivotal trials should anchor efficacy claims. Long-term open-label, switching, satisfaction, biomarker, and modeling studies provide supportive context and should be clearly labelled according to their design and limitations.",
        "comparison": "Every comparison should retain population, comparator, dose, background therapy, duration, endpoint, effect estimate, confidence interval, and limitation. Non-inferiority to alogliptin or vildagliptin does not establish broad superiority, interchangeability, or outcome equivalence.",
        "limitations": "Most pivotal data were generated in Japan, many trials were short, and hard clinical outcomes were not assessed. India evidence is important but currently limited, and the current complete India package insert and post-marketing evidence remain necessary.",
        "long_term": "A 52-week open-label study supports one-year treatment experience across monotherapy and combination settings. Because it lacked a concurrent control group, it supports descriptive durability and tolerability rather than comparative long-term superiority or outcome benefit.",
        "elderly": "Older adults require individualized review of renal function, frailty, hypoglycaemia risk, cognitive capacity, polypharmacy, and caregiver support. Weekly dosing may simplify frequency while increasing the consequences of calendar confusion or an incorrect strength.",
        "renal": "Renal impairment increases exposure, and Japanese regulatory work plus a phase III study support reduced strengths in advanced impairment. The scientific rationale is strong, but exact India thresholds and doses must be reproduced from the current India-approved package insert.",
        "hepatic": "Available evidence suggests moderate hepatic impairment does not markedly alter exposure, consistent with limited hepatic metabolism. Severe hepatic impairment remains insufficiently characterized, and local approved guidance should control use.",
        "pregnancy": "Adequate pregnancy and lactation evidence was not identified. The monograph should avoid reassuring language and should defer to the current approved label, clinical judgment, and the individual's benefit-risk context.",
        "pediatric": "A pediatric efficacy and safety programme adequate to support routine use was not identified. Pediatric use must not be inferred from adult pharmacology or DPP-4 class experience.",
        "concomitant": "Medication reconciliation should identify insulin, sulfonylureas, other secretagogues, previous daily DPP-4 inhibitors, and therapies selected for cardiovascular, renal, or weight objectives. Duplicate DPP-4 inhibition and unintended overlap during switching should be prevented.",
        "selection": "A suitable patient would be locally label-eligible, need DPP-4 inhibition for glycaemic management, have renal function assessed, and understand the weekly routine. The choice should align with comorbidities, treatment objectives, preferences, and alternative therapies.",
        "alternative": "Another strategy may be more appropriate when cardiovascular disease, heart failure, chronic kidney disease, substantial weight reduction, acute metabolic decompensation, or inability to follow a weekly schedule dominates the decision. The monograph should make these boundaries visible.",
        "before": "Before initiation, confirm type 2 diabetes, the current India-approved indication, baseline glycaemia, renal function, concomitant medicines, hypoglycaemia risk, and the patient's ability to follow a weekly schedule. Document the selected weekly day and planned review.",
        "after": "Review glycaemic response, tolerability, hypoglycaemia, renal function where appropriate, correct weekly use, and whether treatment priorities have changed. Lack of adequate response should prompt reassessment rather than an assumption that convenience equals effectiveness.",
        "counselling": "Patients should understand the weekly administration day, label-approved missed-dose instructions, duplicate-dose prevention, signs requiring medical attention, and the continuing role of diet, activity, monitoring, and other prescribed medicines.",
        "qol": "Randomized open-label studies reported improvement in selected treatment-burden or quality-of-life domains, but results were not uniformly significant. These findings support individualized discussion and do not prove improved adherence or clinical outcomes.",
        "outcomes": "No dedicated evidence was identified showing reduction in cardiovascular events, kidney outcomes, hospitalization, mortality, or microvascular complications. The evidence-supported core outcome remains glycaemic control.",
        "statistics": "A p-value or non-inferiority conclusion must be read with the effect size, confidence interval, margin, missing-data method, and clinical context. Numerical results should not be simplified into unsupported adjectives.",
        "neutral": "Neutral biomarker, vascular, weight, or treatment-satisfaction findings are part of the scientific story. Including them clarifies the product's demonstrated scope and prevents selective interpretation.",
        "pv": "Pharmacovigilance should reconcile clinical-trial safety with current post-marketing information, regulator communications, medication-error reports, and sponsor safety data. Weekly exposure and renal impairment warrant particular attention.",
        "guidelines": "Guidelines place DPP-4 inhibitors within individualized diabetes care but often prioritize other classes when proven cardiovascular, kidney, heart-failure, or weight benefits are required. Trelagliptin-specific placement should therefore remain outcome-aware and patient-specific.",
        "foreign": "Foreign-label information can support scientific understanding only when visibly identified by jurisdiction. It cannot define India-facing dosage, contraindications, warnings, or promotional claims.",
        "external": "Before external use, obtain the current India-approved package insert, sponsor product particulars, updated pharmacovigilance data, independently verified study numbers, complete claim-source matrix, rendered visual QA, and medical, regulatory, legal, and compliance approval.",
        "study_table": "Reviewers should verify design, randomization, blinding, population, analysis set, dose, comparator, duration, endpoint hierarchy, statistical method, effect estimate, confidence interval, adverse events, funding, and limitations against the original source.",
        "future": "The most useful future evidence would include long-term India real-world safety, medication-error and persistence studies, broader post-marketing surveillance, and outcome-focused research in clinically relevant populations.",
        "update": "New evidence should enter through a documented update search, source verification, claim-impact assessment, pharmacovigilance review, and controlled approval. Prior wording should remain traceable rather than being silently overwritten.",
        "conclusion": "Trelagliptin provides sustained weekly DPP-4 inhibition and evidence-supported glycaemic efficacy comparable with selected daily DPP-4 inhibitors under studied conditions. Its weekly schedule can be clinically relevant, but outcome superiority and universal adherence benefit remain unestablished.",
    })
    add_scientific_glossary(doc, "trelagliptin", [
        ("DPP-4 inhibition", "Pharmacological inhibition of dipeptidyl peptidase-4, reducing degradation of endogenous incretin hormones. For trelagliptin, sustained inhibition supports weekly administration but does not itself establish long-term clinical outcomes."),
        ("Incretin effect", "The greater insulin response to oral glucose than intravenous glucose, mediated partly by GLP-1 and GIP. DPP-4 inhibitors enhance endogenous incretin activity in a glucose-dependent manner."),
        ("Weekly dosing interval", "A seven-day administration interval supported by the molecule's pharmacodynamic profile. It requires precise medication reconciliation, renal-dose selection, missed-dose instructions, and duplicate-dose prevention."),
        ("Renal exposure increase", "The progressive rise in systemic exposure as renal function declines because trelagliptin is predominantly renally eliminated. It is the scientific basis for jurisdiction-specific reduced-dose strategies."),
        ("Glycaemic efficacy", "Improvement in measures such as HbA1c, fasting glucose, or postprandial glucose. Glycaemic efficacy should not be equated with proven cardiovascular, renal, mortality, or microvascular benefit."),
    ])
    add_clinical_review_lenses(doc, "trelagliptin")
    add_prescribing_placeholder(doc, "trelagliptin", [
        "Exact India-approved indication and combination-use wording.",
        "Approved strengths and renal-dose thresholds.",
        "Missed-dose, duplicate-dose, and switching instructions.",
        "Formal contraindications, warnings, adverse reactions, interactions, and special-population statements.",
        "Current India MAH, pack, storage, legal category, and revision date.",
    ])
    add_claim_matrix(doc, "trelagliptin", [
        ("TR-01", "Trelagliptin is a selective oral DPP-4 inhibitor developed for once-weekly administration.", "PMDA review; PMID 26115728", "Administration frequency does not establish superior outcomes.", "Scientific supported"),
        ("TR-02", "Trelagliptin was non-inferior to daily alogliptin for HbA1c change in a 24-week Japanese phase III study.", "PMID 25609193", "Retain population, comparator, endpoint, duration, margin, and CI.", "Scientific supported"),
        ("TR-03", "An Indian phase III study reported non-inferiority to twice-daily vildagliptin over 16 weeks.", "PMCID PMC12215575", "Open-label and short-duration limitations required.", "Scientific supported"),
        ("TR-04", "A 25-mg weekly dose reduced HbA1c versus placebo in Japanese patients with severe renal impairment or ESRD.", "PMCID PMC7078116", "Japan evidence; do not infer India-approved renal dose.", "Scientific supported"),
        ("TR-05", "Weekly dosing improves adherence and outcomes.", "No adequate direct outcome evidence", "Prohibited unless directly substantiated and approved.", "Prohibited"),
        ("TR-06", "Trelagliptin provides cardiovascular or kidney protection.", "No dedicated outcome evidence identified", "Unsupported.", "Prohibited"),
    ])
    add_source_register(doc, TRELAGLIPTIN_REFS)
    add_heading(doc, "Appendix: Internal Review Questions", 1)
    add_bullets(doc, [
        "Does every efficacy statement identify the studied population, comparator, duration, endpoint, and limitation?",
        "Are Japan regulatory statements clearly separated from India prescribing claims?",
        "Is every weekly-dose instruction verified against the current India package insert?",
        "Are cardiorenal, mortality, adherence, and superiority claims absent unless directly substantiated?",
        "Does safety receive prominence comparable to efficacy?",
        "Has every numerical result been independently reproduced from the original source?",
    ])
    path = OUT / "PM-EXP-Q1_Trelagliptin_Expanded_HCP_Scientific_Monograph_R2.docx"
    doc.save(path)
    return path


def build_seratrodast():
    doc = setup_document("Seratrodast")
    front_matter(
        doc, "Seratrodast",
        "A detailed scientific reference on an oral thromboxane A2/prostaglandin H2 receptor antagonist approved in India as add-on therapy in bronchial asthma.",
        "CDSCO lists 40-mg and 80-mg tablets approved on 30 November 2012 as add-on therapy in bronchial asthma.",
        "The current India-approved package insert was not retrieved; Japanese dosing and safety statements remain clearly jurisdiction-qualified."
    )
    executive_summary(doc, "Seratrodast", [
        ("Why consider the thromboxane pathway in asthma? ", "Thromboxane A2 can contribute to airway smooth-muscle constriction, hyperresponsiveness, inflammation-related signaling, and airway secretion characteristics. ", "Seratrodast antagonizes the TP receptor and provides a mechanistically distinct oral add-on approach. ", "Mechanistic plausibility does not make it a rescue bronchodilator or a replacement for inhaled corticosteroid-containing treatment."),
        ("What does the clinical evidence show? ", "Randomized and supportive studies report effects on selected asthma symptoms, peak-flow variability, airway hyperresponsiveness, sputum characteristics, mucociliary clearance, and selected lung-function measures. ", "An Indian double-blind comparative study evaluated seratrodast against montelukast as add-on therapy. ", "The evidence is generally older, short-duration, limited in size, and does not establish reduction in severe exacerbations, hospitalization, or mortality."),
        ("What is the most important safety signal? ", "Japanese regulatory safety information identifies serious hepatic dysfunction, jaundice, and fulminant hepatitis, with monthly liver-function monitoring in that jurisdiction. ", "This signal requires prominent risk assessment and monitoring discussion. ", "The exact current India warning, monitoring, and contraindication language must be confirmed from the approved India package insert."),
        ("Where might seratrodast fit? ", "The verified India approval-list wording identifies add-on therapy in bronchial asthma. ", "Its place may be relevant only after ensuring that foundational asthma assessment, inhaled corticosteroid-containing therapy, inhaler technique, adherence, and risk management are addressed. ", "It is not part of mainstream GINA treatment pathways and must not be positioned as guideline-preferred or suitable for acute attacks."),
    ], [
        ("Molecule and class", "Seratrodast; oral TP-receptor antagonist; development code AA-2414."),
        ("Verified India role", "40-mg and 80-mg tablets listed by CDSCO as add-on therapy in bronchial asthma."),
        ("Mechanistic focus", "Antagonism of thromboxane A2/prostaglandin H2 receptor-mediated airway effects."),
        ("Evidence base", "Phase II PK/PD, placebo-controlled sputum study, airway-hyperresponsiveness study, biomarker-response work, Indian montelukast comparison, and Japanese regulatory experience."),
        ("Principal safety focus", "Serious hepatic injury signal and liver-function monitoring in Japanese information."),
        ("Critical boundary", "Not a rapid bronchodilator; not a replacement for inhaled corticosteroid-containing controller therapy."),
    ], [
        "Current complete India-approved package insert was not available.",
        "Evidence is generally older, short-duration, and limited in size.",
        "Robust evidence for severe-exacerbation reduction, hospitalization, mortality, and long-term steroid-sparing benefit was not identified.",
        "Current international asthma guidelines do not establish seratrodast as a preferred controller pathway.",
        "India-specific post-marketing safety and long-term effectiveness evidence remain incompletely characterized.",
    ])
    chapters = [
        ("1", "Asthma as a Heterogeneous Chronic Airway Disease", "Understanding asthma heterogeneity prevents an add-on medicine from being treated as a universal solution.", [
            ("Why is asthma control more complex than symptom relief? ", "Asthma combines variable airflow limitation, airway inflammation, bronchial hyperresponsiveness, triggers, comorbidities, and fluctuating future risk. ", "Treatment must address both current symptoms and the risk of exacerbation, loss of lung function, and treatment-related harm. ", "A medicine that changes selected symptoms or biomarkers does not automatically address every dimension of control."),
            ("What should be optimized before adding therapy? ", "Diagnosis, inhaler technique, adherence, exposure to triggers, smoking, comorbid rhinitis, obesity, reflux, and the adequacy of inhaled corticosteroid-containing treatment should be reviewed. ", "This protects patients from escalation that bypasses correctable causes. ", "Seratrodast's add-on role should be interpreted within this broader asthma-management framework."),
            ("Why does phenotype matter? ", "Different inflammatory and mediator pathways can dominate in different patients, and response to targeted add-on approaches may vary. ", "Thromboxane-related biology may be more relevant in selected patients than others. ", "No validated routine biomarker currently establishes who should receive seratrodast in ordinary practice."),
        ]),
        ("2", "The Thromboxane Pathway in Airway Biology", "The pathway explains the scientific rationale while also revealing the limits of inference.", [
            ("What is thromboxane A2? ", "Thromboxane A2 is an unstable arachidonic-acid-derived mediator that acts through the TP receptor and can influence smooth-muscle tone, platelet activity, vascular responses, and inflammatory signaling. ", "Airway effects provide a rationale for receptor antagonism in asthma. ", "Pathway involvement does not establish that blockade will improve all clinically important outcomes."),
            ("How can thromboxane influence asthma? ", "Experimental and clinical observations associate thromboxane signaling with bronchoconstriction, airway hyperresponsiveness, and aspects of airway inflammation and secretion. ", "This supports evaluation of TP-receptor antagonists as controller or add-on therapy. ", "Seratrodast is not a rapid bronchodilator and should not be used to treat an acute attack."),
            ("Why is a receptor antagonist different from standard controllers? ", "Seratrodast blocks TP-receptor signaling, whereas inhaled corticosteroids address broad airway inflammation and modern biologics target selected immune pathways. ", "Mechanistic diversity may create an adjunctive role. ", "It does not justify replacement of therapies with stronger guideline and outcome support."),
        ]),
        ("3", "Discovery, Molecular Identity, and Regulatory History", "The product's history explains both its distinctiveness and the age of much of its evidence base.", [
            ("What is seratrodast? ", "Seratrodast, developed as AA-2414, is an orally active benzoquinone-containing TP-receptor antagonist. ", "It was among the first thromboxane-receptor antagonists used clinically for asthma. ", "Historical first-in-class status is not a claim of current guideline preference."),
            ("What is verified in India? ", "CDSCO records list seratrodast 40-mg and 80-mg tablets, approved on 30 November 2012 as add-on therapy in bronchial asthma. ", "This defines the narrow verified India role available from the public approval record. ", "The current package insert is still required for exact dose, warnings, contraindications, and product particulars."),
            ("What is verified in Japan? ", "Japanese product information documents bronchial-asthma use and provides detailed dosing, PK, and safety context. ", "This information is valuable for scientific understanding and risk identification. ", "It must remain labelled as Japanese information and cannot silently become India prescribing information."),
        ]),
        ("4", "Mechanism of Action and Translational Relevance", "A clinically useful mechanism chapter distinguishes established effects from hypotheses.", [
            ("How does seratrodast act at the receptor? ", "Seratrodast competitively antagonizes TP receptors and inhibits responses to thromboxane-like agonists. ", "Receptor blockade can reduce thromboxane-mediated airway effects. ", "The magnitude of clinical benefit depends on disease biology, dose, background therapy, and patient selection."),
            ("Which airway effects have been investigated? ", "Studies have examined airway hyperresponsiveness, FEV1, peak-flow variability, symptoms, rescue use, sputum amount and viscosity, mucociliary clearance, and mediator biomarkers. ", "This breadth helps explain potential clinical and physiological effects. ", "Most studies were small or short and do not establish broad long-term outcome benefit."),
            ("What emerging biology should remain exploratory? ", "Recent experimental work describes anti-ferroptotic activity after reduction to a hydroquinone form. ", "The finding may expand scientific interest in the molecule. ", "It is preclinical and must not be used to support asthma efficacy, organ protection, or off-label use."),
        ]),
        ("5", "Pharmacodynamics: Airway Response and Biomarker Questions", "Pharmacodynamic evidence helps explain variability in clinical response.", [
            ("Is response related to exposure? ", "A phase II population PK/PD analysis found a concentration-related increase in FEV1 at 120 mg in the studied population. ", "This supports a pharmacological exposure-response relationship. ", "The studied dose and response cannot be converted into an India dosing recommendation without the approved package insert."),
            ("Can thromboxane metabolites predict response? ", "A small biomarker study suggested that patients with markedly high baseline urinary 11-dehydrothromboxane B2 could show stronger response. ", "This raises a biologically plausible patient-selection hypothesis. ", "The finding is exploratory and does not establish a validated clinical test."),
            ("Does seratrodast reduce eosinophilic inflammation? ", "A small study reported reduced airway hyperresponsiveness without significant change in exhaled nitric oxide or sputum eosinophil percentage. ", "This suggests effects may not be captured by common eosinophilic markers. ", "It also reinforces that TP antagonism should not be assumed to replace anti-inflammatory controller therapy."),
        ]),
        ("6", "Pharmacokinetics: From Oral Dose to Clinical Exposure", "PK helps an HCP understand delayed peak, daily use, and special-population caution.", [
            ("How is seratrodast absorbed and eliminated? ", "Japanese information describes a delayed time to maximum concentration of approximately 7.4 hours and a terminal half-life of approximately 25 hours after an 80-mg dose in healthy men. ", "The profile is compatible with once-daily administration in the Japanese context. ", "Exact India dose and meal timing remain package-insert controlled."),
            ("Does repeated dosing lead to accumulation? ", "Japanese data report no meaningful accumulation after 80 mg once daily for seven days. ", "This supports a relatively predictable repeat-dose profile under studied conditions. ", "It does not eliminate risk in hepatic impairment, older adults, or unstudied combinations."),
            ("Which patient factors influence exposure? ", "Population PK associated clearance and distribution parameters with body weight, while Japanese information reports higher exposure and longer half-life in older adults. ", "These observations support careful clinical assessment. ", "Validated India dose-adjustment rules require the local label."),
        ], None, (["PK dimension", "Reported finding", "Clinical meaning"], [
            ("Tmax", "Approximately 7.4 hours in a Japanese single-dose study", "Not a rapid-relief medicine"),
            ("Terminal half-life", "Approximately 25 hours", "Supports daily exposure in Japan-label context"),
            ("Repeat dosing", "No meaningful accumulation reported over 7 days", "Does not remove special-population caution"),
            ("Population PK", "Linear in an 8-week phase II study; parameters associated with body weight", "Exposure-response evidence is contextual, not dosing authority"),
            ("Elderly", "Higher exposure and longer half-life in Japanese information", "Requires careful local-label review"),
        ], [1.2, 2.55, 2.5])),
        ("7", "Approved Role and the Meaning of Add-On Therapy", "The verified India indication is narrow and must remain central.", [
            ("What does the public CDSCO record establish? ", "Seratrodast 40-mg and 80-mg tablets are listed as approved in India as add-on therapy in bronchial asthma. ", "This supports an adjunctive, not replacement, treatment concept. ", "It does not define exact dose, age, background therapy, contraindications, or duration."),
            ("What should add-on therapy imply in practice? ", "Foundational asthma care and appropriate inhaled corticosteroid-containing therapy should be assessed before considering an adjunct. ", "Add-on treatment should address a defined residual problem and be reviewed for benefit. ", "Seratrodast should not be positioned as monotherapy or as a substitute for rescue treatment."),
            ("What should not be inferred from historical studies? ", "Older studies evaluated varying doses and populations, sometimes before current asthma strategies were established. ", "They help characterize the molecule's effects. ", "They do not automatically define contemporary guideline placement or India-approved use."),
        ]),
        ("8", "Dosage, Administration, and the Acute-Asthma Boundary", "The operational chapter must prevent the most dangerous misunderstanding: this is not a rescue medicine.", [
            ("What dosing information is verified? ", "The public India approval record verifies 40-mg and 80-mg tablet strengths but does not provide complete posology; Japanese information describes 80 mg once daily after the evening meal. ", "The Japanese schedule provides scientific context. ", "Only the current India package insert can define India dosing and administration."),
            ("Why must acute asthma be discussed prominently? ", "Seratrodast does not rapidly reverse established bronchospasm. ", "Patients with worsening symptoms or an acute severe attack require appropriate reliever and emergency management. ", "Delayed use of effective acute therapy is a foreseeable safety risk."),
            ("How should benefit be reviewed? ", "An add-on medicine should be continued only when a clinically meaningful objective and response are identifiable. ", "Symptoms, reliever use, lung function, exacerbation risk, and tolerability can inform review. ", "No response threshold should be invented without approved guidance."),
        ]),
        ("9", "Clinical Development Programme: Evidence Map", "The evidence map prevents a small number of positive findings from being mistaken for a modern outcomes programme.", [
            ("What types of evidence exist? ", "The programme includes Japanese regulatory experience, a phase II population PK/PD study, placebo-controlled studies of airway secretion and hyperresponsiveness, biomarker-response work, and an Indian comparative trial with montelukast. ", "Together they characterize mechanism, selected clinical effects, and local comparative evidence. ", "The programme is not equivalent to a contemporary large exacerbation-outcomes development programme."),
            ("What is the strongest India evidence? ", "The double-blind, double-dummy Indian study compared seratrodast 80 mg with montelukast 10 mg in adults continuing low-dose inhaled corticosteroid therapy. ", "This directly informs India scientific context. ", "Its short 28-day duration limits conclusions about exacerbations, long-term control, and safety."),
            ("Which endpoints dominate the evidence? ", "Symptoms, PEF, FEV1, sputum parameters, mucociliary clearance, airway hyperresponsiveness, and biomarkers feature prominently. ", "These endpoints can be clinically and mechanistically informative. ", "They should not be represented as proof of hospitalization, mortality, or severe-exacerbation reduction."),
        ]),
    ]
    for args in chapters:
        chapter(doc, *args)
    studies = [
        {"title": "Phase II Population PK/PD Study", "question": "How did exposure relate to lung-function response in mild-to-moderate asthma?",
         "details": [("Design", "Randomized, double-blind, placebo-controlled, 15-centre phase II study"), ("Population", "183 patients with mild-to-moderate asthma"), ("Interventions", "Placebo, seratrodast 80 mg daily, or 120 mg daily"), ("Duration", "8 weeks"), ("PK model", "Two-compartment model with zero-order input and first-order elimination"), ("Key PD finding", "At 120 mg, FEV1 improvement correlated linearly with plasma concentration"), ("Source", "PMID 9357394")],
         "analysis": [("What did the trial establish? ", "The study characterized linear PK and an exposure-response relationship for FEV1 at 120 mg. ", "It demonstrates pharmacological activity in the studied population. ", "It does not establish the current India-approved dose or superiority to modern controller strategies."),
                      ("What influenced response variability? ", "Initial disease severity influenced the modeled concentration-effect slope. ", "This suggests baseline physiology may affect observed response. ", "The model is not a validated patient-selection tool.")],
         "interpretation": "The study is the principal formal PK/PD trial. It supports mechanistic and dose-development understanding but must not be used as a stand-alone prescribing guide."},
        {"title": "Placebo-Controlled Sputum and Mucociliary Study", "question": "Can TP-receptor blockade improve airway secretion characteristics even when conventional lung-function changes are limited?",
         "details": [("Design", "Multicentre double-blind randomized placebo-controlled study"), ("Population", "45 evaluable patients with mild-to-moderate asthma and >20 g/day sputum"), ("Intervention", "Seratrodast 40 mg/day"), ("Comparator", "Placebo"), ("Duration", "6 weeks after 2-week run-in"), ("Key findings", "No significant FEV1 or PEF difference; reductions in PEF variability, daytime symptoms, rescue beta2-agonist use, sputum amount/viscosity, and improved nasal clearance time"), ("Source", "PMID 10893362")],
         "analysis": [("What makes this study clinically interesting? ", "It separates airway secretion and symptom effects from large changes in standard pulmonary-function measures. ", "This may be relevant to selected patients with prominent sputum symptoms. ", "The selected population was small and the findings should not be generalized to all asthma."),
                      ("What should be stated with equal prominence? ", "FEV1 and PEF changes were not significantly different between groups. ", "This limitation prevents an exaggerated lung-function narrative. ", "Positive secondary findings must remain contextualized by multiplicity and study size.")],
         "interpretation": "The trial supports selected secretion, symptom, and mucociliary effects, while also demonstrating minimal effect on key pulmonary-function measures."},
        {"title": "Airway Hyperresponsiveness and Inflammatory Markers Study", "question": "Does seratrodast alter airway responsiveness independently of common eosinophilic markers?",
         "details": [("Design", "Small prospective study"), ("Population", "14 patients with asthma"), ("Intervention", "Seratrodast 80 mg once daily"), ("Duration", "4 weeks"), ("Findings", "Airway hyperresponsiveness decreased; exhaled nitric oxide and sputum eosinophil percentage did not change significantly"), ("Source", "PMID 9844991")],
         "analysis": [("What did the study suggest? ", "Airway responsiveness improved without parallel change in the measured eosinophilic markers. ", "The result is consistent with a non-eosinophil-specific pathway effect. ", "The small uncontrolled design limits causal and clinical interpretation."),
                      ("How should this influence positioning? ", "The study may help explain a mechanistically distinct add-on rationale. ", "It does not show replacement of inhaled corticosteroids or establish a biomarker-defined treatment pathway. ", "The findings remain supportive and exploratory.")],
         "interpretation": "This study enriches mechanistic understanding but is not adequate for a broad efficacy or anti-inflammatory claim."},
        {"title": "Thromboxane-Metabolite Response Study", "question": "Could thromboxane-metabolite levels identify patients more likely to respond?",
         "details": [("Design", "Small clinical biomarker-response investigation"), ("Population", "Patients with bronchial asthma receiving seratrodast"), ("Biomarkers", "TXB2 and 11-dehydrothromboxane B2 in urine and sputum"), ("Finding", "A small subgroup of marked responders had higher baseline urinary 11-dehydrothromboxane B2"), ("Source", "PMID 11517517")],
         "analysis": [("What hypothesis emerged? ", "Higher baseline urinary thromboxane metabolite levels might identify a subgroup with stronger response. ", "This aligns with the molecule's pathway. ", "The observation is based on a very small subgroup and requires prospective validation."),
                      ("Can this be used routinely? ", "No validated threshold, test strategy, or outcome-guided algorithm was established. ", "The finding belongs in a future-research discussion. ", "It must not be presented as a diagnostic or prescribing biomarker.")],
         "interpretation": "The study suggests a precision-medicine hypothesis but does not provide a clinically validated selection tool."},
        {"title": "Indian Double-Blind Comparison With Montelukast", "question": "How did seratrodast compare with montelukast as add-on therapy in Indian adults with mild-to-moderate asthma?",
         "details": [("Design", "Randomized comparative double-blind double-dummy multicentre parallel-group non-inferiority study"), ("Population", "205 Indian adults with mild-to-moderate asthma continuing lowest-dose inhaled corticosteroid"), ("Interventions", "Seratrodast 80 mg once daily, n=103; montelukast 10 mg once daily, n=102"), ("Duration", "28 days"), ("Outcomes", "Asthma symptom scores, PEF, FVC, FEV1, sputum fucose, ECP, albumin, and safety"), ("Key interpretation", "Both groups improved from baseline; selected PEF and sputum endpoints favored seratrodast"), ("Source", "IMSEAR open-access article; CDSCO NDAC record")],
         "analysis": [("Why is this study central to an India monograph? ", "It directly evaluates seratrodast in Indian patients, on background inhaled corticosteroid therapy, against a relevant oral add-on comparator. ", "This gives local scientific context to the CDSCO add-on indication. ", "The short duration and endpoint set limit long-term and severe-outcome conclusions."),
                      ("How should comparative findings be communicated? ", "Results should be presented endpoint by endpoint with absolute values, denominators, and statistical framework. ", "The study can support defined comparative statements. ", "It must not be converted into a broad claim that seratrodast is superior or generally favored against montelukast.")],
         "interpretation": "This is the most important India comparative evidence identified, but its 28-day duration requires very careful fair-balance language."},
        {"title": "Japanese Regulatory Clinical Experience", "question": "What does the broader Japanese approval experience contribute?",
         "details": [("Population", "Japanese clinical studies included 514 patients with bronchial asthma"), ("Regimens", "Principally 80 mg once daily for 8 weeks; elderly patients could begin at 40 mg in Japanese information"), ("Reported outcome", "Moderate improvement or better reported in 251/514 patients, 48.8%"), ("Safety context", "Preapproval liver-test abnormalities and post-marketing serious hepatic events informed regulatory warnings"), ("Sources", "Japanese product information; PMDA safety bulletin")],
         "analysis": [("What does the regulatory dataset add? ", "It broadens the exposure base beyond individual published studies and documents the historical basis for Japanese use. ", "It is important for understanding efficacy and safety context. ", "The outcome definitions and historical treatment context may not align with contemporary trial standards."),
                      ("Why is the safety history especially important? ", "The regulatory record links preapproval liver-test abnormalities with serious post-marketing hepatic events. ", "This elevates liver safety from a routine adverse-event topic to a central monitoring concern. ", "India-specific requirements must still be verified.")],
         "interpretation": "Japanese regulatory experience is essential scientific context, particularly for hepatic safety, but it must remain jurisdiction-qualified."},
    ]
    idx = 10
    for study in studies:
        study_profile(doc, idx, study)
        idx += 1
    remaining = [
        ("16", "Safety Profile and the Hepatic-Risk Signal", "A seratrodast monograph that does not foreground liver safety is incomplete.", [
            ("What did preapproval experience show? ", "PMDA safety information reports liver-test abnormalities in 28 of 824 patients, or 3.4%, during preapproval trials. ", "This provides a quantitative signal from the regulatory dataset. ", "The denominator and jurisdiction must accompany the number."),
            ("What emerged after marketing? ", "Post-marketing reports included serious hepatic dysfunction, jaundice, fulminant hepatitis, and a fatal case in which seratrodast involvement could not be excluded. ", "The seriousness of these events justifies prominent monitoring and patient-counselling discussion. ", "Case reports do not by themselves quantify incidence or prove causality in every case."),
            ("What monitoring is described in Japan? ", "Japanese information requires liver-function testing once monthly. ", "This is a clear risk-mitigation instruction in that jurisdiction. ", "Whether identical monitoring is mandated in India must be verified from the current India package insert."),
        ]),
        ("17", "Other Adverse Reactions, Interactions, and Special Populations", "The safety profile extends beyond the headline hepatic signal.", [
            ("Which other reactions are reported in Japanese information? ", "Reported reactions include rash, pruritus, gastrointestinal symptoms, hematologic changes, dizziness, drowsiness, palpitations, malaise, flushing, and edema, among others. ", "These observations inform clinical vigilance and counselling. ", "India frequencies and approved wording remain unverified."),
            ("Which interaction concerns are described? ", "Japanese information advises observation with medicines associated with hemolytic anemia and describes in-vitro displacement of protein binding by aspirin. ", "These signals support careful medication review. ", "In-vitro findings do not automatically establish a clinically important interaction."),
            ("Which populations require caution? ", "Japanese information highlights hepatic impairment, pregnancy, lactation, lack of pediatric trials, and increased exposure in older adults. ", "These considerations should shape risk assessment. ", "Final India recommendations require the India-approved label."),
        ]),
        ("18", "Current Asthma Guidelines and Place in Therapy", "A modern monograph must place historical molecule evidence inside contemporary asthma care.", [
            ("What do contemporary guidelines prioritize? ", "Current international asthma strategies emphasize inhaled corticosteroid-containing treatment, reliever strategy, risk reduction, adherence, inhaler technique, and phenotype-directed escalation. ", "This framework should remain the foundation of treatment selection. ", "Seratrodast is not established as a mainstream preferred pathway in the current GINA strategy."),
            ("What does that mean for seratrodast? ", "The verified India role is add-on therapy in bronchial asthma, which may be considered only within local approval and individualized care. ", "A distinct mechanism can be scientifically relevant when a defined residual need exists. ", "Absence from mainstream guidelines requires restrained positioning, not dismissal or exaggeration."),
            ("Which patients might be unsuitable? ", "Patients needing rapid relief, those with active or significant hepatic disease, those without optimized foundational therapy, and those for whom evidence-based alternatives better address future risk may require another strategy. ", "Explicit unsuitable-patient discussion improves safe selection. ", "Exact contraindications must still follow the India package insert."),
        ]),
        ("19", "Practical HCP Assessment and Monitoring Framework", "The practical framework should help the HCP define a reason to start, a reason to continue, and a reason to stop.", [
            ("What should be assessed before initiation? ", "Confirm asthma diagnosis, severity, control, exacerbation risk, inhaler technique, adherence, background inhaled corticosteroid-containing therapy, liver status, concomitant medicines, and the locally approved indication. ", "This prevents an add-on medicine from substituting for foundational assessment. ", "It is not a replacement for the package insert or clinical judgment."),
            ("What should be monitored? ", "Clinical response should include symptoms, reliever use, lung function where appropriate, exacerbations, tolerability, and liver-function monitoring according to the approved local requirements. ", "Monitoring should test whether the intended add-on objective is being met. ", "A lack of meaningful benefit should prompt reassessment."),
            ("What should the patient understand? ", "Patients should know that seratrodast is not a rescue medicine, should continue prescribed controller therapy, should recognize worsening asthma, and should report symptoms potentially suggestive of hepatic injury. ", "Education is a direct safety intervention. ", "Final counselling wording must be approved and label-aligned."),
        ], ["Verify optimized inhaled corticosteroid-containing treatment.", "Review liver history and baseline testing requirements.", "Reconcile interacting and concomitant medicines.", "Define the intended add-on treatment objective.", "Explain that the medicine is not for acute attacks.", "Review response and liver monitoring according to the approved India label."]),
        ("20", "Fair Comparison With Montelukast and Other Controllers", "Fair comparison clarifies evidence without manufacturing a winner.", [
            ("What direct comparison exists? ", "The Indian double-blind study directly compared seratrodast with montelukast over 28 days on background low-dose inhaled corticosteroid therapy. ", "It permits discussion of the studied symptom, lung-function, sputum, and safety endpoints. ", "Its short duration prevents broad comparison of long-term control or exacerbation outcomes."),
            ("What comparison is not valid? ", "Cross-trial comparisons with inhaled corticosteroids, LABAs, LAMAs, biologics, or other controller strategies cannot establish relative efficacy. ", "Different populations, endpoints, and treatment contexts make such rankings unreliable. ", "Seratrodast must not be presented as replacing guideline-preferred controller therapy."),
            ("How should mechanism be compared? ", "Seratrodast blocks TP-receptor signaling, while montelukast blocks cysteinyl-leukotriene signaling and inhaled corticosteroids act broadly on airway inflammation. ", "Mechanistic comparison can improve understanding. ", "Mechanistic difference alone does not establish superior clinical benefit."),
        ]),
        ("21", "Evidence Gaps and Research Priorities", "The evidence gaps are substantial and should be treated as part of the product's scientific identity.", [
            ("Which outcomes remain uncertain? ", "Large, modern studies demonstrating reduction in severe exacerbations, hospitalization, mortality, long-term oral corticosteroid exposure, or sustained lung-function decline were not identified. ", "This limits outcome-based positioning. ", "Selected symptom or biomarker improvements must not fill these gaps."),
            ("What safety evidence is needed? ", "Current India post-marketing safety data, hepatic-event incidence, risk factors, monitoring effectiveness, and outcomes after discontinuation would strengthen benefit-risk assessment. ", "The known hepatic signal makes this especially important. ", "Absence of public data should not be interpreted as absence of risk."),
            ("What would modern clinical research look like? ", "A contemporary programme could evaluate phenotype-informed selection, longer-term add-on efficacy, exacerbations, patient-reported outcomes, hepatic monitoring, and real-world India effectiveness. ", "Such studies could clarify whether a distinct responder population exists. ", "Until then, positioning must remain conservative."),
        ]),
        ("22", "Balanced Scientific Conclusions", "The conclusion must preserve clinical interest and clinical restraint at the same time.", [
            ("What is established? ", "Seratrodast is an oral TP-receptor antagonist with a verified India approval-list role as add-on therapy in bronchial asthma. ", "Clinical studies support activity across selected symptom, airway-response, secretion, and lung-function measures. ", "The evidence is older and narrower than that supporting mainstream contemporary asthma pathways."),
            ("What is clinically distinctive? ", "The molecule targets thromboxane-mediated airway biology and may offer an adjunctive option in locally approved circumstances. ", "Its mechanism and India comparative study provide legitimate scientific interest. ", "It must not be portrayed as a rescue bronchodilator, monotherapy, or guideline-preferred controller."),
            ("What must remain prominent? ", "Serious hepatic injury and Japanese monthly liver-function monitoring are central safety considerations, while India-specific wording remains to be verified. ", "A balanced monograph should help HCPs assess both potential role and risk. ", "External use remains blocked until the current India package insert and required approvals are incorporated."),
        ]),
    ]
    for args in remaining:
        chapter(doc, *args)
    add_evidence_methodology_annex(doc, "seratrodast")
    add_hcp_faq_annex(doc, "seratrodast", {
        "one_line": "Seratrodast is an oral thromboxane A2/prostaglandin H2 receptor antagonist with a verified India approval-list role as add-on therapy in bronchial asthma. Its evidence supports selected physiological and clinical effects, while major long-term outcomes remain unestablished.",
        "role": "CDSCO records identify 40-mg and 80-mg tablets as add-on therapy in bronchial asthma. This narrow role should remain central until the current India package insert defines exact dosing, contraindications, warnings, and patient boundaries.",
        "distinction": "The molecule targets TP-receptor-mediated airway effects, providing a mechanistically distinct oral add-on concept. This distinction is scientifically relevant but does not make seratrodast a rescue bronchodilator or a replacement for inhaled corticosteroid-containing therapy.",
        "efficacy": "The evidence base includes a phase II PK/PD study, placebo-controlled airway-secretion work, an airway-hyperresponsiveness study, Japanese regulatory clinical experience, and an Indian double-blind comparison with montelukast. Effects are most consistently described for selected symptoms, sputum, and physiological measures.",
        "india": "The Indian study randomized 205 adults on low-dose inhaled corticosteroid therapy to seratrodast 80 mg or montelukast 10 mg for 28 days. Both groups improved from baseline, with selected PEF, expectoration, and sputum-marker outcomes favoring seratrodast.",
        "international": "International evidence includes a randomized phase II exposure-response study and small placebo-controlled or exploratory studies of sputum, mucociliary clearance, airway hyperresponsiveness, and biomarkers. Most studies are older, small, or short.",
        "safety": "The most important safety issue is serious hepatic injury. Japanese regulatory information reports liver-test abnormalities, serious hepatic dysfunction, jaundice, fulminant hepatitis, and fatal cases, making liver safety central to benefit-risk assessment.",
        "monitoring": "Japanese information requires monthly liver-function monitoring, but the exact India monitoring requirement must be verified. Monitoring must also address asthma control, reliever use, lung function where appropriate, tolerability, and whether the defined add-on objective is being achieved.",
        "never_claim": "The evidence does not support claims that seratrodast replaces inhaled corticosteroids, treats acute attacks, is a preferred global controller, or reduces severe exacerbations, hospitalization, mortality, or long-term corticosteroid exposure.",
        "mechanism": "Seratrodast competitively antagonizes TP receptors and inhibits thromboxane-mediated airway effects. Mechanistic plausibility helps explain selected effects on airway responsiveness and secretions, but it does not establish comprehensive asthma control or long-term outcomes.",
        "pk": "Japanese information describes a delayed Tmax of approximately 7.4 hours and a terminal half-life of approximately 25 hours, which reinforces that seratrodast is not a rapid-relief medicine. Older adults may have higher exposure and longer half-life.",
        "hierarchy": "The Indian comparative study and randomized phase II study should carry more weight than small uncontrolled biomarker studies. Japanese regulatory experience adds breadth and safety context, while recent preclinical repurposing findings remain investigational.",
        "comparison": "Comparison with montelukast should reproduce the 28-day design, background inhaled corticosteroid therapy, individual endpoints, numerical results, and limitations. It must not become a generalized superiority or guideline-ranking claim.",
        "limitations": "The evidence is generally old, short, and limited in sample size. Robust exacerbation, hospitalization, mortality, long-term control, and contemporary comparative evidence were not identified, and a seratrodast-specific systematic review was not located.",
        "long_term": "Controlled long-term evidence is limited. Japanese regulatory experience adds exposure context, but modern long-duration randomized evidence assessing exacerbations, liver safety, and sustained benefit remains an important gap.",
        "elderly": "Japanese PK information describes higher exposure and longer half-life in older adults and suggests cautious dosing. India-specific elderly instructions must come from the current package insert, with liver status and polypharmacy carefully reviewed.",
        "renal": "Adequate renal-impairment PK and dosing evidence was not identified. Renal recommendations should not be invented or inferred from general pharmacology.",
        "hepatic": "Hepatic safety is the central special-population concern. Patients with hepatic impairment require careful assessment, and Japanese regulatory information describes serious injury and monthly monitoring; exact India restrictions and actions remain label-controlled.",
        "pregnancy": "Adequate human pregnancy and lactation evidence was not identified. Japanese information uses a benefit-risk approach and notes animal transfer into milk, but India-approved wording must control local use.",
        "pediatric": "Japanese information states that pediatric clinical trials have not been conducted. Pediatric use should not be inferred from adult asthma biology or tablet strengths.",
        "concomitant": "Review inhaled corticosteroid-containing therapy, reliever treatment, systemic corticosteroids, medicines associated with hepatic injury or hemolytic anemia, and the overall asthma plan. Steroid reduction must never be abrupt or based solely on adding seratrodast.",
        "selection": "A potential patient must be within the verified add-on role, have foundational asthma management optimized, have liver risk assessed, and have a clear residual treatment objective. Selection should be individualized and periodically reviewed.",
        "alternative": "Another strategy may be required for acute attacks, severe or unstable asthma, significant hepatic disease, unoptimized inhaled therapy, or when a guideline-supported pathway better addresses exacerbation risk or phenotype.",
        "before": "Before treatment, confirm asthma diagnosis and control, review inhaler technique and adherence, verify background controller therapy, assess liver history and testing requirements, review concomitant medicines, and define the add-on objective.",
        "after": "Review symptoms, reliever use, lung function where appropriate, exacerbations, tolerability, and liver-function monitoring according to local requirements. Continue only when benefit and safety remain clinically acceptable.",
        "counselling": "Patients should understand that seratrodast is not for acute attacks, should continue prescribed controller and reliever therapy, should report worsening asthma, and should promptly report symptoms potentially suggestive of hepatic injury.",
        "qol": "The available studies emphasize symptoms, sputum, and physiological outcomes more than validated long-term quality-of-life outcomes. Any patient-experience discussion should remain specific to the measured endpoints.",
        "outcomes": "Evidence for severe-exacerbation reduction, hospitalization, mortality, sustained oral-steroid reduction, or prevention of lung-function decline was not identified. These outcomes must not be implied from symptom or biomarker findings.",
        "statistics": "Selected significant secondary endpoints should be interpreted alongside neutral FEV1 or PEF findings, multiple comparisons, sample size, and short duration. Statistical significance does not establish broad clinical superiority.",
        "neutral": "The sputum study found minimal pulmonary-function effect despite positive secretion and symptom findings, and the hyperresponsiveness study found no significant change in exhaled nitric oxide or sputum eosinophils. These neutral results are essential fair balance.",
        "pv": "Pharmacovigilance is central because serious hepatic events emerged after marketing. The monograph must reconcile current India safety data, sponsor reports, regulator communications, and monitoring requirements before external use.",
        "guidelines": "Current GINA strategy emphasizes inhaled corticosteroid-containing therapy and does not establish seratrodast as a mainstream global pathway. The local add-on approval can be discussed without implying guideline-preferred status.",
        "foreign": "Japanese dosing and safety information can support scientific understanding and risk identification only when clearly labelled by jurisdiction. It cannot define India-facing prescribing instructions without confirmation.",
        "external": "Before external use, obtain the current India-approved package insert, exact hepatic monitoring and contraindication language, sponsor product particulars, updated India pharmacovigilance data, verified study numbers, and all required approvals.",
        "study_table": "Reviewers should verify population, background inhaled corticosteroid therapy, dose, comparator, duration, endpoint hierarchy, absolute results, p-values, adverse events, hepatic findings, and limitations against the original publication.",
        "future": "Priority research includes long-term India effectiveness and safety, hepatic-risk characterization, validated responder biomarkers, exacerbation outcomes, contemporary comparator studies, and a formally conducted systematic review.",
        "update": "New evidence should be added only after source verification, assessment of label relevance, pharmacovigilance review, claim-impact analysis, and controlled approval. Emerging preclinical findings should remain clearly investigational.",
        "conclusion": "Seratrodast is a mechanistically distinct, locally approved oral add-on option with evidence for selected asthma symptoms and airway measures. Its place is constrained by limited contemporary outcome evidence and a serious hepatic-safety signal requiring prominent assessment.",
    })
    add_scientific_glossary(doc, "seratrodast", [
        ("TP receptor", "The thromboxane A2/prostaglandin H2 receptor through which thromboxane-related signaling can affect smooth-muscle tone, platelets, vascular responses, and airway biology. Seratrodast acts as an antagonist at this receptor."),
        ("Airway hyperresponsiveness", "An exaggerated airway-narrowing response to stimuli. Improvement in a provocation measure can support physiological activity but is not identical to improvement in symptoms, exacerbations, or long-term outcomes."),
        ("Mucociliary clearance", "The coordinated movement of airway mucus by cilia. Selected seratrodast evidence describes changes in sputum characteristics and nasal clearance time in a small, specifically selected study population."),
        ("Add-on asthma therapy", "Treatment used in addition to foundational asthma management. It must not be interpreted as a rescue medicine or as permission to stop inhaled corticosteroid-containing therapy."),
        ("Hepatic safety signal", "A body of evidence suggesting possible serious liver injury, including post-marketing reports and regulatory action. The signal requires prominent assessment, monitoring, and jurisdiction-specific prescribing controls."),
    ])
    add_clinical_review_lenses(doc, "seratrodast")
    add_seratrodast_hepatic_annex(doc)
    add_prescribing_placeholder(doc, "seratrodast", [
        "Exact India-approved dose, meal timing, maximum dose, duration, and elderly instructions.",
        "India contraindications and hepatic-impairment wording.",
        "India liver-function monitoring requirement and action thresholds.",
        "India-approved adverse-reaction frequencies, interactions, pregnancy, lactation, pediatric, and geriatric wording.",
        "Current India MAH, brand, pack, storage, legal category, and revision date.",
    ])
    add_claim_matrix(doc, "seratrodast", [
        ("SE-01", "Seratrodast is an oral TP-receptor antagonist.", "PMID 9357394; J-STAGE discovery review", "Mechanism does not establish superiority.", "Scientific supported"),
        ("SE-02", "CDSCO lists seratrodast 40 mg and 80 mg as add-on therapy in bronchial asthma.", "CDSCO approved-drug list", "Exact current package-insert wording still required.", "Scientific supported"),
        ("SE-03", "A small placebo-controlled study reported selected sputum and symptom improvements with minimal pulmonary-function effect.", "PMID 10893362", "Small selected population and short duration required.", "Scientific supported"),
        ("SE-04", "Japanese safety information identifies serious hepatic injury and monthly liver-function monitoring.", "PMDA safety bulletin", "Jurisdiction-qualified; India wording unverified.", "Scientific supported"),
        ("SE-05", "Seratrodast is superior to or should replace standard controller therapy.", "No adequate support", "Prohibited.", "Prohibited"),
        ("SE-06", "Seratrodast reduces severe exacerbations, hospitalization, or mortality.", "No adequate outcome evidence identified", "Unsupported.", "Prohibited"),
    ])
    add_source_register(doc, SERATRODAST_REFS)
    add_heading(doc, "Appendix: Internal Review Questions", 1)
    add_bullets(doc, [
        "Is the India add-on indication reproduced narrowly and accurately?",
        "Are Japanese dose and safety statements visibly jurisdiction-qualified?",
        "Is the hepatic-risk signal prominent in the executive summary, safety chapter, practical monitoring, and claim matrix?",
        "Are positive secondary endpoints balanced by study size, duration, and pulmonary-function limitations?",
        "Is seratrodast clearly distinguished from rescue therapy and guideline-preferred foundational controller therapy?",
        "Has every numerical result been independently reproduced from the original source?",
    ])
    path = OUT / "PM-EXP-Q2_Seratrodast_Expanded_HCP_Scientific_Monograph_R2.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_trelagliptin())
    print(build_seratrodast())
