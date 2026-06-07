from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "product-monograph-qualification"
OUT.mkdir(parents=True, exist_ok=True)

NAVY, TEAL, PALE, GRAY, WHITE, AMBER = "18324A", "137C7B", "E5F1F0", "F2F4F6", "FFFFFF", "F6E8BE"

products = [
{
"code":"PM-Q1","name":"Trelagliptin","subtitle":"Scientific Product Monograph - Qualification Draft",
"control":"Exact current India-approved prescribing information was not supplied. An India approval listing was identified, while Japan regulatory information and peer-reviewed evidence provide the detailed scientific context. All India prescribing statements require verification against the current approved package insert before use.",
"identity":[("Generic / INN","Trelagliptin; regulatory materials commonly identify trelagliptin succinate"),("Pharmacological class","Dipeptidyl peptidase-4 (DPP-4) inhibitor"),("Therapeutic area","Type 2 diabetes mellitus"),("Differentiating characteristic","Oral once-weekly dosing"),("Verified regulatory context","Approved in Japan; an India approval listing was identified, but the current India package insert was not retrieved"),("Qualification use","Internal scientific draft only")],
"summary":"Trelagliptin is a long-acting oral DPP-4 inhibitor developed for once-weekly administration in type 2 diabetes mellitus. Its principal practical distinction is reduced dosing frequency. Evidence includes Japanese phase 3 studies and a published Indian phase 3 non-inferiority study. Clinical use must remain governed by the current local approved label.",
"need":"Type 2 diabetes requires durable glycemic management and sustained treatment engagement. A once-weekly oral option may reduce dosing burden for selected patients, but improved adherence or outcomes must not be claimed without direct evidence.",
"moa":"DPP-4 inhibition prolongs endogenous incretin activity, supporting glucose-dependent insulin secretion and reducing inappropriate glucagon secretion. Mechanistic explanation must not be extended into unsupported outcome claims.",
"pk":[("Administration frequency","Once weekly in verified Japan regulatory material"),("Usual Japan adult dose","100 mg orally once weekly"),("Moderate renal impairment","Japan safety revision describes 50 mg once weekly"),("Renal handling","Primarily excreted through the kidneys; renal function materially affects exposure"),("Local-use boundary","Use only the effective local label for dose, missed-dose instructions, contraindications, and precautions")],
"indication":"Verified Japan context: treatment of type 2 diabetes mellitus. The precise India-approved indication, population, combination use, and restrictions must be confirmed from authoritative local prescribing information.",
"dose":"Do not use this qualification draft for prescribing. Japan sources describe 100 mg once weekly and a reduced 50 mg once-weekly dose in moderate renal impairment. Exact local dose, timing, missed-dose management, and restrictions require current India-label verification.",
"safety":["Hypoglycemia risk requires attention, particularly with insulin secretagogues or insulin, according to applicable class/local-label controls.","Renal impairment affects exposure and may require dose adjustment or avoidance according to the effective label.","Class-relevant safety topics and product-specific adverse reactions must be taken from the current approved local label.","Once-weekly dosing requires clear education to prevent duplicate dosing or incorrect switching."],
"trials":[("Japanese phase 3 non-inferiority study","Randomized, double-blind, double-dummy; trelagliptin 100 mg weekly versus alogliptin 25 mg daily and placebo; 24 weeks","Trelagliptin showed similar efficacy and safety to alogliptin in the studied Japanese population","Generalization and local-label use require caution"),("Indian phase 3 non-inferiority study","Randomized, open-label, active-controlled; 240 treatment-naive patients; trelagliptin weekly versus vildagliptin twice daily; 16 weeks","Provides locally relevant comparative efficacy and safety evidence","Open-label design and duration should be considered"),("Switch study","Open-label exploratory phase 3; patients switched from daily sitagliptin to weekly trelagliptin","Supports understanding of switching in the studied Japanese population","Exploratory design; not a universal switching recommendation")],
"place":"Potential role as a once-weekly oral DPP-4 option for appropriately selected adults with type 2 diabetes when permitted by the local label and clinical judgment. It should not be positioned as superior on adherence, outcomes, or safety without substantiation.",
"sources":[("PMDA","Zafatek / trelagliptin deliberation report","https://www.pmda.go.jp/files/000213963.pdf"),("PMDA","Trelagliptin dosage and renal-impairment safety revision","https://www.pmda.go.jp/files/000231402.pdf"),("PubMed PMID 25609193","Once-weekly trelagliptin versus daily alogliptin","https://pubmed.ncbi.nlm.nih.gov/25609193/"),("PubMed PMID 40605903","Trelagliptin versus vildagliptin in Indian patients","https://pubmed.ncbi.nlm.nih.gov/40605903/"),("PubMed PMID 28836351","Switching from daily DPP-4 inhibitor to trelagliptin","https://pubmed.ncbi.nlm.nih.gov/28836351/")],
"claims":[("TR-01","Trelagliptin is an oral once-weekly DPP-4 inhibitor","PMDA deliberation report; PMID 25609193","Scientific draft: supported"),("TR-02","Trelagliptin demonstrated non-inferiority to alogliptin in the studied Japanese phase 3 population","PMID 25609193","Scientific draft: supported with population qualifier"),("TR-03","Published Indian phase 3 evidence compared weekly trelagliptin with twice-daily vildagliptin","PMID 40605903","Scientific draft: supported"),("TR-04","Trelagliptin improves adherence","No direct source established in this qualification run","Prohibited unless substantiated")],
},
{
"code":"PM-Q2","name":"Seratrodast","subtitle":"Scientific Product Monograph - Qualification Draft",
"control":"A CDSCO record verifies India approval of seratrodast 40 mg and 80 mg tablets as add-on therapy in bronchial asthma, but the current India-approved package insert was not retrieved. Japanese prescribing information and safety material are jurisdiction-specific and must not be treated as India-label wording.",
"identity":[("Generic / INN","Seratrodast"),("Development code","AA-2414"),("Pharmacological class","Thromboxane A2 / prostaglandin H2 receptor antagonist"),("Therapeutic area","Bronchial asthma"),("Verified India regulatory context","CDSCO lists 40 mg and 80 mg tablets, approved 30 November 2012 as add-on therapy in bronchial asthma"),("Qualification use","Internal scientific draft only")],
"summary":"Seratrodast is an orally active thromboxane A2 receptor antagonist approved in India as add-on therapy in bronchial asthma. Evidence is older and narrower than the evidence base supporting contemporary guideline-preferred asthma controllers. Important Japanese safety information includes serious hepatic injury and monthly liver-function monitoring; the exact current India safety wording remains unverified.",
"need":"Asthma control requires anti-inflammatory controller therapy, appropriate reliever strategy, adherence, inhaler technique, trigger management, and risk assessment. Seratrodast addresses thromboxane-mediated bronchoconstrictive and inflammatory pathways, but its role must be interpreted within current local guidance and approved use.",
"moa":"Seratrodast competitively antagonizes thromboxane A2/prostaglandin H2 receptors, reducing thromboxane-mediated airway effects. Mechanistic plausibility does not establish superiority over guideline-preferred controller therapy.",
"pk":[("Route","Oral"),("Japan single-dose study","After 80 mg in healthy men: Tmax approximately 7.4 hours and terminal half-life approximately 25 hours"),("Japan repeat dosing","No meaningful accumulation reported after 80 mg once daily for 7 days"),("Elderly exposure","Japanese information reports higher exposure and longer half-life; use requires careful local-label review"),("Local-use boundary","Exact India dose, contraindications, hepatic/renal precautions, and interactions require the current India package insert")],
"indication":"Verified India approval-list wording: as add-on therapy in bronchial asthma. The exact current India package-insert wording, eligible age group, combinations, restrictions, and posology remain to be verified.",
"dose":"Do not use this qualification draft for prescribing. Japanese information describes 80 mg once daily after the evening meal and a lower elderly starting dose, but these instructions must not be transferred to India. Only the effective India-approved package insert may define dose and administration.",
"safety":["Japanese regulatory safety information reports serious hepatic dysfunction, jaundice, and fulminant hepatitis, and requires monthly liver-function monitoring. The exact current India warning and monitoring wording must be verified.","The complete India adverse-reaction, contraindication, interaction, hepatic, renal, pregnancy, and monitoring profile must be sourced from the current approved package insert.","Older and relatively small studies limit confidence about uncommon or long-term harms.","Seratrodast is not a rapid bronchodilator and must not be presented as a substitute for guideline-recommended inhaled corticosteroid-containing asthma treatment.","Worsening asthma, acute symptoms, or exacerbations require appropriate standard clinical assessment and treatment; this draft provides no acute-use recommendation."],
"trials":[("Population PK/PD study","Phase 2 randomized, double-blind, placebo-controlled; 183 mild-to-moderate asthma patients; placebo, 80 mg, or 120 mg daily for 8 weeks","Reported linear PK and a concentration-related FEV1 response at 120 mg in the studied population","Historical study; exact approved dose and modern treatment context require confirmation"),("Sputum and mucociliary study","Multicenter, double-blind, randomized, placebo-controlled; 45 patients; 40 mg daily or placebo for 6 weeks","Reported improvements in selected symptom, sputum, and mucociliary measures, with minimal pulmonary-function effect","Small, selected population and short duration"),("Seratrodast versus montelukast study","Randomized double-blind comparative non-inferiority study in adults with mild-to-moderate asthma; seratrodast 80 mg evaluated","Provides comparative evidence in the studied population","Source and local-label relevance require full review before claim use")],
"place":"Seratrodast may have a locally approved adjunctive role in asthma, but its current place in therapy must be confirmed against the effective label and contemporary guidelines. The monograph must not imply guideline-preferred status or use for acute asthma.",
"sources":[("CDSCO","Approved-drug list: seratrodast 40 mg and 80 mg, add-on therapy in bronchial asthma","https://cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadApprovalMarketingFDC/LIST-OF-APPROVED-DRUG-2012%281%29.pdf"),("PMDA","Seratrodast hepatic safety bulletin","https://www.pmda.go.jp/safety/info-services/drugs/calling-attention/safety-info/0148.html"),("PubMed PMID 9357394","Population PK/PD analysis in mild-to-moderate asthma","https://pubmed.ncbi.nlm.nih.gov/9357394/"),("PubMed PMID 10893362","Effect on sputum production and physicochemical properties","https://pubmed.ncbi.nlm.nih.gov/10893362/"),("WHO IMSEAR","Comparative clinical trial with montelukast","https://imsear.searo.who.int/items/3683f1dc-4268-464d-9b58-424ae89f4cd3")],
"claims":[("SE-01","Seratrodast is a thromboxane A2 receptor antagonist studied in asthma","PMID 9357394","Scientific draft: supported"),("SE-02","CDSCO lists seratrodast 40 mg and 80 mg as add-on therapy in bronchial asthma","CDSCO approved-drug list","Supported; exact package-insert wording still required"),("SE-03","Seratrodast improved selected sputum and symptom measures in a small placebo-controlled study","PMID 10893362","Scientific draft: supported with limitations"),("SE-04","Japanese safety information identifies serious hepatic injury and monthly liver-function monitoring","PMDA hepatic safety bulletin","Jurisdiction-qualified; verify India wording"),("SE-05","Seratrodast should replace standard controller therapy","No authoritative support established","Prohibited")],
}
]

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); s=pr.find(qn("w:shd"))
    if s is None: s=OxmlElement("w:shd"); pr.append(s)
    s.set(qn("w:fill"),fill)
def cell(cell,text,bold=False,color=None,size=8.5):
    cell.text=""; r=cell.paragraphs[0].add_run(str(text)); r.bold=bold;r.font.name="Aptos";r.font.size=Pt(size)
    if color:r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def table(doc,heads,rows,widths=None,fill=NAVY):
    t=doc.add_table(rows=1,cols=len(heads));t.style="Table Grid";t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    for i,h in enumerate(heads):cell(t.rows[0].cells[i],h,True,WHITE);shade(t.rows[0].cells[i],fill)
    for n,row in enumerate(rows):
        cs=t.add_row().cells
        for i,v in enumerate(row):cell(cs[i],v);shade(cs[i],WHITE if n%2==0 else GRAY)
    doc.add_paragraph()
def heading(doc,text,n=1): doc.add_paragraph(text,style=f"Heading {n}")
def bullets(doc,items):
    for x in items: doc.add_paragraph(x,style="List Bullet")
def setup():
    d=Document();s=d.sections[0];s.top_margin=Inches(.7);s.bottom_margin=Inches(.7);s.left_margin=Inches(.75);s.right_margin=Inches(.75)
    d.styles["Normal"].font.name="Aptos";d.styles["Normal"].font.size=Pt(9.5);d.styles["Normal"].paragraph_format.space_after=Pt(5)
    for n,z,c in [("Title",25,NAVY),("Heading 1",15,NAVY),("Heading 2",11.5,TEAL)]:
        st=d.styles[n];st.font.name="Aptos Display";st.font.size=Pt(z);st.font.bold=True;st.font.color.rgb=RGBColor.from_string(c)
    f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.CENTER;f.add_run("ClinCommand OS | Product Monograph Qualification Draft | Dr. Bhupesh Dewan").font.size=Pt(7)
    return d
def build(p):
    d=setup();q=d.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_before=Pt(45);q.add_run("CLINCOMMAND OS").bold=True
    x=d.add_paragraph(style="Title");x.alignment=WD_ALIGN_PARAGRAPH.CENTER;x.add_run(f"Product Monograph\n{p['name']}")
    x=d.add_paragraph();x.alignment=WD_ALIGN_PARAGRAPH.CENTER;x.add_run(p["subtitle"]).italic=True
    table(d,["Control Field","Value"],[("Qualification run",p["code"]),("Intended audience","Internal medical and qualified HCP-content reviewers"),("Evidence cut-off","6 June 2026"),("Status","Draft; not approved for prescribing, promotion, or external distribution"),("Controlled package","Owner-supplied Product Monograph SOP and HCP Product Monograph skill consulted")],fill=TEAL)
    x=d.add_paragraph();x.add_run("LABEL AND USE BOUNDARY: ").bold=True;x.add_run(p["control"])
    d.add_page_break()
    heading(d,"1. Executive Clinical Summary");d.add_paragraph(p["summary"])
    table(d,["Quick Reference","Qualified Finding"],p["identity"],fill=TEAL)
    heading(d,"2. Therapeutic Gap and Clinical Context");d.add_paragraph(p["need"])
    heading(d,"3. Product Identity and Molecule Profile");table(d,["Item","Finding"],p["identity"])
    heading(d,"4. Approved Indication and Label Boundaries");d.add_paragraph(p["indication"])
    heading(d,"5. Mechanism of Action and Translational Relevance");d.add_paragraph(p["moa"])
    heading(d,"6. Pharmacology, PK, and Practical Use");table(d,["Parameter","Qualified Finding"],p["pk"])
    heading(d,"7. Dosage and Administration Boundary");d.add_paragraph(p["dose"])
    heading(d,"8. Clinical Evidence Architecture");table(d,["Study","Design / Population","Key finding","Limitation"],p["trials"])
    heading(d,"9. Safety, Tolerability, and Monitoring");bullets(d,p["safety"])
    heading(d,"10. Place in Therapy and Patient Selection");d.add_paragraph(p["place"])
    heading(d,"11. Fair Comparative Framing")
    d.add_paragraph("Comparisons must be factual, like-for-like, source-backed, jurisdiction-appropriate, and balanced. No superiority, adherence, safety, guideline-preference, or outcome claim may be made unless specifically supported and approved.")
    heading(d,"12. Claim Substantiation Matrix");table(d,["Claim ID","Claim","Evidence anchor","Use status"],p["claims"],fill=TEAL)
    heading(d,"13. Required Data Gaps Before Approval")
    bullets(d,["Current authoritative India prescribing information and approval status.","Exact contraindications, warnings, precautions, adverse reactions, interactions, and special-population language.","Current guideline positioning and permitted claims.","Medical, regulatory, legal, compliance, and pharmacovigilance review.","Approved prescribing-information appendix, version, effective date, and distribution classification."])
    heading(d,"14. References and Source Register");table(d,["Source","Title","URL"],p["sources"])
    heading(d,"15. Review and Approval Status");table(d,["Control","Status"],[("SOP and skill consultation","Passed for draft planning; consultation receipt recorded"),("Scientific content","Pending qualified medical review"),("Label verification","Pending current India label"),("Claim-source matrix","Draft complete; reviewer disposition pending"),("Visual page review","Pending"),("External distribution","Prohibited")],fill=TEAL)
    path=OUT/f"{p['code']}_{p['name']}_Product_Monograph.docx";d.save(path);return path
if __name__=="__main__":
    for p in products: print(build(p))
