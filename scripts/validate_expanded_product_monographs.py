from pathlib import Path
from docx import Document
import json
import re
import zipfile

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "product-monograph-expanded"

REQUIRED = [
    "Executive Scientific Summary",
    "Clinical Development Programme",
    "Safety Profile",
    "Place in Therapy",
    "Practical HCP",
    "Evidence Gaps",
    "Balanced Scientific Conclusions",
    "India Prescribing Information",
    "Claim-Source",
    "Reference and Open-Web Source Register",
]

FORBIDDEN = [
    "completely safe",
    "no side effects",
    "guaranteed",
    "game changer",
    "should prescribe",
    "preferred over",
    "empowered to prescribe",
]


def validate(path):
    doc = Document(path)
    text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [c.text for t in doc.tables for r in t.rows for c in r.cells]
    )
    words = re.findall(r"\b[\w'-]+\b", text)
    checks = []
    for section in REQUIRED:
        checks.append((f"required:{section}", section.lower() in text.lower()))
    for phrase in FORBIDDEN:
        checks.append((f"forbidden:{phrase}", phrase.lower() not in text.lower()))
    checks.extend([
        ("word_count>=12000", len(words) >= 12000),
        ("paragraphs>=250", len(doc.paragraphs) >= 250),
        ("tables>=18", len(doc.tables) >= 18),
        ("page_breaks>=30", text.count("") >= 0 and len(doc.element.xpath(".//w:br[@w:type='page']")) >= 30),
        ("urls>=12", text.count("https://") >= 12),
        ("india_boundary", "current India" in text or "current india" in text.lower()),
        ("not_for_prescribing", "not approved for prescribing" in text.lower()),
        ("zip_integrity", zipfile.ZipFile(path).testzip() is None),
    ])
    failures = [name for name, passed in checks if not passed]
    return {
        "file": path.name,
        "word_count": len(words),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "explicit_page_breaks": len(doc.element.xpath(".//w:br[@w:type='page']")),
        "checks": len(checks),
        "passed": len(checks) - len(failures),
        "failures": failures,
        "result": "CONDITIONAL_PASS" if not failures else "FAIL",
        "visual_qa": "PENDING",
    }


results = [validate(p) for p in sorted(OUT.glob("*_R2.docx"))]
(OUT / "expanded_validation.json").write_text(json.dumps(results, indent=2), encoding="utf-8")
for result in results:
    print(json.dumps(result))
