from pathlib import Path
from datetime import datetime, timezone
from hashlib import sha256
from shutil import copy2
from docx import Document
import json
import re

ROOT = Path(__file__).resolve().parents[1]
REPO = ROOT / "controlled-repository"
ORIGINALS = REPO / "originals"
EXTRACTED = REPO / "extracted"
MANIFESTS = REPO / "manifests"

SOURCES = [
    {
        "activity": "product_appraisal",
        "asset_type": "skill",
        "title": "Pharma Product Appraisal Skill Template",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\pharma_product_appraisal_skill_template.md"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_appraisal",
        "asset_type": "sop",
        "title": "Preparation of Product Appraisal",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\Expanded Auro _SOP_Product_Appraisal.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_monograph",
        "asset_type": "sop",
        "title": "Preparation of Product Monograph",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\SOP of Preparation of Product Monograph.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "cme_slides",
        "asset_type": "checklist",
        "title": "CME Slide Preparation Checklist",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\CME slide Preparation- Check List .docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "medical_queries",
        "asset_type": "checklist",
        "title": "Answering Queries Checklist",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Checklist\ANSWERING QUERIES.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "cme_slides",
        "asset_type": "checklist",
        "title": "CME Slides Checklist",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Checklist\Check list CME slides.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_monograph",
        "asset_type": "checklist",
        "title": "Product Monograph Checklist",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Checklist\Checklist Monograph.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "cme_slides",
        "asset_type": "checklist",
        "title": "CME Slide Preparation Checklist Master",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Checklist\CME slide Preparation- Check List .docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_appraisal",
        "asset_type": "checklist",
        "title": "Master Auro Checklist for Product Appraisal",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Checklist\MASTER Auro CHECKLIST FOR PRODUCT APPRAISAL.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_appraisal",
        "asset_type": "checklist",
        "title": "Product Appraisal Checklist",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Checklist\PRODUCT APPRAISAL.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "training_manual",
        "asset_type": "checklist",
        "title": "Training Manual Checklist",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Checklist\Training Manual Checklist.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "promotional_literature",
        "asset_type": "sop",
        "title": "Approval of Promotional Literature",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\SOP of Approval of Promotional Literature.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "package_insert",
        "asset_type": "sop",
        "title": "Preparation of Package Insert",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\SOP of Preparation of  Package Inesrt.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "abpi",
        "asset_type": "sop",
        "title": "Preparation of ABPI",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\SOP of Preparation of ABPI.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_rationale",
        "asset_type": "sop",
        "title": "Preparation of Product Rationale",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\SOP of Product Rationale.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "queries",
        "asset_type": "sop",
        "title": "Answering Queries",
        "source": Path(r"D:\OneDrive\Professional\SOP\SOP of Medico Marketing\Ver 1.0\SOP of Answering Queries.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_monograph",
        "asset_type": "skill",
        "title": "World-Class HCP Product Monograph Skill",
        "source": Path(r"D:\OneDrive\Professional\SOP\Skills\world_class_hcp_product_monograph_skill.md"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "product_training_slides",
        "asset_type": "skill",
        "title": "Pharmaceutical Product Training Slides Skill",
        "source": Path(r"D:\OneDrive\Professional\SOP\Skills\product_training_slides_skill.prompt.md"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
    {
        "activity": "cross_activity_medical_writing",
        "asset_type": "meta_prompt",
        "title": "Meta Prompt for Product Monograph and Appraisal",
        "source": Path(r"D:\OneDrive\Professional\Auro Pharma\META prompt for Product Monograph and Appraisal.docx"),
        "status": "OWNER_SUPPLIED_SOURCE",
    },
]

def slug(value):
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")

def digest(path):
    h = sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()

def extract_docx(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(redact_brand(text))
    for index, table in enumerate(doc.tables, start=1):
        parts.append(f"\n[TABLE {index}]")
        for row in table.rows:
            parts.append(" | ".join(redact_brand(cell.text.strip().replace("\n", " / ")) for cell in row.cells))
    return "\n".join(parts)

def redact_brand(text):
    text = re.sub(r"(?i)zuventus", "Company", text)
    text = re.sub(r"(?i)\bzuv\b", "Company", text)
    text = re.sub(r"(?i)\bZUV\b", "Company", text)
    text = re.sub(r"(?i)\bZUV_", "COMPANY_", text)
    text = re.sub(r"(?i)_ZUV\b", "_COMPANY", text)
    return text

def sanitize_docx(src, dst):
    doc = Document(src)
    for p in doc.paragraphs:
        if p.text:
            p.text = redact_brand(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    cell.text = redact_brand(cell.text)
    for section in doc.sections:
        for attr in ("header", "footer"):
            container = getattr(section, attr)
            for p in container.paragraphs:
                if p.text:
                    p.text = redact_brand(p.text)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            cell.text = redact_brand(cell.text)
    doc.save(dst)

def ingest(source):
    src = source["source"]
    if not src.exists():
        raise FileNotFoundError(src)

    asset_id = f"{source['activity']}--{source['asset_type']}--{slug(source['title'])}"
    original_dir = ORIGINALS / source["activity"] / source["asset_type"]
    extracted_dir = EXTRACTED / source["activity"] / source["asset_type"]
    original_dir.mkdir(parents=True, exist_ok=True)
    extracted_dir.mkdir(parents=True, exist_ok=True)

    original_path = original_dir / f"{asset_id}{src.suffix.lower()}"
    if src.suffix.lower() == ".docx":
        sanitize_docx(src, original_path)
    else:
        copy2(src, original_path)
    original_hash = digest(original_path)

    if src.suffix.lower() == ".docx":
        extracted_text = extract_docx(original_path)
    else:
        extracted_text = original_path.read_text(encoding="utf-8", errors="replace")

    extracted_path = extracted_dir / f"{asset_id}.txt"
    extracted_path.write_text(extracted_text, encoding="utf-8")
    extracted_hash = digest(extracted_path)

    return {
        "assetId": asset_id,
        "activity": source["activity"],
        "assetType": source["asset_type"],
        "title": source["title"],
        "owner": "Dr. Bhupesh Dewan",
        "status": source["status"],
        "sourcePath": str(src),
        "originalRepositoryPath": str(original_path.relative_to(ROOT)).replace("\\", "/"),
        "extractedRepositoryPath": str(extracted_path.relative_to(ROOT)).replace("\\", "/"),
        "originalSha256": original_hash,
        "extractedSha256": extracted_hash,
        "sourceModifiedAt": datetime.fromtimestamp(src.stat().st_mtime, timezone.utc).isoformat(),
        "ingestedAt": datetime.now(timezone.utc).isoformat(),
        "effectiveForExecution": False,
        "approvalRequired": True,
        "supersedesAssetId": None,
    }

def main():
    MANIFESTS.mkdir(parents=True, exist_ok=True)
    records = [ingest(source) for source in SOURCES]
    manifest = {
        "repository": "ClinCommand OS Controlled Activity Source Repository",
        "manifestVersion": "1.0",
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "executionRule": "Agents must consult approved effective activity packages before planning or execution.",
        "assets": records,
    }
    manifest_path = MANIFESTS / "owner-supplied-activity-sources.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(manifest_path)
    for record in records:
        print(record["assetId"], record["originalSha256"])

if __name__ == "__main__":
    main()
