from pathlib import Path
from docx import Document
import json

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "product-appraisal-qualification"

required_headings = [
    "Executive Decision Summary",
    "Product Identification and Assessment Scope",
    "Scientific and Strategic Synopsis",
    "SWOT Assessment",
    "Evidence Landscape",
    "Weighted Appraisal Scorecard",
    "Mandatory Red-Flag Review",
    "Clinical, Safety, and Regulatory Interpretation",
    "Cross-Functional Diligence Plan",
    "Preliminary Go/No-Go Interpretation",
    "Source Register",
    "Qualification and Review Status",
]

required_control_phrases = [
    "not approved for external use",
    "data gaps",
    "Pending qualification review",
    "External use",
    "Prohibited",
]

forbidden_unqualified_phrases = [
    "completely safe",
    "no side effects",
    "guaranteed efficacy",
    "approved in India",
]

def validate(path):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_text = [
        cell.text.strip()
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    text = "\n".join(paragraphs + table_text)
    checks = []

    for heading in required_headings:
        checks.append({
            "check": f"Required section: {heading}",
            "critical": True,
            "passed": heading.lower() in text.lower()
        })
    for phrase in required_control_phrases:
        checks.append({
            "check": f"Control phrase present: {phrase}",
            "critical": True,
            "passed": phrase.lower() in text.lower()
        })
    for phrase in forbidden_unqualified_phrases:
        checks.append({
            "check": f"Forbidden unsupported phrase absent: {phrase}",
            "critical": True,
            "passed": phrase.lower() not in text.lower()
        })

    checks.extend([
        {"check": "Contains at least 10 tables", "critical": False, "passed": len(doc.tables) >= 10},
        {"check": "Contains at least 60 table rows", "critical": False, "passed": sum(len(t.rows) for t in doc.tables) >= 60},
        {"check": "Contains source URLs", "critical": True, "passed": "https://" in text},
        {"check": "Contains red-flag override language", "critical": True, "passed": "numeric score does not override unresolved critical red flags" in text.lower()},
        {"check": "Contains India-specific verification gap", "critical": True, "passed": "india" in text.lower() and "verify" in text.lower()},
    ])

    critical_failures = [c for c in checks if c["critical"] and not c["passed"]]
    all_failures = [c for c in checks if not c["passed"]]
    return {
        "file": path.name,
        "checks": len(checks),
        "passed_checks": len(checks) - len(all_failures),
        "critical_failures": len(critical_failures),
        "all_failures": len(all_failures),
        "content_control_result": "PASS" if not critical_failures else "FAIL",
        "visual_render_result": "UNVERIFIED",
        "activity_qualification_result": "CONDITIONAL_PASS" if not critical_failures else "FAIL",
        "failed_checks": all_failures,
    }

results = [
    validate(OUT / "PA-Q1_Trelagliptin_Product_Appraisal_R2.docx"),
    validate(OUT / "PA-Q2_Empagliflozin_Product_Appraisal_R2.docx"),
    validate(OUT / "PA-Q3_Fevipiprant_Product_Appraisal_R2.docx"),
    validate(OUT / "PA-Q4_Olokizumab_Product_Appraisal_R2.docx"),
]

result_path = OUT / "product_appraisal_automated_validation_results.json"
result_path.write_text(json.dumps(results, indent=2), encoding="utf-8")
for result in results:
    print(json.dumps({
        "file": result["file"],
        "content_control_result": result["content_control_result"],
        "activity_qualification_result": result["activity_qualification_result"],
        "passed_checks": result["passed_checks"],
        "checks": result["checks"],
        "critical_failures": result["critical_failures"],
        "visual_render_result": result["visual_render_result"],
    }))
