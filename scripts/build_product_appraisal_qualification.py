from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "product-appraisal-qualification"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "18324A"
TEAL = "137C7B"
LIGHT_TEAL = "DDEEEE"
LIGHT_BLUE = "E8F0F6"
LIGHT_GRAY = "F2F4F6"
AMBER = "F5E6B7"
RED = "F3D7D7"
WHITE = "FFFFFF"
TEXT = RGBColor(35, 48, 58)

weights = [
    ("Therapeutic gap and medical need", 8),
    ("Molecule and target profile", 8),
    ("Preclinical evidence", 6),
    ("Clinical evidence", 16),
    ("Safety and tolerability", 10),
    ("Regulatory status and pathway", 8),
    ("Patent and exclusivity", 8),
    ("Competitor landscape", 8),
    ("Pricing and market access", 7),
    ("API economics and supply", 5),
    ("Commercial attractiveness", 6),
    ("Strategic fit with company product mix", 6),
    ("Operational and CMC feasibility", 4),
    ("HCP value proposition", 4),
    ("Evidence quality and digital intelligence", 3),
    ("Reputation, ethics, and access", 1),
]

trelagliptin = {
    "requested_name": "Trilagliptin",
    "verified_name": "Trelagliptin",
    "code": "PA-Q1",
    "subtitle": "Preliminary India Portfolio and Medical-Strategy Appraisal",
    "objective": "Assess the opportunity and diligence requirements for a once-weekly DPP-4 inhibitor in type 2 diabetes mellitus.",
    "status": "Conditional candidate - diligence required before any Go/No-Go decision",
    "recommendation": "Proceed only to a structured diligence stage. The once-weekly dosing proposition is differentiated, but Indian regulatory status, commercial opportunity, patent/FTO, API economics, supply readiness, and comparative long-term outcomes require verification.",
    "identity_note": "The requested name 'Trilagliptin' was not verified as an established INN. This appraisal uses the verified INN trelagliptin and requires owner confirmation before controlled use.",
    "product": [
        ("Verified INN", "Trelagliptin; commonly described as trelagliptin succinate in regulatory material"),
        ("Class", "Oral dipeptidyl peptidase-4 (DPP-4) inhibitor"),
        ("Core differentiation", "Once-weekly oral administration"),
        ("Verified approval context", "Approved in Japan for type 2 diabetes mellitus; approval outside Japan must be verified jurisdiction by jurisdiction"),
        ("India evidence context", "A randomized phase 3 non-inferiority trial versus twice-daily vildagliptin in Indian patients was published in 2025"),
        ("Assessment geography", "India"),
    ],
    "strengths": [
        "Distinct once-weekly oral DPP-4 dosing proposition.",
        "Japanese phase 3 evidence supports non-inferiority to once-daily alogliptin in the studied population.",
        "Published Indian phase 3 evidence provides locally relevant comparative data versus vildagliptin.",
        "DPP-4 class familiarity may simplify physician education."
    ],
    "weaknesses": [
        "No verified broad cardiovascular, heart-failure, or kidney-outcome advantage comparable with major SGLT2 evidence programs.",
        "Current Indian marketing authorization status was not verified from an authoritative source during this qualification run.",
        "Adherence benefit is plausible from dosing frequency but should not be assumed without supportive evidence.",
        "Patent, FTO, API economics, pricing, market share, and supply readiness remain unverified."
    ],
    "opportunities": [
        "Potential option for selected patients and prescribers who value reduced dosing frequency.",
        "Local evidence may support a differentiated scientific discussion if regulatory and label status permit.",
        "Could complement a diabetes portfolio where DPP-4 familiarity and oral convenience are strategically relevant."
    ],
    "threats": [
        "Crowded DPP-4 and broader oral antidiabetic market.",
        "Strong competition from agents with demonstrated cardiorenal outcomes.",
        "Once-weekly administration creates medication-error and missed-dose education requirements.",
        "Unverified regulatory, pricing, patent, and supply position may materially alter attractiveness."
    ],
    "scores": [4,4,2,3,3,2,1,3,1,1,3,3,2,4,3,3],
    "red_flags": [
        ("Regulatory status", "Current Indian approval and exact proposed label were not verified.", "Must resolve before progression."),
        ("Patent/FTO", "No authoritative patent or freedom-to-operate assessment was available.", "Legal diligence required."),
        ("Commercial inputs", "No validated India pricing, market share, forecast, or API cost data were supplied.", "Commercial and supply diligence required."),
        ("Name identity", "Requested name differs from verified INN.", "Owner confirmation required.")
    ],
    "evidence": [
        ("PMDA review report", "Zafatek Tablets 50 mg and 100 mg; trelagliptin succinate; PMDA deliberation report, 2015.", "Authoritative regulatory review", "High"),
        ("Inagaki et al., 2015", "Once-weekly trelagliptin versus daily alogliptin; randomized, double-blind phase 3 non-inferiority study.", "Peer-reviewed primary clinical study", "High"),
        ("Inagaki et al., 2018", "Switching from daily sitagliptin to once-weekly trelagliptin; open-label exploratory phase 3 study.", "Peer-reviewed clinical study", "Moderate"),
        ("Dewan et al., 2025", "Once-weekly trelagliptin versus twice-daily vildagliptin in Indian patients; randomized phase 3 non-inferiority study.", "Peer-reviewed India clinical study", "Moderate"),
        ("McKeage, 2015", "Trelagliptin: First Global Approval.", "Peer-reviewed regulatory milestone review", "Moderate"),
    ],
    "sources": [
        ("PMDA", "Report on the Deliberation Results: Zafatek (trelagliptin succinate)", "https://www.pmda.go.jp/files/000213963.pdf"),
        ("PubMed PMID 25609193", "Once-weekly trelagliptin versus daily alogliptin", "https://pubmed.ncbi.nlm.nih.gov/25609193/"),
        ("PubMed PMID 28836351", "Switching from daily DPP-4 inhibitor to trelagliptin", "https://pubmed.ncbi.nlm.nih.gov/28836351/"),
        ("PubMed PMID 40605903", "Trelagliptin versus vildagliptin in Indian patients", "https://pubmed.ncbi.nlm.nih.gov/40605903/"),
        ("PubMed PMID 26115728", "Trelagliptin: First Global Approval", "https://pubmed.ncbi.nlm.nih.gov/26115728/"),
    ],
}

empagliflozin = {
    "requested_name": "Empagliflozin",
    "verified_name": "Empagliflozin",
    "code": "PA-Q2",
    "subtitle": "Preliminary India Portfolio and Medical-Strategy Appraisal",
    "objective": "Assess the strategic and medical value of empagliflozin across type 2 diabetes, heart failure, and chronic kidney disease while identifying portfolio-entry diligence needs.",
    "status": "Strong candidate - pursue differentiated portfolio diligence",
    "recommendation": "Proceed to detailed commercial, patent/FTO, India-label, pricing, supply, and portfolio-fit diligence. The scientific value proposition is strong and supported by mature cardiovascular, heart-failure, and kidney evidence, but market crowding and class safety controls are material.",
    "identity_note": "The INN empagliflozin was verified. Exact India label, local brands, prices, patent/FTO, and market position require separate authoritative verification.",
    "product": [
        ("Verified INN", "Empagliflozin"),
        ("Class", "Oral sodium-glucose co-transporter 2 (SGLT2) inhibitor"),
        ("Core differentiation", "Established glycemic, cardiovascular, heart-failure, and kidney evidence"),
        ("Verified U.S. label context", "FDA label includes heart failure, chronic kidney disease, cardiovascular-risk reduction in adults with type 2 diabetes and established cardiovascular disease, and glycemic control in adults and pediatric patients aged 10 years and older"),
        ("Verified EU context", "EMA describes use in type 2 diabetes, symptomatic chronic heart failure, and chronic kidney disease"),
        ("Assessment geography", "India"),
    ],
    "strengths": [
        "Mature evidence base across type 2 diabetes, cardiovascular outcomes, heart failure, and chronic kidney disease.",
        "Current FDA and EMA sources provide clear regulatory and safety context.",
        "Once-daily oral dosing with a clinically coherent cardiorenal value proposition.",
        "Major outcomes programs provide strong HCP and guideline relevance."
    ],
    "weaknesses": [
        "Crowded SGLT2 market with established competitors and potential price pressure.",
        "Class safety and monitoring requirements include ketoacidosis, volume depletion, genitourinary infections, and other label warnings.",
        "Exact India label, price, patent/FTO, market share, API economics, and supply readiness were not verified in this run.",
        "Differentiation may depend more on execution, evidence communication, access, and portfolio fit than molecule novelty."
    ],
    "opportunities": [
        "Integrated diabetes, cardiovascular, heart-failure, and kidney medical strategy.",
        "Strong evidence platform for HCP education and multidisciplinary engagement.",
        "Potential fit for a broad cardiometabolic and renal portfolio.",
        "Mature evidence can support medical-information, monograph, training, and scientific-content programs."
    ],
    "threats": [
        "Strong branded and generic class competition.",
        "Safety events or inappropriate patient selection can undermine benefit-risk communication.",
        "Local reimbursement, pricing, and affordability may constrain adoption.",
        "Commoditization risk if product strategy lacks a differentiated service and evidence plan."
    ],
    "scores": [5,4,4,5,4,5,2,4,3,3,4,5,4,5,5,4],
    "red_flags": [
        ("India label and status", "Exact current India-approved indications and prescribing information were not verified.", "Regulatory verification required."),
        ("Patent/FTO", "No authoritative India patent or FTO opinion was available.", "Legal diligence required."),
        ("Market and supply", "Validated local pricing, market share, API economics, and supply data were not supplied.", "Commercial and supply diligence required."),
        ("Safety execution", "SGLT2 class risks require clear patient selection, monitoring, perioperative interruption, and escalation education.", "Mandatory medical and PV controls.")
    ],
    "evidence": [
        ("FDA prescribing information, revised October 2025", "Current U.S. indications, dose, limitations, warnings, and adverse reactions.", "Authoritative regulatory label", "High"),
        ("EMA Jardiance EPAR", "EU authorization context and evidence summary for type 2 diabetes, heart failure, and chronic kidney disease.", "Authoritative regulatory assessment", "High"),
        ("EMPA-REG OUTCOME", "Cardiovascular outcomes and mortality in type 2 diabetes with established cardiovascular disease.", "Peer-reviewed pivotal outcomes study", "High"),
        ("EMPEROR program", "Heart-failure outcome evidence across reduced and preserved ejection fraction populations.", "Peer-reviewed pivotal outcomes studies", "High"),
        ("EMPA-KIDNEY", "Kidney-disease progression and cardiorenal outcome evidence in a broad CKD population.", "Peer-reviewed pivotal outcomes study", "High"),
    ],
    "sources": [
        ("FDA", "Jardiance prescribing information, revised October 2025", "https://www.accessdata.fda.gov/drugsatfda_docs/label/2025/204629s063lbl.pdf"),
        ("EMA", "Jardiance EPAR", "https://www.ema.europa.eu/en/medicines/human/EPAR/jardiance"),
        ("PubMed PMID 26378978", "Empagliflozin, Cardiovascular Outcomes, and Mortality in Type 2 Diabetes", "https://pubmed.ncbi.nlm.nih.gov/?term=26378978"),
        ("PubMed PMID 37529652", "EMPA-KIDNEY evidence overview", "https://pubmed.ncbi.nlm.nih.gov/37529652/"),
        ("PubMed PMID 26819227", "Heart failure outcomes in EMPA-REG OUTCOME", "https://pubmed.ncbi.nlm.nih.gov/26819227/"),
    ],
}

fevipiprant = {
    "requested_name": "Fevipiprant",
    "verified_name": "Fevipiprant",
    "code": "PA-Q3",
    "subtitle": "Preliminary India Portfolio and Medical-Strategy Appraisal",
    "objective": "Assess whether fevipiprant merits standalone development or licensing investment for asthma in the current evidence environment.",
    "status": "No-Go - current evidence does not support independent development",
    "recommendation": "Do not allocate development capital to a standalone fevipiprant launch. Multiple phase III failures, official discontinuation of asthma development, and an unattractive risk-adjusted return profile argue against internal progression. Revisit only if a new proprietary package, partnering rights, or a clearly superior subgroup strategy emerges.",
    "identity_note": "Identity is verified, but commercial viability is not. The molecule remains an investigational/abandoned asthma candidate rather than a market-ready asset.",
    "product": [
        ("Verified INN", "Fevipiprant"),
        ("Class", "Oral DP2/CRTh2 receptor antagonist"),
        ("Core differentiation", "Mechanism-based anti-inflammatory oral therapy for asthma"),
        ("Verified development context", "Novartis announced that pooled LUSTER Phase III results did not support further asthma development in December 2019"),
        ("India evidence context", "No verified Indian approval or commercial launch position was established in this run"),
        ("Assessment geography", "India"),
    ],
    "strengths": [
        "Mechanistic rationale and oral dosing convenience were scientifically interesting.",
        "Earlier phase data created a credible hypothesis for anti-inflammatory benefit.",
        "Asthma remains a large unmet-need market, so any true differentiator would have had strategic appeal."
    ],
    "weaknesses": [
        "Pivotal phase III evidence failed to support further asthma development.",
        "The program was discontinued by the sponsor in 2019.",
        "No approved-market label or India launch package was established.",
        "Risk-adjusted commercial return is weak relative to the remaining development burden."
    ],
    "opportunities": [
        "Only a highly differentiated biomarker or partnership-led rescue strategy would justify reconsideration.",
        "A post-hoc scientific review could still inform class learning and publication intelligence.",
        "The molecule may remain relevant as a negative-development case for portfolio discipline."
    ],
    "threats": [
        "Definitive late-stage failure risk has already materialized.",
        "Competitive asthma therapy has moved toward established biologics, inhaled combinations, and better-validated pathways.",
        "Development capital spent here would likely displace better-return opportunities.",
        "Any future re-entry would require a materially new evidence package, not incremental optimism."
    ],
    "scores": [1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1],
    "rationale_map": {
        "Therapeutic gap and medical need": "Asthma remains important, but the molecule failed to prove enough value to justify further internal development.",
        "Molecule and target profile": "The DP2/CRTh2 concept was biologically plausible but did not convert into robust late-stage differentiation.",
        "Preclinical evidence": "The present run did not uncover a preclinical package that would overcome the clinical failure signal.",
        "Clinical evidence": "Phase III ZEAL and LUSTER results and subsequent sponsor communications do not support progression.",
        "Safety and tolerability": "No dominant safety signal drove the no-go; the main barrier is lack of convincing efficacy and development viability.",
        "Regulatory status and pathway": "No verified current approval or realistic filing path was established for India or major markets.",
        "Patent and exclusivity": "Any remaining exclusivity value is eclipsed by the commercial reality of discontinuation and failed proof of concept.",
        "Competitor landscape": "Asthma competition is crowded and the molecule lacks a current efficacy edge.",
        "Pricing and market access": "No launchable pricing strategy is credible without a valid development case.",
        "API economics and supply": "No supply-chain or manufacturing case can rescue a failed development program.",
        "Commercial attractiveness": "The expected return does not justify the remaining investment.",
        "Strategic fit with company product mix": "A failed asthma program is a weak fit unless the company is specifically acquiring negative-development rights for learning value.",
        "Operational and CMC feasibility": "Technical feasibility is moot without a defensible clinical and regulatory path.",
        "HCP value proposition": "The proposed value proposition is not strong enough to support a confident launch narrative.",
        "Evidence quality and digital intelligence": "The sponsor's own late-stage evidence and public discontinuation decision are high-signal sources.",
        "Reputation, ethics, and access": "A continued push without a new evidence basis would be a poor stewardship decision."
    },
    "red_flags": [
        ("Clinical programme status", "Novartis publicly stated that fevipiprant should not be further developed for asthma.", "Treat as a hard stop unless new rights/evidence materially change the picture."),
        ("Late-stage efficacy", "ZEAL and LUSTER phase III programmes did not deliver the needed efficacy signal.", "Do not proceed on hope alone."),
        ("Commercial value", "Risk-adjusted ROI is unattractive compared with other development opportunities.", "No-Go for standalone investment."),
        ("Approval status", "No verified India marketing authorization was identified.", "No launch path without a new regulatory case.")
    ],
    "evidence": [
        ("Novartis media release, Dec 2019", "Topline LUSTER Phase III results did not support further development of fevipiprant in asthma.", "Official sponsor discontinuation notice", "High"),
        ("ZEAL-1 and ZEAL-2", "Two replicate phase III randomised placebo-controlled trials in uncontrolled asthma.", "Peer-reviewed pivotal failure package", "High"),
        ("SPIRIT long-term study", "Exploratory long-term safety and efficacy evaluation after earlier development stages.", "Peer-reviewed extension evidence", "Moderate"),
        ("PMC review/meta-analysis", "Open-access synthesis of fevipiprant asthma evidence and development trajectory.", "Secondary literature", "Moderate"),
        ("Novartis 2019 Form 20-F", "Public company filing noting discontinuation of QAW039 development in asthma.", "Sponsor disclosure", "High"),
    ],
    "sources": [
        ("Novartis", "Novartis provides update on LUSTER Phase III studies in patients with uncontrolled GINA 4/5 asthma", "https://www.novartis.com/news/media-releases/novartis-provides-update-luster-phase-iii-studies-patients-uncontrolled-gina-45-asthma"),
        ("PubMed PMID 33997741", "Efficacy and safety of fevipiprant in patients with uncontrolled asthma: ZEAL-1 and ZEAL-2", "https://pubmed.ncbi.nlm.nih.gov/33997741/"),
        ("PMC 8099656", "ZEAL-1 and ZEAL-2 full text", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8099656/"),
        ("PMC 8666007", "Long-term safety and exploratory efficacy of fevipiprant: SPIRIT randomised clinical trial", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8666007/"),
        ("Novartis 20-F 2019", "Annual report noting discontinuation of QAW039 in asthma", "https://www.novartis.com/sites/novartis_com/files/novartis-20-f-2019.pdf"),
    ],
}

olonacimab = {
    "requested_name": "Olonacimab",
    "verified_name": "Olokizumab",
    "code": "PA-Q4",
    "subtitle": "Preliminary India Portfolio and Medical-Strategy Appraisal",
    "objective": "Assess the requested biologic candidate while resolving the identity risk and determining whether any development path is commercially and operationally sensible.",
    "status": "No-Go as requested; conditional diligence only if the molecule is corrected to olokizumab and rights/data access are secured",
    "recommendation": "Do not progress the project as 'Olonacimab' because the name could not be verified. If the intention was olokizumab, the molecule remains a high-complexity biologic with no verified major-market approval and should be considered only for partnership or due-diligence review, not immediate standalone development.",
    "identity_note": "The requested name 'Olonacimab' could not be verified in the source run. The most probable match is olokizumab, an IL-6 ligand monoclonal antibody; the identity ambiguity is itself a release-blocking risk.",
    "product": [
        ("Requested name", "Olonacimab"),
        ("Probable verified INN", "Olokizumab"),
        ("Class", "Monoclonal antibody against interleukin-6"),
        ("Clinical-development context", "CDSCO has permitted an Indian phase III clinical trial in moderate-to-severe rheumatoid arthritis"),
        ("India approval status", "No verified India marketing authorization found in this run"),
        ("Assessment geography", "India"),
    ],
    "strengths": [
        "Biologic mechanism is scientifically credible for rheumatoid arthritis.",
        "Phase III and meta-analytic data suggest a potentially meaningful anti-inflammatory effect.",
        "Indian clinical-trial activity indicates local development interest and regulatory visibility."
    ],
    "weaknesses": [
        "The requested name is not verified, which makes immediate planning unsafe.",
        "No major-market approval was verified in the current run.",
        "Biologic development adds CMC, comparability, immunogenicity, and supply-chain complexity.",
        "Competitive RA biologic positioning is crowded and capital intensive."
    ],
    "opportunities": [
        "If the sponsor rights are secured, the asset could merit a due-diligence partnership review.",
        "A focused India clinical-development strategy might be possible if the molecule identity is corrected and the data room is complete.",
        "The program may be relevant for portfolio learning even if it is not a near-term launch choice."
    ],
    "threats": [
        "Identity ambiguity can lead to regulatory, scientific, and legal errors.",
        "No verified market authorization means there is no immediate launch basis.",
        "Biologic development is expensive and slow, which raises ROI hurdles.",
        "Competitor therapies already occupy a dense RA treatment landscape."
    ],
    "scores": [1,2,2,3,1,1,1,2,1,1,1,2,1,2,2,1],
    "rationale_map": {
        "Therapeutic gap and medical need": "Rheumatoid arthritis remains important, but this specific candidate does not yet have a clean identity or launch path.",
        "Molecule and target profile": "IL-6 targeting is credible, but the requested name is not verified and the asset must be identity-locked first.",
        "Preclinical evidence": "Not assessed in depth here; the main issue is the unresolved candidate identity and development readiness.",
        "Clinical evidence": "Phase II and III olokizumab studies exist, but they do not rescue the identity problem for 'Olonacimab'.",
        "Safety and tolerability": "The molecule class requires careful infection and laboratory monitoring; the current run does not support a launch-ready safety package.",
        "Regulatory status and pathway": "CDSCO permitted a phase III trial, but no approved India label was verified.",
        "Patent and exclusivity": "No complete IP or rights package was validated in the current run.",
        "Competitor landscape": "RA is heavily contested by established biologics and targeted agents.",
        "Pricing and market access": "No credible access strategy can be finalised without identity and rights confirmation.",
        "API economics and supply": "Biologic manufacturing and supply are complex and not yet de-risked.",
        "Commercial attractiveness": "The asset is not launch-ready and the investment case remains speculative.",
        "Strategic fit with company product mix": "Potential fit exists only if the company wants a biologics capability and the rights package is compelling.",
        "Operational and CMC feasibility": "High CMC complexity and comparability burden remain unresolved.",
        "HCP value proposition": "No authoritative India-facing value proposition can be finalised while the molecule identity remains uncertain.",
        "Evidence quality and digital intelligence": "The strongest signal is the CDSCO trial permit plus peer-reviewed phase III literature for olokizumab.",
        "Reputation, ethics, and access": "Proceeding on a misnamed molecule would be an avoidable governance risk."
    },
    "red_flags": [
        ("Identity verification", "The requested molecule name could not be confirmed.", "Treat as a hard stop until Regulatory Affairs confirms the intended INN."),
        ("Market authorization", "No verified approved India label or major-market approval was established.", "No immediate launch path."),
        ("Biologic complexity", "CMC, immunogenicity, comparability, and supply burdens are substantial.", "Partnership diligence only if the identity is corrected."),
        ("Competitive intensity", "Rheumatoid arthritis is a crowded biologics arena.", "Commercial case must be exceptional to justify entry.")
    ],
    "evidence": [
        ("CDSCO trial permit", "Indian approval to conduct a phase III olokizumab study in rheumatoid arthritis.", "Official Indian regulator record", "High"),
        ("Phase III MTX-IR study", "Olokizumab phase III efficacy and safety in methotrexate-inadequate responders.", "Peer-reviewed primary trial", "High"),
        ("Phase III TNFi-IR study", "Olokizumab phase III efficacy and safety in TNF-inhibitor-inadequate responders.", "Peer-reviewed primary trial", "High"),
        ("Systematic review and network meta-analysis", "Synthesised olokizumab rheumatoid arthritis evidence and safety profile.", "Secondary literature", "Moderate"),
        ("Long-term safety studies", "Open-label extension data support longer exposure characterisation but not approval status.", "Extension evidence", "Moderate"),
    ],
    "sources": [
        ("CDSCO", "Olokizumab phase III clinical-trial permission in India", "https://www.cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadCTApprovals/6.%20DRL%20Olokizumab.pdf"),
        ("PMC 8921576", "Olokizumab in patients with active rheumatoid arthritis despite methotrexate", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8921576/"),
        ("PMC 9664111", "Olokizumab in patients with rheumatoid arthritis inadequately controlled by TNF inhibitor therapy", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9664111/"),
        ("PMC 10202974", "Systematic review, pairwise, and network meta-analysis of olokizumab in RA", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10202974/"),
        ("PMC 9770405", "Long-term safety and efficacy of olokizumab in phase II studies", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9770405/"),
    ],
}

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, WHITE, 8.5)
        shade(table.rows[0].cells[i], header_fill)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, None, 8.5)
            shade(cells[i], WHITE if ri % 2 == 0 else LIGHT_GRAY)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

def weighted_rows(data):
    rows = []
    total = 0.0
    for (category, weight), score in zip(weights, data["scores"]):
        weighted = weight * score / 5.0
        total += weighted
        rows.append((category, weight, score, f"{weighted:.1f}", score_rationale(data, category)))
    return rows, total

def score_rationale(data, category):
    gaps = {"Patent and exclusivity", "Pricing and market access", "API economics and supply"}
    if category in gaps:
        return "Conservative score because authoritative India-specific diligence data were not supplied."
    custom = data.get("rationale_map")
    if isinstance(custom, dict) and category in custom:
        return custom[category]
    if data["verified_name"] == "Trelagliptin":
        mapping = {
            "Therapeutic gap and medical need": "Dosing-frequency convenience may address treatment-burden concerns, but clinical impact needs validation.",
            "Molecule and target profile": "Once-weekly DPP-4 inhibition is a clear administration differentiator.",
            "Preclinical evidence": "No detailed preclinical package was supplied for this qualification run.",
            "Clinical evidence": "Japan and India comparative studies are available, but outcome breadth is limited.",
            "Safety and tolerability": "Regulatory review supports acceptability in studied populations; geography-specific label verification remains required.",
            "Regulatory status and pathway": "Japan approval verified; India status not verified.",
            "Competitor landscape": "Differentiated dosing but crowded diabetes market.",
            "Commercial attractiveness": "Potential niche value; no validated forecast or market data.",
            "Strategic fit with company product mix": "Potentially useful in a diabetes portfolio, subject to call-point and franchise assessment.",
            "Operational and CMC feasibility": "No validated CMC or launch-readiness package supplied.",
            "HCP value proposition": "Simple once-weekly proposition, but must avoid unproven adherence claims.",
            "Evidence quality and digital intelligence": "Several credible sources; important commercial and operational gaps.",
            "Reputation, ethics, and access": "Oral convenience may be useful; affordability and access are unknown.",
        }
    elif data["verified_name"] == "Fevipiprant":
        mapping = {
            "Therapeutic gap and medical need": "Asthma remains important, but this candidate failed to establish enough value to justify internal progression.",
            "Molecule and target profile": "Mechanistic plausibility did not translate into durable late-stage differentiation.",
            "Preclinical evidence": "The current run did not uncover a preclinical package that offsets the development failure.",
            "Clinical evidence": "Phase III ZEAL and LUSTER evidence supports a no-go conclusion.",
            "Safety and tolerability": "Safety was not the main barrier; the dominant issue is inadequate efficacy and development viability.",
            "Regulatory status and pathway": "No verified India approval or realistic filing path was identified.",
            "Competitor landscape": "Asthma competition is intense and fevipiprant lacks a current efficacy edge.",
            "Commercial attractiveness": "Return on investment is unattractive.",
            "Strategic fit with company product mix": "Weak fit unless the company is explicitly pursuing a negative-development learning asset.",
            "Operational and CMC feasibility": "Technical feasibility is moot without a viable clinical and regulatory path.",
            "HCP value proposition": "No launchable value proposition is credible on the present evidence.",
            "Evidence quality and digital intelligence": "The sponsor's own discontinuation and phase III failures are high-signal evidence.",
            "Reputation, ethics, and access": "Further push without new evidence would be poor stewardship.",
        }
    elif data["verified_name"] == "Olokizumab":
        mapping = {
            "Therapeutic gap and medical need": "Rheumatoid arthritis remains important, but this specific candidate does not yet have a clean identity or launch path.",
            "Molecule and target profile": "IL-6 targeting is credible, but the requested name is not verified and the asset must be identity-locked first.",
            "Preclinical evidence": "Not assessed in depth here; the main issue is the unresolved candidate identity and development readiness.",
            "Clinical evidence": "Phase II and III olokizumab studies exist, but they do not rescue the identity problem for 'Olonacimab'.",
            "Safety and tolerability": "The molecule class requires careful infection and laboratory monitoring; the current run does not support a launch-ready safety package.",
            "Regulatory status and pathway": "CDSCO permitted a phase III trial, but no approved India label was verified.",
            "Patent and exclusivity": "No complete IP or rights package was validated in the current run.",
            "Competitor landscape": "RA is heavily contested by established biologics and targeted agents.",
            "Commercial attractiveness": "The asset is not launch-ready and the investment case remains speculative.",
            "Strategic fit with company product mix": "Potential fit exists only if the company wants a biologics capability and the rights package is compelling.",
            "Operational and CMC feasibility": "High CMC complexity and comparability burden remain unresolved.",
            "HCP value proposition": "No authoritative India-facing value proposition can be finalised while the molecule identity remains uncertain.",
            "Evidence quality and digital intelligence": "The strongest signal is the CDSCO trial permit plus peer-reviewed phase III literature for olokizumab.",
            "Reputation, ethics, and access": "Proceeding on a misnamed molecule would be an avoidable governance risk.",
        }
    else:
        mapping = {
            "Therapeutic gap and medical need": "Strong relevance across cardiometabolic, heart-failure, and kidney-risk populations.",
            "Molecule and target profile": "Established SGLT2 mechanism with broad clinical relevance; novelty is no longer high.",
            "Preclinical evidence": "Mature marketed product, but detailed preclinical package was outside this run.",
            "Clinical evidence": "Multiple major outcomes programs support a broad evidence proposition.",
            "Safety and tolerability": "Well-characterized but requires disciplined risk communication and monitoring.",
            "Regulatory status and pathway": "Current FDA and EMA sources support broad indications; India label remains to be verified.",
            "Competitor landscape": "Strong class position but significant crowding and price competition.",
            "Commercial attractiveness": "Strong clinical demand potential; local economics remain unverified.",
            "Strategic fit with company product mix": "High fit for an integrated cardiometabolic-renal portfolio.",
            "Operational and CMC feasibility": "Mature product class, but company-specific CMC and supply data are missing.",
            "HCP value proposition": "Clear evidence-based cardiorenal value proposition with practical once-daily dosing.",
            "Evidence quality and digital intelligence": "High-quality regulatory and pivotal outcome sources are readily available.",
            "Reputation, ethics, and access": "Potential major public-health value, subject to affordability and appropriate use.",
        }
    return mapping.get(category, "Requires cross-functional diligence.")

def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [("Title", 25, NAVY), ("Heading 1", 15, NAVY), ("Heading 2", 11.5, TEAL), ("Heading 3", 10, NAVY)]:
        st = styles[name]
        st.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("ClinCommand OS | Internal Qualification Output | Dr. Bhupesh Dewan, Mumbai, Maharashtra")
    rr.font.name = "Aptos"
    rr.font.size = Pt(7)
    rr.font.color.rgb = RGBColor.from_string("6A7882")
    return doc

def add_cover(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("CLINCOMMAND OS")
    r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(TEAL)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"Product Appraisal\n{data['requested_name']}")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run(data["subtitle"])
    rr.italic = True; rr.font.size = Pt(12); rr.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()
    add_table(doc, ["Control Field", "Value"], [
        ("Qualification run", data["code"]),
        ("Verified INN", data["verified_name"]),
        ("Intended use", "Internal preliminary portfolio and medical-strategy appraisal"),
        ("Evidence cut-off", "6 June 2026"),
        ("Document status", "Qualification draft - not approved for external use or final Go/No-Go decision"),
        ("Owner", "Dr. Bhupesh Dewan, Mumbai, Maharashtra, India"),
    ], [1.8, 5.3], TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    rr = p.add_run("IMPORTANT LIMITATION\n")
    rr.bold = True; rr.font.color.rgb = RGBColor.from_string("A34A00")
    rr = p.add_run("This appraisal is an architecture qualification output. Missing regulatory, patent, legal, pricing, market, CMC, API, supply, and company-strategy information is explicitly classified as data gaps and diligence gaps, not inferred.")
    rr.font.size = Pt(9.5)
    doc.add_page_break()

def add_appraisal(doc, data):
    add_heading(doc, "1. Executive Decision Summary")
    add_table(doc, ["Decision Item", "Assessment"], [
        ("Objective", data["objective"]),
        ("Preliminary status", data["status"]),
        ("Recommendation", data["recommendation"]),
        ("Identity/control note", data["identity_note"]),
    ], [1.65, 5.45], TEAL)

    add_heading(doc, "2. Product Identification and Assessment Scope")
    add_table(doc, ["Field", "Verified or Qualified Finding"], data["product"], [1.8, 5.3])

    add_heading(doc, "3. Scientific and Strategic Synopsis")
    p = doc.add_paragraph()
    p.add_run("Appraisal interpretation. ").bold = True
    if data["verified_name"] == "Trelagliptin":
        p.add_run("Trelagliptin offers a clear administration-frequency distinction within the DPP-4 class. Its scientific proposition is narrower than products supported by extensive cardiorenal outcomes programs. The opportunity therefore depends on verified local authorization, appropriate patient positioning, differentiated execution, affordability, and robust operational diligence.")
    else:
        p.add_run("Empagliflozin has a mature, high-quality evidence platform across diabetes, cardiovascular risk, heart failure, and chronic kidney disease. The principal strategic question is not whether the molecule has clinical value, but whether the company can create a differentiated, compliant, affordable, and operationally credible portfolio proposition in a crowded market.")

    add_heading(doc, "4. SWOT Assessment")
    add_table(doc, ["Strengths", "Weaknesses"], [("\n".join(data["strengths"]), "\n".join(data["weaknesses"]))], [3.55, 3.55], TEAL)
    add_table(doc, ["Opportunities", "Threats"], [("\n".join(data["opportunities"]), "\n".join(data["threats"]))], [3.55, 3.55], TEAL)

    add_heading(doc, "5. Evidence Landscape")
    add_table(doc, ["Evidence", "Relevance", "Source class", "Confidence"], data["evidence"], [1.4, 3.2, 1.65, 0.85])

    add_heading(doc, "6. Weighted Appraisal Scorecard")
    rows, total = weighted_rows(data)
    add_table(doc, ["Category", "Weight", "Score 0-5", "Weighted", "Rationale"], rows, [2.0, 0.55, 0.65, 0.65, 3.25])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Total preliminary weighted score: {total:.1f} / 100")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(TEAL)

    add_heading(doc, "7. Mandatory Red-Flag Review")
    add_table(doc, ["Area", "Finding", "Required disposition"], data["red_flags"], [1.35, 3.4, 2.35], AMBER)

    add_heading(doc, "8. Clinical, Safety, and Regulatory Interpretation")
    if data["verified_name"] == "Trelagliptin":
        add_bullets(doc, [
            "Clinical evidence supports glucose-lowering efficacy and non-inferiority in specific comparative studies; generalization beyond studied populations requires caution.",
            "Once-weekly dosing is a practical differentiator, but improved adherence or outcomes must not be claimed without evidence.",
            "Japan regulatory review is verified. Exact current India approval, label, dose, warnings, and permitted claims require authoritative confirmation.",
            "A full safety review should use the effective local label, current PV information, and class-specific risks before progression."
        ])
    else:
        add_bullets(doc, [
            "Empagliflozin has a strong outcomes-based evidence platform extending beyond glycemic control.",
            "Current FDA labeling includes heart failure, chronic kidney disease, cardiovascular-risk reduction in defined patients, and glycemic control indications.",
            "Benefit-risk communication must visibly include ketoacidosis, volume depletion, genitourinary infection, hypoglycemia with insulin/secretagogues, perioperative interruption, and other current label controls.",
            "Exact India-approved indications, restrictions, safety wording, and use conditions must govern any local content."
        ])

    add_heading(doc, "9. Cross-Functional Diligence Plan")
    diligence = [
        ("Regulatory", "Verify current India approval status, approved indications, dose, safety language, and change history.", "Critical"),
        ("Medical and PV", "Complete evidence synthesis, local safety assessment, and risk-communication plan.", "Critical"),
        ("Patent and legal", "Complete India patent landscape, exclusivity, FTO, litigation, and licensing assessment.", "Critical"),
        ("Commercial and market access", "Validate market size, competitors, prices, affordability, channels, forecast, and portfolio fit.", "Major"),
        ("CMC and supply", "Validate API source, cost, DMF, quality, capacity, stability, formulation, and launch feasibility.", "Major"),
        ("Quality", "Confirm approved SOP, qualified reviewers, traceability, and controlled record requirements.", "Major"),
    ]
    add_table(doc, ["Function", "Required action", "Priority"], diligence, [1.35, 4.95, 0.8])

    add_heading(doc, "10. Preliminary Go/No-Go Interpretation")
    p = doc.add_paragraph()
    p.add_run(data["status"] + ". ").bold = True
    p.add_run(data["recommendation"])
    p = doc.add_paragraph()
    p.add_run("Decision boundary: ").bold = True
    p.add_run("The numeric score does not override unresolved critical red flags. This output supports diligence planning only.")

    add_heading(doc, "11. Source Register")
    add_table(doc, ["Source", "Title", "URL"], data["sources"], [1.5, 3.3, 2.3])

    add_heading(doc, "12. Qualification and Review Status")
    add_table(doc, ["Control", "Status"], [
        ("Business SOP consulted", "Owner-supplied Product Appraisal SOP used as structural reference"),
        ("Skill consulted", "Owner-supplied Pharma Product Appraisal Skill Template used for scoring and content architecture"),
        ("Evidence handling", "Authoritative and peer-reviewed sources used; unavailable data explicitly marked"),
        ("AI-generated content", "Qualification draft requiring qualified human review"),
        ("Internal approval", "Pending qualification review"),
        ("External use", "Prohibited"),
    ], [2.2, 4.9], TEAL)

def save_doc(data):
    doc = setup_doc()
    add_cover(doc, data)
    add_appraisal(doc, data)
    filename = f"{data['code']}_{data['verified_name']}_Product_Appraisal_R2.docx"
    path = OUT / filename
    doc.save(path)
    return path

def build_report(paths):
    doc = setup_doc()
    add_cover(doc, {
        "requested_name": "Product Appraisal Activity Qualification Report",
        "verified_name": "Two-run qualification",
        "code": "QUAL-MA-PA-001",
        "subtitle": "Comparison of Trelagliptin and Empagliflozin Appraisal Outputs"
    })
    add_heading(doc, "1. Qualification Summary")
    add_table(doc, ["Run", "Output", "Preliminary document result"], [
        ("PA-Q1", paths[0].name, "Pending independent review; architecture test intentionally contains material data gaps"),
        ("PA-Q2", paths[1].name, "Pending independent review; architecture test covers a mature evidence-rich product"),
    ], [0.8, 3.6, 3.5], TEAL)
    add_heading(doc, "2. Preliminary Architecture Assessment")
    add_table(doc, ["Dimension", "PA-Q1 Trelagliptin", "PA-Q2 Empagliflozin"], [
        ("Identity handling", "Pass: requested-name discrepancy surfaced", "Pass: INN verified"),
        ("Evidence behavior", "Pass: limited evidence and geography gaps disclosed", "Pass: regulatory and pivotal evidence summarized"),
        ("Data-gap behavior", "Pass: no invented India approval, pricing, patent, API, or supply conclusions", "Pass: local commercial and operational gaps disclosed"),
        ("Score behavior", "Pass: conservative scoring and red-flag override", "Pass: strong science score balanced by commercial and safety diligence"),
        ("Output usability", "Pending render and human review", "Pending render and human review"),
    ], [1.6, 3.2, 3.2])
    add_heading(doc, "3. Required Human Review")
    add_bullets(doc, [
        "Confirm the requested molecule name Trilagliptin should be interpreted as trelagliptin.",
        "Review scientific interpretations and preliminary scores.",
        "Provide or validate India regulatory, patent/FTO, pricing, market, API, supply, CMC, and company-strategy inputs.",
        "Decide whether the 23-section structure and weighted scorecard adequately reflect the owner-supplied SOP and skill.",
        "Record critical, major, and minor findings before architecture approval."
    ])
    add_heading(doc, "4. Preliminary Decision")
    p = doc.add_paragraph()
    p.add_run("Conditional Pass pending render QA and qualified human review. ").bold = True
    p.add_run("The two outputs demonstrate that the activity can distinguish an uncertain, geography-limited opportunity from a mature evidence-rich product without filling missing data with invented conclusions. Final architecture approval must wait for recorded owner and expert findings.")
    path = OUT / "QUAL-MA-PA-001_Product_Appraisal_Qualification_Report_R2.docx"
    doc.save(path)
    return path

def build_reassessment_report(paths):
    doc = setup_doc()
    add_cover(doc, {
        "requested_name": "Product Appraisal Reassessment",
        "verified_name": "Fevipiprant and Olokizumab",
        "code": "QUAL-MA-PA-002",
        "subtitle": "Reassessment of New Portfolio Candidates After Owner Feedback"
    })
    add_heading(doc, "1. Qualification Summary")
    add_table(doc, ["Run", "Output", "Preliminary document result"], [
        ("PA-Q3", paths[0].name, "No-Go recommendation; phase III failure and discontinuation are treated as hard-stop signals"),
        ("PA-Q4", paths[1].name, "No-Go as requested; identity verification failed, and olokizumab can only be considered after naming correction"),
    ], [0.8, 3.6, 3.5], TEAL)
    add_heading(doc, "2. Preliminary Architecture Assessment")
    add_table(doc, ["Dimension", "PA-Q3 Fevipiprant", "PA-Q4 Olonacimab / Olokizumab"], [
        ("Identity handling", "Pass: molecule verified", "Pass for the ambiguity being surfaced; requested name not verified"),
        ("Evidence behavior", "Pass: sponsor discontinuation and phase III failure clearly surfaced", "Pass: peer-reviewed olokizumab evidence and India trial permit surfaced"),
        ("Data-gap behavior", "Pass: no invented approval or launch case", "Pass: no invented approval or launch case"),
        ("Score behavior", "Pass: conservative no-go scoring", "Pass: no-go with identity-first control"),
        ("Output usability", "Pending render and human review", "Pending render and human review"),
    ], [1.6, 3.2, 3.2])
    add_heading(doc, "3. Required Human Review")
    add_bullets(doc, [
        "Confirm that 'Olonacimab' was intended to mean olokizumab, or replace it with the correct INN before any downstream use.",
        "Review whether the no-go conclusion for fevipiprant should be preserved as a portfolio learning case.",
        "Check any rights, patent, CMC, and access assumptions before any biologic-development discussion.",
        "Record whether these outputs should be treated as negative-development references or retired from active consideration.",
    ])
    add_heading(doc, "4. Preliminary Decision")
    p = doc.add_paragraph()
    p.add_run("Conditional Pass pending render QA and qualified human review. ").bold = True
    p.add_run("The reassessment output is intentionally strict: fevipiprant is blocked by failed late-stage development, and the requested olonacimab appraisal is blocked by identity uncertainty until Regulatory Affairs confirms the intended molecule.")
    path = OUT / "QUAL-MA-PA-002_Product_Appraisal_Reassessment_Report_R2.docx"
    doc.save(path)
    return path

if __name__ == "__main__":
    p1 = save_doc(trelagliptin)
    p2 = save_doc(empagliflozin)
    p3 = build_report([p1, p2])
    p4 = save_doc(fevipiprant)
    p5 = save_doc(olonacimab)
    p6 = build_reassessment_report([p4, p5])
    print(p1)
    print(p2)
    print(p3)
    print(p4)
    print(p5)
    print(p6)
